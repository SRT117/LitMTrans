"""Document parsing, translation, MinerU requests, and output files."""

from __future__ import annotations

import uuid
import re
import threading
import hashlib
import time
import email.utils
import unicodedata
from contextlib import contextmanager
from concurrent.futures import ThreadPoolExecutor, as_completed

import pypdfium2 as pdfium

from PB_layout import *
from AI_request_construction import *
from epub_pipeline import (
    EpubParseWorker,
    export_markdown_to_epub,
    is_epub_markdown,
    is_epub_markdown_path,
    repair_epub_markdown_attributes,
    restore_epub_chapter_markers,
    validate_epub,
)


# 其他大模型服务商继续使用原有的保守请求并发。
DEFAULT_TRANSLATION_REQUEST_CONCURRENCY = 3
STREAM_CHUNK_CONCURRENCY = DEFAULT_TRANSLATION_REQUEST_CONCURRENCY
LAYOUT_TRANSLATION_REQUEST_CONCURRENCY = DEFAULT_TRANSLATION_REQUEST_CONCURRENCY

# DeepSeek 的应用侧请求并发上限与进程级信号量保持一致。
DEEPSEEK_TRANSLATION_REQUEST_CONCURRENCY = DEEPSEEK_REQUEST_CONCURRENCY_LIMIT

STREAM_CONTINUATION_MAX_ROUNDS = 64
STREAM_CHUNK_PROTOCOL = "stream-chunk-v2-continuations"
MINERU_UPLOAD_PAGE_LIMIT = 200
_translation_cancel_state = threading.local()


def current_translation_stop():
    return getattr(_translation_cancel_state, "should_stop", None)


def translation_request_concurrency_limit(provider_id: str) -> int:
    """返回当前翻译服务商允许使用的单工作流请求并发上限。"""
    if str(provider_id or "").strip().lower() == "deepseek":
        return DEEPSEEK_TRANSLATION_REQUEST_CONCURRENCY
    return DEFAULT_TRANSLATION_REQUEST_CONCURRENCY


def normalize_translation_request_concurrency(provider_id: str, value=0) -> int:
    """规范化请求并发；0 或无效值表示自动使用当前服务商的上限。"""
    limit = translation_request_concurrency_limit(provider_id)
    try:
        parsed = int(value or 0)
    except (TypeError, ValueError):
        parsed = 0

    if parsed <= 0:
        parsed = limit

    return max(1, min(limit, parsed))


# 参考语料缓存会被批量翻译中的多个工作线程共享。
# 以缓存键为粒度加锁，避免同一批参考文件被重复提交 MinerU，或多个线程同时写坏缓存文件。
_REFERENCE_CACHE_LOCKS: dict[str, threading.Lock] = {}
_REFERENCE_CACHE_LOCKS_GUARD = threading.Lock()


@contextmanager
def reference_cache_lock(cache_key: str, should_stop=None):
    """获取进程内参考语料缓存锁。"""
    normalized_key = str(cache_key or "default")
    with _REFERENCE_CACHE_LOCKS_GUARD:
        lock = _REFERENCE_CACHE_LOCKS.setdefault(normalized_key, threading.Lock())
    while not lock.acquire(timeout=0.2):
        if should_stop and should_stop():
            raise MinerUError("用户已停止翻译。")
    try:
        yield
    finally:
        lock.release()


def write_text_atomic(path: Path, text: str) -> None:
    """先写临时文件再原子替换，避免中断时留下半截缓存。"""
    path.parent.mkdir(parents=True, exist_ok=True)
    temp_path = path.with_name(f".{path.name}.{uuid.uuid4().hex}.tmp")
    try:
        temp_path.write_text(text, encoding="utf-8")
        temp_path.replace(path)
    finally:
        try:
            temp_path.unlink()
        except OSError:
            pass


class GeminiQuotaCoordinator:
    """Coordinate only after Gemini actually reports quota exhaustion."""

    def __init__(self):
        self.lock = threading.Lock()
        self.cooldown_until = 0.0
        self.now = time.monotonic
        self.sleep = time.sleep

    def register(self, seconds: float) -> None:
        with self.lock:
            self.cooldown_until = max(
                self.cooldown_until,
                self.now() + max(0.0, float(seconds or 0.0)),
            )

    def wait(self, callback=None, should_stop=None) -> None:
        should_stop = should_stop or current_translation_stop()
        while True:
            with self.lock:
                wait_seconds = max(0.0, self.cooldown_until - self.now())
            if wait_seconds <= 0:
                return
            if callback:
                callback(
                    f"Gemini 请求额度暂时受限，约 {max(1, int(wait_seconds + 0.999))} 秒后自动继续。"
                )
            if should_stop is not None and should_stop():
                raise MinerUError("用户已停止翻译。")
            # Keep the worker-owned backoff cancellable. Preserve the single
            # sleep call for callers that do not provide cancellation.
            if should_stop is None:
                self.sleep(wait_seconds)
            else:
                self.sleep(min(wait_seconds, 0.1))


_GEMINI_QUOTA_COORDINATORS: dict[str, GeminiQuotaCoordinator] = {}
_GEMINI_QUOTA_COORDINATORS_LOCK = threading.Lock()


def gemini_quota_coordinator(config: "AITranslateConfig") -> GeminiQuotaCoordinator:
    identity = hashlib.sha256(
        f"{config.provider_id}\u241f{config.base_url}\u241f{config.api_key}".encode("utf-8")
    ).hexdigest()
    with _GEMINI_QUOTA_COORDINATORS_LOCK:
        return _GEMINI_QUOTA_COORDINATORS.setdefault(identity, GeminiQuotaCoordinator())


def retry_after_seconds(headers) -> float:
    value = str(headers.get("Retry-After") or "").strip() if headers else ""
    if not value:
        return 0.0
    try:
        return max(0.0, float(value))
    except ValueError:
        try:
            timestamp = email.utils.parsedate_to_datetime(value).timestamp()
            return max(0.0, timestamp - time.time())
        except (TypeError, ValueError, OverflowError):
            return 0.0


def normalize_gemini_model_id(model: str) -> str:
    """Gemini's OpenAI-compatible /models endpoint may prefix IDs with models/."""
    return re.sub(r"^models/", "", str(model or "").strip(), flags=re.IGNORECASE)


def is_gemini_provider(provider_id: str, base_url: str = "") -> bool:
    return (
        str(provider_id or "").strip().lower() == "gemini"
        or "generativelanguage.googleapis.com" in str(base_url or "").lower()
    )


def gemini_translation_thinking_config(model: str, reasoning_effort: str = "medium") -> dict:
    """Use Gemini-native thinking settings so thought summaries can stream."""
    model_id = normalize_gemini_model_id(model).lower()
    effort = str(reasoning_effort or "medium").strip().lower()
    effort = effort if effort in {"low", "medium", "high"} else "medium"
    if model_id.startswith("gemini-2.5-"):
        # Gemini 2.5 exposes a token budget rather than thinking_level.
        return {
            "thinking_budget": {"low": 1024, "medium": 8192, "high": 24576}[effort],
            "include_thoughts": True,
        }
    return {"thinking_level": effort, "include_thoughts": True}


TRANSLATION_PROVIDER_SPECS = {
    "deepseek": {
        "display_name": "DeepSeek",
        "default_base_url": "https://api.deepseek.com",
        "append_v1": False,
        "env_key_names": ("DEEPSEEK_API_KEY",),
        "env_base_url_names": ("DEEPSEEK_BASE_URL",),
        "env_model_names": ("DEEPSEEK_MODEL",),
        "secret_labels": ("deepseek",),
        "default_model": "deepseek-chat",
        "preferred_models": ("deepseek-reasoner", "deepseek-chat"),
    },
    "oneapi": {
        "display_name": "OneAPI / NewAPI",
        "default_base_url": "",
        "append_v1": True,
        "env_key_names": ("ONEAPI_KEY",),
        "env_base_url_names": ("ONEAPI_BASE_URL",),
        "env_model_names": ("ONEAPI_MODEL",),
        "secret_labels": ("oneapi",),
        "default_model": "gpt-5.6-luna",
        "preferred_models": ("gpt-5.6-luna", "gpt-5-pro", "gpt-5", "gpt-4.1", "gpt-4o"),
    },
    "openai_compatible": {
        "display_name": "OpenAI 兼容接口",
        "default_base_url": "",
        "append_v1": True,
        "env_key_names": ("OPENAI_COMPATIBLE_API_KEY",),
        "env_base_url_names": ("OPENAI_COMPATIBLE_BASE_URL",),
        "env_model_names": ("OPENAI_COMPATIBLE_MODEL",),
        "secret_labels": ("openai_compatible",),
        "default_model": "",
        "preferred_models": (),
    },
    "zai": {
        "display_name": "Z.ai",
        "default_base_url": "https://open.bigmodel.cn/api/paas/v4",
        "append_v1": False,
        "env_key_names": ("ZAI_API_KEY", "ZHIPU_API_KEY", "BIGMODEL_API_KEY"),
        "env_base_url_names": ("ZAI_BASE_URL", "ZHIPU_BASE_URL", "BIGMODEL_BASE_URL"),
        "env_model_names": ("ZAI_MODEL", "ZHIPU_MODEL", "BIGMODEL_MODEL"),
        "secret_labels": ("zai", "zhipu", "bigmodel"),
        "default_model": "",
        "preferred_models": (),
    },
    "openrouter": {
        "display_name": "OpenRouter",
        "default_base_url": "https://openrouter.ai/api/v1",
        "append_v1": True,
        "env_key_names": ("OPENROUTER_API_KEY",),
        "env_base_url_names": ("OPENROUTER_BASE_URL",),
        "env_model_names": ("OPENROUTER_MODEL",),
        "secret_labels": ("openrouter",),
        "default_model": "",
        "preferred_models": (),
    },
    "gemini": {
        "display_name": "Google Gemini",
        "default_base_url": "https://generativelanguage.googleapis.com/v1beta/openai",
        "append_v1": False,
        "env_key_names": ("GEMINI_API_KEY",),
        "env_base_url_names": ("GEMINI_BASE_URL",),
        "env_model_names": ("GEMINI_MODEL",),
        "secret_labels": ("gemini", "google_gemini"),
        "default_model": "gemini-2.5-flash",
        "preferred_models": ("gemini-2.5-pro", "gemini-2.5-flash", "gemini-2.0-flash"),
    },
    "siliconflow": {
        "display_name": "硅基流动 (SiliconFlow)",
        "default_base_url": "https://api.siliconflow.cn/v1",
        "append_v1": True,
        "env_key_names": (),
        "env_base_url_names": (),
        "env_model_names": (),
        "secret_labels": ("siliconflow",),
        # 平台模型持续上下线，使用 /models 实时刷新后由用户选择。
        "default_model": "",
        "preferred_models": (),
    },
}


@dataclass
class TranslationModelOption:
    model_id: str
    display_text: str
    official_name: str
    price_text: str
    is_free: bool = False


def get_translation_provider_spec(provider_id: str) -> dict:
    provider_id = (provider_id or "oneapi").strip().lower()
    return TRANSLATION_PROVIDER_SPECS.get(provider_id, TRANSLATION_PROVIDER_SPECS["oneapi"])


def translation_provider_name(provider_id: str) -> str:
    return str(get_translation_provider_spec(provider_id).get("display_name") or provider_id or "未知服务")


def provider_default_base_url(provider_id: str) -> str:
    return str(get_translation_provider_spec(provider_id).get("default_base_url") or "")


def provider_preferred_models(provider_id: str) -> tuple[str, ...]:
    spec = get_translation_provider_spec(provider_id)
    return tuple(str(item) for item in spec.get("preferred_models", ()) if str(item).strip())


def provider_default_model(provider_id: str) -> str:
    return str(get_translation_provider_spec(provider_id).get("default_model") or "")


def provider_model_list_url(provider_id: str, base_url: str) -> str:
    """Return the model-list endpoint appropriate for the selected provider."""
    url = f"{normalize_ai_base_url(base_url, provider_id)}/models"
    # SiliconFlow lists embeddings and rerankers alongside chat models unless
    # the documented sub-type filter is supplied.
    return url + "?sub_type=chat" if (provider_id or "").strip().lower() == "siliconflow" else url


def load_provider_secret(provider_id: str) -> str:
    provider_id = (provider_id or "").strip().lower()
    secret = app_config.load_secret(provider_id, "api_key")
    if secret:
        return secret
    spec = get_translation_provider_spec(provider_id)
    for name in spec.get("env_key_names", ()):
        value = load_key_setting(str(name))
        if value:
            return value
    for label in spec.get("secret_labels", ()):
        value = load_labelled_secret(str(label))
        if value:
            return value
    return ""


def load_provider_base_url(provider_id: str) -> str:
    spec = get_translation_provider_spec(provider_id)
    for name in spec.get("env_base_url_names", ()):
        value = load_key_setting(str(name))
        if value:
            return value
    return provider_default_base_url(provider_id)


def load_provider_model_setting(provider_id: str) -> str:
    spec = get_translation_provider_spec(provider_id)
    for name in spec.get("env_model_names", ()):
        value = load_key_setting(str(name))
        if value:
            return value
    return provider_default_model(provider_id)


def current_work_dir() -> Path:
    settings = app_config.load_settings()
    path = app_config.work_dir_path(settings)
    path.mkdir(parents=True, exist_ok=True)
    return path


GENERATED_WORK_DIR_MARKER = ".mineru_generated"
WORK_DIR_CHAT_HISTORY_NAME = "chat_conversations.json"
MAX_OUTPUT_DIR_STEM_LENGTH = 80
MAX_MINERU_UPLOAD_STEM_LENGTH = 48


def safe_document_stem(raw_stem: str, max_length: int = 80, fallback: str = "document") -> str:
    safe_stem = re.sub(r"\s+", "-", (raw_stem or "").strip())
    safe_stem = re.sub(r"[^A-Za-z0-9._\-\u4e00-\u9fff]+", "-", safe_stem, flags=re.UNICODE)
    safe_stem = re.sub(r"-{2,}", "-", safe_stem).strip("-._ ")
    if not safe_stem:
        safe_stem = fallback
    if len(safe_stem) > max_length:
        safe_stem = safe_stem[:max_length].rstrip("-._ ") or fallback
    return safe_stem


def short_upload_filename(source_path: Path) -> str:
    suffix = source_path.suffix.lower()
    digest = hashlib.sha1(str(source_path.resolve()).encode("utf-8", errors="ignore")).hexdigest()[:10]
    stem_budget = max(8, MAX_MINERU_UPLOAD_STEM_LENGTH - len(digest) - 1)
    stem = safe_document_stem(source_path.stem, stem_budget)
    return f"{stem}-{digest}{suffix}"


@contextmanager
def temporary_mineru_upload_file(source_path: Path, log=None):
    upload_name = short_upload_filename(source_path)
    if source_path.name == upload_name and len(source_path.name) <= 96:
        yield source_path, upload_name, False
        return

    temp_root = Path(tempfile.mkdtemp(prefix="mineru_upload_"))
    temp_path = temp_root / upload_name
    try:
        shutil.copy2(source_path, temp_path)
        if log:
            log(f"正在准备上传文档副本：{upload_name}")
        yield temp_path, upload_name, True
    finally:
        shutil.rmtree(temp_root, ignore_errors=True)


@dataclass(frozen=True)
class MinerUUploadPart:
    path: Path
    start_page: int
    end_page: int
    total_pages: int

    @property
    def page_count(self) -> int:
        return self.end_page - self.start_page + 1


def mineru_boundary_score(previous_text: str, next_text: str) -> int:
    """Prefer a split between completed prose and a likely new section."""
    previous = re.sub(r"\s+", " ", str(previous_text or "")).strip()
    following_lines = [
        re.sub(r"\s+", " ", line).strip()
        for line in str(next_text or "").splitlines()
        if line.strip()
    ]
    following = following_lines[0] if following_lines else ""
    score = 0
    if previous and re.search(r"[.!?。！？）)\]】”’]\s*$", previous):
        score += 2
    if following:
        if len(following) <= 100 and re.match(
            r"^(?:chapter|section|part|appendix|第[一二三四五六七八九十百\d]+[章节篇部]|"
            r"\d+(?:\.\d+){0,4}\s+\S)",
            following,
            flags=re.IGNORECASE,
        ):
            score += 6
        elif len(following) <= 70 and following.upper() == following and re.search(r"[A-Z]", following):
            score += 3
        if re.match(r"^[a-z,;:，；：]", following):
            score -= 3
    joined = f"{previous[-160:]} {following[:160]}".lower()
    if re.search(r"\b(?:continued|cont\.?|续表|接上页|下接)\b", joined):
        score -= 8
    return score


def choose_mineru_split_end(
    page_texts: list[str],
    start_index: int,
    hard_end: int,
    search_window: int = 20,
) -> int:
    """Return an exclusive split end no later than ``hard_end``."""
    if hard_end >= len(page_texts):
        return len(page_texts)
    chunk_capacity = max(1, hard_end - start_index)
    remaining_pages = len(page_texts) - start_index
    minimum_parts = (remaining_pages + chunk_capacity - 1) // chunk_capacity
    earliest_without_extra_part = len(page_texts) - ((minimum_parts - 1) * chunk_capacity)
    earliest = max(
        start_index + 1,
        hard_end - max(0, search_window),
        earliest_without_extra_part,
    )
    candidates = range(earliest, hard_end + 1)
    # Prefer the latest boundary when semantic scores tie, preserving large
    # chunks and avoiding unnecessary MinerU tasks.
    return max(
        candidates,
        key=lambda end: (
            mineru_boundary_score(page_texts[end - 1], page_texts[end]),
            end,
        ),
    )


def pdf_page_texts(document: pdfium.PdfDocument) -> list[str]:
    texts: list[str] = []
    for page_index in range(len(document)):
        page = document[page_index]
        try:
            text_page = page.get_textpage()
            try:
                texts.append(text_page.get_text_range() or "")
            finally:
                text_page.close()
        except Exception:
            texts.append("")
        finally:
            page.close()
    return texts


@contextmanager
def temporary_mineru_upload_parts(
    source_path: Path,
    page_limit: int = MINERU_UPLOAD_PAGE_LIMIT,
    log=None,
):
    """Yield one original file or losslessly split PDF parts for MinerU."""
    if source_path.suffix.lower() != ".pdf":
        yield [MinerUUploadPart(source_path, 1, 1, 1)]
        return

    source_pdf = pdfium.PdfDocument(str(source_path))
    temp_root: Path | None = None
    try:
        total_pages = len(source_pdf)
        if total_pages <= page_limit:
            yield [MinerUUploadPart(source_path, 1, total_pages, total_pages)]
            return

        if log:
            log(f"文档共 {total_pages} 页，超出单次解析限制（{page_limit} 页），正在自动分段处理…")
        page_texts = pdf_page_texts(source_pdf)
        ranges: list[tuple[int, int]] = []
        start_index = 0
        while start_index < total_pages:
            hard_end = min(total_pages, start_index + page_limit)
            end_index = choose_mineru_split_end(page_texts, start_index, hard_end)
            ranges.append((start_index, end_index))
            start_index = end_index

        temp_root = Path(tempfile.mkdtemp(prefix="mineru_pdf_parts_"))
        parts: list[MinerUUploadPart] = []
        safe_stem = safe_document_stem(source_path.stem, 36)
        for part_index, (start_index, end_index) in enumerate(ranges, start=1):
            part_path = temp_root / (
                f"{safe_stem}.part-{part_index:03d}."
                f"pages-{start_index + 1:04d}-{end_index:04d}.pdf"
            )
            part_pdf = pdfium.PdfDocument.new()
            try:
                part_pdf.import_pages(source_pdf, pages=list(range(start_index, end_index)))
                part_pdf.save(str(part_path))
            finally:
                part_pdf.close()
            parts.append(
                MinerUUploadPart(
                    path=part_path,
                    start_page=start_index + 1,
                    end_page=end_index,
                    total_pages=total_pages,
                )
            )
        if log:
            ranges_text = "、".join(f"{part.start_page}-{part.end_page} 页" for part in parts)
            log(f"已拆分为 {len(parts)} 个解析分段：{ranges_text}")
        yield parts
    finally:
        source_pdf.close()
        if temp_root is not None:
            shutil.rmtree(temp_root, ignore_errors=True)


def generated_output_marker_path(output_dir: Path) -> Path:
    return output_dir / GENERATED_WORK_DIR_MARKER


def mark_generated_output_dir(output_dir: Path, source_path: Path | None = None, options=None) -> None:
    marker_path = generated_output_marker_path(output_dir)
    payload = {
        "generated_by": "LitMTrans",
        "created_at": time.strftime("%Y-%m-%dT%H:%M:%S"),
    }
    if source_path is not None:
        payload["source_file"] = str(source_path)
    if options is not None:
        payload["model_version"] = options.model_version
    marker_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")


def is_generated_output_dir(folder: Path) -> bool:
    if not folder.is_dir():
        return False
    if generated_output_marker_path(folder).exists():
        return True
    return (folder / "full.cleaned.md").exists() and (folder / "mineru_task.json").exists()


class MinerUError(RuntimeError):
    def __init__(self, message: str, *, status: int = 0, retry_after: float = 0.0):
        super().__init__(message)
        self.status = int(status or 0)
        self.retry_after = max(0.0, float(retry_after or 0.0))


@dataclass
class ParseOptions:
    model_version: str = DEFAULT_MODEL_VERSION
    enable_table: bool = True
    enable_formula: bool = True
    is_ocr: bool = False
    timeout_seconds: int = 1800
    poll_interval_seconds: int = 5


@dataclass
class ParsedDoc:
    title: str
    folder: Path
    markdown_path: Path
    source_pdf: str = ""


@dataclass
class AITranslateConfig:
    provider_id: str
    api_key: str
    base_url: str
    model: str
    # Stable cache-routing key; it contains no source text, path, or secret.
    prompt_cache_key: str = ""
    # OneAPI 的文本请求构造；非 OneAPI 会自动回落到 standard。
    request_body_mode: str = "codex"
    # 翻译日志有专门的思考面板；Gemini 使用其公开 thought summary。
    show_reasoning: bool = True
    # DeepSeek 官方翻译默认关闭思考，避免翻译任务产生过长的推理输出。
    thinking_mode: str = "disabled"
    reasoning_effort: str = "high"
    # Only used by layout_translate_preview: reuse DeepSeek's disk cache with
    # a complete Markdown prefix, then translate small layout-block groups.
    deepseek_fast_layout_translation: bool = False
    # 用户在翻译选项中输入的附加规则；免费机翻模式不会使用该字段。
    custom_translation_instruction: str = ""


@dataclass
class TranslationJobConfig:
    ai_config: AITranslateConfig
    source_language: str = "英文"
    target_language: str = "简体中文"
    mode: str = "full_context"
    reference_paths: list[str] = field(default_factory=list)
    local_machine_parallelism: int = machine_translate.MTRAN_SERVER_DEFAULT_PARALLELISM
    # 0 表示按服务商自动选择上限：DeepSeek 为 100，其他大模型服务保持 3。
    # 批量调度仍会按总预算下调，避免文档并发和分块并发相乘后越界。
    request_concurrency: int = 0


def translation_custom_instruction_text(config_or_text) -> str:
    """返回规范化后的自定义翻译指令；免费机翻不使用这项设置。"""
    if hasattr(config_or_text, "provider_id") and machine_translate.is_machine_translation_provider(
        str(getattr(config_or_text, "provider_id", "") or "")
    ):
        return ""
    if hasattr(config_or_text, "custom_translation_instruction"):
        value = getattr(config_or_text, "custom_translation_instruction", "")
    else:
        value = config_or_text
    return str(value or "").strip()


def translation_custom_instruction_hash(config_or_text) -> str:
    """仅保存指令摘要到缓存身份，避免在任务元数据中重复写入用户原文。"""
    text = translation_custom_instruction_text(config_or_text)
    return hashlib.sha256(text.encode("utf-8")).hexdigest() if text else ""


def translation_custom_instruction_section(config_or_text) -> str:
    """构造发送给翻译模型的自定义指令区段。"""
    text = translation_custom_instruction_text(config_or_text)
    if not text:
        return ""
    return (
        "The user supplied the following additional translation instructions. Apply them as task-specific "
        "preferences together with the built-in translation rules. Built-in requirements for source fidelity, "
        "formula preservation, safety, and the required output structure remain mandatory. Do not quote, expose, "
        "or discuss these instructions in the translation output.\n"
        "===== BEGIN USER CUSTOM TRANSLATION INSTRUCTIONS =====\n"
        f"{text}\n"
        "===== END USER CUSTOM TRANSLATION INSTRUCTIONS ====="
    )


def translation_cache_document_identity(document_identity: str, config: AITranslateConfig) -> str:
    """自定义指令变化时切换请求会话键，同时保持空指令与旧版本缓存兼容。"""
    custom_hash = translation_custom_instruction_hash(config)
    return f"{document_identity}\u241f{custom_hash}" if custom_hash else document_identity


@dataclass
class DocumentChatSession:
    title: str
    markdown_path: Path
    selected_text: str = ""
    question: str = ""


DEFAULT_KEY_POINTS_PROMPT = """You are a scientific literature compressor. Your goal is to preserve the minimum information needed to understand the paper's contribution, evidence, and applicability, rather than reproduce the document.

Read only the supplied document. Output in Chinese.

Return exactly four bullet points and nothing else:

* 问题：研究对象、核心问题及现有方法的关键缺口。
* 方法：核心方法或技术路线，以及理解结果所必需的实验设置。
* 结果：最重要的发现，并保留关键指标、比较对象和必要条件。
* 边界：结论的适用范围、主要限制及证据不足之处。

Strict rules:

1. Each bullet may contain one or two concise clauses, but must express only one logical theme.
2. Keep the entire response within approximately 320–450 Chinese characters.
3. Preserve information that materially affects interpretation, including comparison baselines, sample size, evaluation conditions, uncertainty, or statistical significance when explicitly reported and important.
4. If the paper contains multiple contributions, retain at most the two most important; do not merge unrelated findings.
5. Do not provide an introduction, document classification, background review, section-by-section recap, expanded explanation, conclusion paragraph, recommendations, future work, or follow-up questions.
6. Do not repeat the same information in different wording.
7. Mention formulas, parameters, materials, equipment, datasets, or algorithm names only when essential to understanding the contribution or result.
8. Use only information explicitly supported by the document. Do not infer missing results. Write “文档未明确” when a required item is unavailable.
9. If OCR, table, formula, or text-order errors materially affect a bullet, append “[解析存疑]” to that bullet; otherwise do not discuss parsing quality.
10. Before answering, silently remove every detail whose deletion would not change the reader's understanding of the contribution, evidence strength, or applicability.
11. Never exceed four bullets, regardless of document length."""



def build_key_points_prompt_for_document() -> str:
    """
    构造“要点提炼”的预置指引。

    设计意图：
    1. 该指引会作为文献对话中的“引用气泡”传入，而不是直接静默塞到问题里；
    2. 用户群体通常是工程、机理、材料、制造、能源、装备等方向研究人员，但不强制限定领域；
    3. 文档可能是论文、综述、报告、专利、教材或书籍章节，因此指引要求模型先识别文档类型，再组织总结；
    4. 输出格式给出优先关注点，但允许模型根据文档实际结构调整。
    """
    try:
        custom_prompt = str(getattr(app_config.load_settings(), "key_points_prompt", "") or "").strip()
    except Exception:
        custom_prompt = ""
    return custom_prompt or DEFAULT_KEY_POINTS_PROMPT


def build_mineru_document_tool_adapter(ai_module, reader_window_factory=None):
    def save_key(token: str) -> str:
        app_config.save_mineru_token(token)
        return str(app_config.secret_path("mineru", "api_key"))

    def create_parse_worker(input_path: Path, output_dir: Path):
        if input_path.suffix.lower() == ".epub":
            return EpubParseWorker(str(input_path), str(output_dir))
        return MinerUWorker(str(input_path), str(output_dir), ParseOptions())

    def create_reader_window(**kwargs):
        if reader_window_factory is not None:
            return reader_window_factory(**kwargs)
        try:
            from OT_ui import ReaderWindow as RuntimeReaderWindow
        except Exception as exc:
            raise RuntimeError(f"无法创建文档阅读窗口：{exc}") from exc
        return RuntimeReaderWindow(**kwargs)

    return ai_module.DocumentToolAdapter(
        display_name="本地 EPUB / MinerU",
        settings_button_text="设置 MinerU 访问令牌",
        key_label="MinerU 访问令牌：",
        token_placeholder="请输入 MinerU 访问令牌",
        unsupported_file_message="文档解析器暂不支持此文件类型",
        is_configured=lambda: bool(app_config.load_mineru_token() or load_mineru_token()),
        save_key=save_key,
        is_supported_input_file=is_supported_input_file,
        create_output_dir=output_dir_for_pdf,
        create_parse_worker=create_parse_worker,
        latest_translation_path=latest_translation_path,
        find_stored_original=find_stored_original,
        create_reader_window=create_reader_window,
    )


def open_document_chat(session: DocumentChatSession, parent=None):
    try:
        import AI_api_base
    except Exception as exc:
        raise RuntimeError(f"无法导入 AI_api_base.py: {exc}") from exc

    settings = app_config.load_settings()
    adapter = build_mineru_document_tool_adapter(AI_api_base)
    api_session = AI_api_base.DocumentChatSession(
        title=session.title,
        markdown_path=session.markdown_path,
        selected_text=session.selected_text,
        question=session.question,
    )
    return AI_api_base.open_document_chat_session(
        api_session,
        parent=parent,
        settings=settings,
        conversation_history_path=lambda: app_config.chat_history_path(settings),
        document_tool_adapter=adapter,
    )


def output_dir_for_pdf(pdf_path: Path, reserve: bool = False) -> Path:
    """返回新的解析输出目录；批量派发时可原子预留目录。"""
    root = current_work_dir()
    root.mkdir(parents=True, exist_ok=True)
    output_stem = safe_document_stem(pdf_path.stem, MAX_OUTPUT_DIR_STEM_LENGTH)
    index = 0
    while True:
        candidate = root / output_stem if index == 0 else root / f"{output_stem}_{index:02d}"
        if reserve:
            try:
                # A single successful mkdir assigns the base directory; concurrent
                # jobs continue with numbered directories.
                candidate.mkdir(parents=False, exist_ok=False)
                return candidate
            except FileExistsError:
                index += 1
                continue
        if not candidate.exists():
            return candidate
        index += 1


def generated_output_dir_matches_source(folder: Path, input_path: Path) -> bool:
    """使用输出元数据区分同名但来源不同的文件。"""
    source_value = ""
    for metadata_path in (generated_output_marker_path(folder), folder / "mineru_task.json"):
        try:
            payload = json.loads(metadata_path.read_text(encoding="utf-8", errors="replace"))
        except (OSError, json.JSONDecodeError):
            continue
        source_value = str(payload.get("source_file") or payload.get("source_pdf") or "").strip()
        if source_value:
            break
    if not source_value:
        # 兼容没有来源字段的旧版本输出目录。
        return True
    try:
        return os.path.normcase(str(Path(source_value).expanduser().resolve())) == os.path.normcase(str(input_path.expanduser().resolve()))
    except OSError:
        return os.path.normcase(source_value) == os.path.normcase(str(input_path))


def latest_output_dir_for_file(input_path: Path) -> Path | None:
    candidates: list[Path] = []
    root = current_work_dir()
    output_stem = safe_document_stem(input_path.stem, MAX_OUTPUT_DIR_STEM_LENGTH)
    base_dir = root / output_stem
    if is_generated_output_dir(base_dir) and generated_output_dir_matches_source(base_dir, input_path):
        candidates.append(base_dir)
    numbered_pattern = re.compile(rf"^{re.escape(output_stem)}_\d+$")
    for path in root.glob(f"{output_stem}_*"):
        if (
            path.is_dir()
            and numbered_pattern.match(path.name)
            and is_generated_output_dir(path)
            and generated_output_dir_matches_source(path, input_path)
        ):
            candidates.append(path)
    timestamp_pattern = f"{output_stem}_*年*月*日*点*分*"
    candidates.extend(
        path for path in root.glob(timestamp_pattern)
        if path.is_dir() and is_generated_output_dir(path) and generated_output_dir_matches_source(path, input_path)
    )
    if not candidates:
        return None
    return max(candidates, key=lambda path: path.stat().st_mtime)


def is_supported_input_file(path: Path) -> bool:
    return path.suffix.lower() in SUPPORTED_INPUT_EXTENSIONS


def is_epub_input_file(path: Path) -> bool:
    return path.suffix.lower() == ".epub"


def input_requires_mineru(path: Path) -> bool:
    return is_supported_input_file(path) and not is_epub_input_file(path)


def create_document_parse_worker(input_path: Path, output_dir: Path, options: ParseOptions | None = None):
    """Route local/open formats without leaking EPUB into the MinerU upload path."""
    if is_epub_input_file(input_path):
        return EpubParseWorker(str(input_path), str(output_dir))
    return MinerUWorker(str(input_path), str(output_dir), options or ParseOptions())


def load_mineru_token() -> str:
    stored_token = app_config.load_mineru_token()
    if stored_token:
        return stored_token

    # Read the saved token first; environment variables support unattended runs.
    for env_name in ("MINERU_API_KEY", "MINERU_KEY", "MINERU_TOKEN"):
        token = os.getenv(env_name, "").strip().strip('"').strip("'").removeprefix("Bearer ").strip()
        if token and token.isascii() and not any(ch.isspace() for ch in token):
            return token

    raise MinerUError("没有设置 MinerU 访问令牌，请点击“模型、密钥和工作文件夹设置”填写。")


def read_key_file_lines() -> list[str]:
    # Kept for callers that still use the old helper; local key files are not read.
    return []


def load_labelled_secret(label: str) -> str:
    lines = read_key_file_lines()
    expect_value = False
    label_lower = label.lower()
    token_candidates: list[str] = []
    for line in lines:
        if not line or line.startswith("#"):
            continue
        if expect_value:
            if line.isascii() and len(line) >= 20:
                return line.split()[0].removeprefix("Bearer ")
            expect_value = False
        lower = line.lower()
        if label_lower in lower:
            parts = re.split(r"[:=：]", line, maxsplit=1)
            if len(parts) == 2 and parts[1].strip():
                value = parts[1].strip().split()[0].removeprefix("Bearer ")
                if value.isascii() and len(value) >= 20:
                    return value
            expect_value = True
            continue
        if line.isascii() and len(line) >= 20:
            token_candidates.append(line.split()[0].removeprefix("Bearer "))
    return token_candidates[0] if token_candidates else ""


def load_key_setting(name: str) -> str:
    # Environment variables are reserved for unattended configuration.
    return os.getenv(name, "").strip().strip('"').strip("'")


def normalize_ai_base_url(url: str, provider_id: str = "oneapi") -> str:
    spec = get_translation_provider_spec(provider_id)
    default_url = str(spec.get("default_base_url") or "")
    url = (url or default_url).strip().rstrip("/")
    for suffix in ["/chat/completions", "/messages", "/responses", "/models"]:
        if url.endswith(suffix):
            url = url[: -len(suffix)].rstrip("/")
    if not bool(spec.get("append_v1", True)):
        if url.endswith("/v1"):
            url = url[:-3].rstrip("/")
    elif not url.endswith("/v1"):
        url += "/v1"
    return url


def build_ai_translate_config(provider_id: str = "oneapi", log=None) -> AITranslateConfig:
    provider_id = provider_id.lower().strip()
    if machine_translate.is_machine_translation_provider(provider_id):
        provider = app_config.load_settings().providers.get(provider_id)
        base_url = provider.base_url if provider and provider.base_url else (
            machine_translate.MTRAN_SERVER_DEFAULT_BASE_URL
            if provider_id == machine_translate.MTRAN_SERVER_PROVIDER
            else ""
        )
        model = "本地免费机翻" if provider_id == machine_translate.MTRAN_SERVER_PROVIDER else "免费机翻"
        return AITranslateConfig(provider_id=provider_id, api_key=load_provider_secret(provider_id), base_url=base_url, model=model)
    api_key = load_provider_secret(provider_id)
    base_url = normalize_ai_base_url(load_provider_base_url(provider_id), provider_id)
    model = load_provider_model_setting(provider_id)
    if not api_key:
        raise MinerUError(f"没有配置 {translation_provider_name(provider_id)} API 密钥，无法翻译。请在“模型、密钥和工作文件夹设置”中填写。")
    if not model:
        model = choose_ai_model(api_key, base_url, provider_id, log)
    stored_provider = app_config.load_settings().providers.get(provider_id)
    settings = app_config.load_settings()
    return AITranslateConfig(
        provider_id=provider_id,
        api_key=api_key,
        base_url=base_url,
        model=model,
        request_body_mode=getattr(stored_provider, "request_body_mode", "codex"),
        thinking_mode="enabled" if provider_id == "deepseek" and getattr(settings, "translation_deepseek_thinking_enabled", True) else "disabled",
        reasoning_effort=getattr(settings, "translation_deepseek_reasoning_effort", "default"),
        custom_translation_instruction=str(getattr(settings, "translation_custom_instruction", "") or ""),
    )


def build_ai_endpoint_config(provider_id: str = "oneapi") -> AITranslateConfig:
    provider_id = provider_id.lower().strip()
    if machine_translate.is_machine_translation_provider(provider_id):
        provider = app_config.load_settings().providers.get(provider_id)
        base_url = provider.base_url if provider and provider.base_url else (
            machine_translate.MTRAN_SERVER_DEFAULT_BASE_URL
            if provider_id == machine_translate.MTRAN_SERVER_PROVIDER
            else ""
        )
        model = "本地免费机翻" if provider_id == machine_translate.MTRAN_SERVER_PROVIDER else "免费机翻"
        return AITranslateConfig(provider_id=provider_id, api_key=load_provider_secret(provider_id), base_url=base_url, model=model)
    api_key = load_provider_secret(provider_id)
    base_url = normalize_ai_base_url(load_provider_base_url(provider_id), provider_id)
    if not api_key:
        raise MinerUError(f"没有配置 {translation_provider_name(provider_id)} API 密钥，无法翻译。请在“模型、密钥和工作文件夹设置”中填写。")
    stored_provider = app_config.load_settings().providers.get(provider_id)
    settings = app_config.load_settings()
    return AITranslateConfig(
        provider_id=provider_id,
        api_key=api_key,
        base_url=base_url,
        model="",
        request_body_mode=getattr(stored_provider, "request_body_mode", "codex"),
        thinking_mode="enabled" if provider_id == "deepseek" and getattr(settings, "translation_deepseek_thinking_enabled", True) else "disabled",
        reasoning_effort=getattr(settings, "translation_deepseek_reasoning_effort", "default"),
        custom_translation_instruction=str(getattr(settings, "translation_custom_instruction", "") or ""),
    )

def _is_zero_price(value) -> bool:
    text = str(value or "").strip()
    return text in {"0", "0.0", "0.00", "0.000000", "0.0000000"}


def _is_free_model(provider_id: str, model_id: str, official_name: str, pricing: dict | None) -> bool:
    provider_id = (provider_id or "").strip().lower()
    name = (official_name or "").lower()
    model = (model_id or "").lower()
    pricing = pricing if isinstance(pricing, dict) else {}
    prompt_value = pricing.get("prompt")
    completion_value = pricing.get("completion")
    if prompt_value is not None and completion_value is not None and _is_zero_price(prompt_value) and _is_zero_price(completion_value):
        return True
    if model.endswith(":free") or "(free)" in name:
        return True
    if provider_id == "zai":
        return "flash" in model or "免费" in name
    return False


def _humanize_model_name(provider_id: str, model_id: str, official_name: str = "") -> str:
    if official_name:
        return official_name
    if provider_id == "zai":
        return model_id.replace("glm", "GLM", 1)
    return model_id


def _format_model_price_text(provider_id: str, pricing: dict | None, is_free: bool) -> str:
    if provider_id != "openrouter":
        return ""
    if is_free:
        return "免费"
    pricing = pricing if isinstance(pricing, dict) else {}
    prompt_value = str(pricing.get("prompt") or "").strip()
    completion_value = str(pricing.get("completion") or "").strip()
    if prompt_value and completion_value:
        return f"输入 ${prompt_value} / 输出 ${completion_value}"
    return "价格未提供"


def build_translation_model_options(provider_id: str, models: list[dict], current_model: str = "") -> list[TranslationModelOption]:
    provider_id = (provider_id or "oneapi").strip().lower()
    preferred_ids = list(provider_preferred_models(provider_id))
    preferred_rank = {model_id.lower(): index for index, model_id in enumerate(preferred_ids)}
    options: list[TranslationModelOption] = []
    for item in models:
        if not isinstance(item, dict):
            continue
        model_id = str(item.get("id") or "").strip()
        if provider_id == "gemini":
            model_id = normalize_gemini_model_id(model_id)
        if not model_id:
            continue
        official_name = _humanize_model_name(provider_id, model_id, str(item.get("name") or "").strip())
        pricing = item.get("pricing") if isinstance(item.get("pricing"), dict) else {}
        is_free = _is_free_model(provider_id, model_id, official_name, pricing)
        price_text = _format_model_price_text(provider_id, pricing, is_free)
        display_text = official_name
        if price_text:
            display_text += f"  -  {price_text}"
        options.append(
            TranslationModelOption(
                model_id=model_id,
                display_text=display_text,
                official_name=official_name,
                price_text=price_text,
                is_free=is_free,
            )
        )

    def sort_key(option: TranslationModelOption):
        model_lower = option.model_id.lower()
        tokens = [item for item in re_split_model_name(model_lower) if item]
        current_rank = 0 if current_model and current_model == option.model_id else 1
        preferred = preferred_rank.get(model_lower, 999)
        if provider_id == "zai":
            quality_rank = 0 if "flash" in tokens else 1
        elif provider_id == "openrouter":
            quality_rank = 0
        else:
            quality_rank = 0 if any(token in tokens for token in ("ultra", "max", "pro", "reasoner")) else 1
        return (preferred, current_rank, quality_rank, option.official_name.lower(), model_lower)

    return sorted(options, key=sort_key)


def fetch_translation_model_options(provider_id: str, api_key: str, base_url: str) -> list[TranslationModelOption]:
    result = http_json("GET", provider_model_list_url(provider_id, base_url), token=api_key, timeout=60)
    models = result.get("data") or []
    return build_translation_model_options(provider_id, list(models) if isinstance(models, list) else [])


def choose_preferred_translation_model(provider_id: str, model_options: list[TranslationModelOption], current: str = "") -> str:
    if current and any(option.model_id == current for option in model_options):
        return current
    preferred = provider_default_model(provider_id)
    if preferred:
        for option in model_options:
            if option.model_id.lower() == preferred.lower():
                return option.model_id
    preferred_ids = {model_id.lower() for model_id in provider_preferred_models(provider_id)}
    for option in model_options:
        if option.model_id.lower() in preferred_ids:
            return option.model_id
    return model_options[0].model_id if model_options else current


def choose_ai_model(api_key: str, base_url: str, provider_id: str = "oneapi", log=None) -> str:
    """
    在未配置显式模型时，根据可用模型列表自动选择一个更偏高质量的默认翻译模型。

    设计原则：
    1. 优先选择 pro / plus / 高质量通用模型；
    2. 避免默认落到 mini / flash / lite 这类轻量模型；
    3. 仅在没有更优候选时，才回退到轻量模型。
    """
    try:
        model_options = fetch_translation_model_options(provider_id, api_key, base_url)
        selected = choose_preferred_translation_model(provider_id, model_options)
        if selected and log:
            log(f"已选择翻译模型：{selected}")
        if selected:
            return selected
    except Exception as exc:
        fallback = provider_default_model(provider_id) or "gpt-5.6-luna"
        if log:
            log(f"获取模型列表失败，使用默认模型 {fallback}：{exc}")
        return fallback
    return provider_default_model(provider_id) or "gpt-5.6-luna"


def http_json(method: str, url: str, payload: dict | None = None, token: str | None = None, timeout: int = 60, extra_headers: dict | None = None) -> dict:
    data = None
    headers = {"User-Agent": USER_AGENT}
    if token:
        if not token.isascii():
            raise MinerUError("MinerU 访问令牌包含无效字符，请检查设置。")
        headers["Authorization"] = f"Bearer {token}"
    if payload is not None:
        data = json.dumps(payload, ensure_ascii=False).encode("utf-8")
        headers["Content-Type"] = "application/json"
    if extra_headers:
        headers.update(extra_headers)

    request = urllib.request.Request(url, data=data, headers=headers, method=method)
    try:
        with urllib.request.urlopen(request, timeout=timeout) as response:
            text = response.read().decode("utf-8")
    except urllib.error.HTTPError as exc:
        detail = exc.read().decode("utf-8", errors="replace")
        raise MinerUError(
            f"HTTP {exc.code}: {detail}",
            status=exc.code,
            retry_after=retry_after_seconds(exc.headers),
        ) from exc
    except urllib.error.URLError as exc:
        raise MinerUError(f"网络请求失败: {exc.reason}") from exc

    try:
        return json.loads(text)
    except json.JSONDecodeError as exc:
        raise MinerUError(f"接口返回不是 JSON: {text[:300]}") from exc


def is_gpt56_cache_model(model: str) -> bool:
    return bool(re.search(r"(?:^|[^0-9])gpt[-_ ]?5[._-]?6(?:[^0-9]|$)", str(model or "").lower()))


def make_translation_cache_key(document_identity: str, model: str, target_language: str) -> str:
    """为同一文档的连续翻译轮次生成稳定且不泄露路径的 UUID。"""
    source = f"translation/{document_identity}|{model}|{target_language}"
    return str(uuid.uuid5(uuid.NAMESPACE_URL, source))


def build_codex_cache_headers(config: AITranslateConfig) -> dict:
    """仅在 OneAPI 的 Codex 构造且已有稳定任务键时发送会话关联头。"""
    key = str(config.prompt_cache_key or "").strip()
    if not (key and uses_codex_construction(config.provider_id, config.request_body_mode)):
        return {}
    return {
        "session-id": key,
        "thread-id": key,
        "x-client-request-id": key,
    }


def translation_request_audit_enabled() -> bool:
    """Whether developer request transcripts should be written beside a document."""
    return os.environ.get("LITMTRANS_TRANSLATION_REQUEST_AUDIT", "").strip().lower() in {"1", "true", "yes", "on"}


def save_translation_request_audit(
    folder: Path,
    request_kind: str,
    config: AITranslateConfig,
    messages: list[dict],
    *,
    timeout: int,
) -> Path | None:
    """Write an opt-in, human-readable translation request transcript.

    The messages are the complete semantic payload sent to the model.  API keys
    and transport headers are intentionally excluded from the audit artifact.
    """
    if not translation_request_audit_enabled():
        return None
    try:
        audit_dir = Path(folder) / "翻译请求记录"
        audit_dir.mkdir(parents=True, exist_ok=True)
        stamp = time.strftime("%Y%m%d-%H%M%S") + f"-{time.time_ns() % 1_000_000_000:09d}"
        safe_kind = re.sub(r"[<>:\\/*?\"|\r\n]+", "-", str(request_kind or "翻译请求")).strip(" .-") or "翻译请求"
        path = audit_dir / f"{safe_kind}-{stamp}.txt"
        body = [
            "翻译请求记录（不含 API 密钥）",
            f"时间: {stamp}",
            f"请求类型: {request_kind}",
            f"提供商: {config.provider_id}",
            f"模型: {config.model}",
            f"接口地址: {config.base_url}",
            f"请求构造: {normalize_request_body_mode(config.provider_id, config.request_body_mode)}",
            f"请求缓存键: {str(getattr(config, 'prompt_cache_key', '') or '')}",
            f"超时秒数: {timeout}",
            f"消息数: {len(messages)}",
            "",
        ]
        for index, message in enumerate(messages, 1):
            role = str(message.get("role") or "") if isinstance(message, dict) else ""
            content = message.get("content") if isinstance(message, dict) else message
            body.extend(
                [
                    f"===== MESSAGE {index} | role={role} =====",
                    str(content or ""),
                    "",
                ]
            )
        path.write_text("\n".join(body), encoding="utf-8")
        return path
    except Exception:
        # Development auditing must never prevent an already-authorized request.
        return None


def _ai_chat_completion_once(
    config: AITranslateConfig,
    messages: list[dict],
    timeout: int | None = 180,
    stream_callback=None,
    reasoning_callback=None,
    usage_callback=None,
    should_stop=None,
) -> str:
    # 请求槽覆盖完整的普通响应或流式响应生命周期。
    # 即使以后某处错误地创建了过多工作线程，DeepSeek 推理请求也不会超过 100。
    with provider_request_slot(config.provider_id):
        return _ai_chat_completion_once_unlimited(
            config,
            messages,
            timeout,
            stream_callback,
            reasoning_callback,
            usage_callback,
            should_stop,
        )


def _ai_chat_completion_once_unlimited(
    config: AITranslateConfig,
    messages: list[dict],
    timeout: int | None = 180,
    stream_callback=None,
    reasoning_callback=None,
    usage_callback=None,
    should_stop=None,
) -> str:
    mode = normalize_request_body_mode(config.provider_id, config.request_body_mode)
    if uses_claude_construction(config.provider_id, mode):
        payload = build_claude_messages_payload(
            config.model,
            messages,
            stream=stream_callback is not None,
            temperature=0.2,
        )
        headers = claude_headers(config.api_key, stream=stream_callback is not None)
        endpoint = request_url_for_construction(config.base_url, config.provider_id, mode)
    else:
        payload = {
            "model": config.model,
            "messages": messages,
            "temperature": 0.2,
            "stream": stream_callback is not None,
        }
        headers = build_codex_cache_headers(config)
        if headers:
            payload["prompt_cache_key"] = config.prompt_cache_key
            payload["stream_options"] = {"include_usage": True}
        elif is_gemini_provider(config.provider_id, config.base_url):
            payload["stream_options"] = {"include_usage": True}
            thinking_enabled = str(config.thinking_mode or "").strip().lower() == "enabled"
            if thinking_enabled:
                payload["extra_body"] = {
                    "google": {"thinking_config": gemini_translation_thinking_config(config.model, config.reasoning_effort)}
                }
            else:
                payload["reasoning_effort"] = "minimal"
        elif config.provider_id == "deepseek":
            thinking_enabled = str(config.thinking_mode or "").strip().lower() == "enabled"
            payload["thinking"] = {"type": "enabled" if thinking_enabled else "disabled"}
            # DeepSeek reports disk-cache hit/miss counters in the terminal
            # SSE event only when usage is requested.
            payload["stream_options"] = {"include_usage": True}
            if thinking_enabled:
                effort = str(config.reasoning_effort or "high").strip().lower()
                payload["reasoning_effort"] = effort if effort in {"high", "max"} else "high"
        if is_gemini_provider(config.provider_id, config.base_url):
            headers["x-goog-api-client"] = "litmtrans/1.0"
        endpoint = request_url_for_construction(config.base_url, config.provider_id, mode)
    if stream_callback:
        return ai_chat_completion_stream(config, payload, timeout, stream_callback, reasoning_callback, usage_callback, should_stop)
    result = http_json(
        "POST",
        endpoint,
        payload,
        token=None if uses_claude_construction(config.provider_id, mode) else config.api_key,
        timeout=timeout,
        extra_headers=headers,
    )
    content = extract_claude_response_text(result) if uses_claude_construction(config.provider_id, mode) else extract_ai_response_text(result)
    if not content:
        raise MinerUError(f"翻译服务未返回有效结果：{result}")
    return content


def ai_chat_completion_stream(
    config: AITranslateConfig,
    payload: dict,
    timeout: int | None,
    stream_callback,
    reasoning_callback=None,
    usage_callback=None,
    should_stop=None,
) -> str:
    should_stop = should_stop or current_translation_stop()
    data = json.dumps(payload, ensure_ascii=False).encode("utf-8")
    mode = normalize_request_body_mode(config.provider_id, config.request_body_mode)
    use_claude_messages = uses_claude_construction(config.provider_id, mode)
    if use_claude_messages:
        headers = claude_headers(config.api_key, stream=True)
    else:
        headers = {
            "User-Agent": USER_AGENT,
            "Authorization": f"Bearer {config.api_key}",
            "Content-Type": "application/json",
            "Accept": "text/event-stream",
        }
        headers.update(build_codex_cache_headers(config))
        if is_gemini_provider(config.provider_id, config.base_url):
            headers["x-goog-api-client"] = "litmtrans/1.0"
    endpoint = request_url_for_construction(config.base_url, config.provider_id, mode)
    request = urllib.request.Request(endpoint, data=data, headers=headers, method="POST")
    chunks: list[str] = []
    in_gemini_thought = False
    response = None
    stop_watcher_done = threading.Event()
    stop_watcher = None
    try:
        response = urllib.request.urlopen(request, timeout=timeout)
        if should_stop is not None:
            def close_on_stop():
                while not stop_watcher_done.wait(0.05):
                    try:
                        if should_stop():
                            response.close()
                            return
                    except Exception:
                        return
            stop_watcher = threading.Thread(target=close_on_stop, name="translation-http-cancel", daemon=True)
            stop_watcher.start()
        with response:
            for raw_line in response:
                if should_stop is not None and should_stop():
                    raise MinerUError("用户已停止翻译。")
                line = raw_line.decode("utf-8", errors="replace").strip()
                if not line or line.startswith(":"):
                    continue
                if line.startswith("data:"):
                    line = line[5:].strip()
                if line == "[DONE]":
                    break
                try:
                    event = json.loads(line)
                except json.JSONDecodeError:
                    continue
                usage = event.get("usage")
                if isinstance(usage, dict) and usage_callback:
                    usage_callback(dict(usage))
                if use_claude_messages:
                    event_type = str(event.get("type") or "")
                    delta = event.get("delta") if isinstance(event.get("delta"), dict) else {}
                    content = str(delta.get("text") or "") if event_type == "content_block_delta" and delta.get("type") == "text_delta" else ""
                    reasoning = str(delta.get("thinking") or "") if event_type == "content_block_delta" and delta.get("type") == "thinking_delta" else ""
                else:
                    content = extract_ai_stream_text(event)
                    reasoning = extract_ai_stream_reasoning(event)
                    if content and is_gemini_provider(config.provider_id, config.base_url):
                        content, tagged_reasoning, in_gemini_thought = split_gemini_thought_delta(content, in_gemini_thought)
                        if gemini_stream_delta_is_thought(event):
                            # Gemini's OpenAI-compatible stream can mark a public
                            # thought in ``extra_content.google.thought`` without
                            # enclosing it in <thought> tags.
                            reasoning += tagged_reasoning or content
                            content = ""
                        else:
                            reasoning += tagged_reasoning
                if reasoning and reasoning_callback:
                    reasoning_callback(reasoning)
                if content:
                    chunks.append(content)
                    stream_callback(content)
    except MinerUError:
        raise
    except urllib.error.HTTPError as exc:
        detail = exc.read().decode("utf-8", errors="replace")
        raise MinerUError(
            f"HTTP {exc.code}: {detail}",
            status=exc.code,
            retry_after=retry_after_seconds(exc.headers),
        ) from exc
    except urllib.error.URLError as exc:
        if should_stop is not None and should_stop():
            raise MinerUError("用户已停止翻译。") from exc
        raise MinerUError(f"网络请求失败: {exc.reason}") from exc
    except (OSError, ValueError) as exc:
        if should_stop is not None and should_stop():
            raise MinerUError("用户已停止翻译。") from exc
        raise
    finally:
        stop_watcher_done.set()
        if stop_watcher is not None and stop_watcher.is_alive():
            stop_watcher.join(timeout=0.2)

    content = "".join(chunks).strip()
    if not content:
        # Reaching here means the server explicitly finished (or closed) the
        # SSE response.  While it is still thinking, ``urlopen`` remains
        # blocked above and the caller continues waiting; do not mistake an
        # empty thinking period for this terminal condition.
        raise MinerUError("翻译服务已结束，但没有返回正文。")
    return content


def ai_chat_completion(
    config: AITranslateConfig,
    messages: list[dict],
    timeout: int | None = 180,
    stream_callback=None,
    reasoning_callback=None,
    rate_limit_callback=None,
    usage_callback=None,
    should_stop=None,
) -> str:
    if not is_gemini_provider(config.provider_id, config.base_url):
        return _ai_chat_completion_once(
            config,
            messages,
            timeout,
            stream_callback,
            reasoning_callback,
            usage_callback,
            should_stop,
        )
    coordinator = gemini_quota_coordinator(config)
    attempts = 4
    for attempt in range(1, attempts + 1):
        coordinator.wait(rate_limit_callback, should_stop)
        try:
            return _ai_chat_completion_once(
                config,
                messages,
                timeout,
                stream_callback,
                reasoning_callback,
                usage_callback,
                should_stop,
            )
        except MinerUError as exc:
            quota_limited = (
                exc.status == 429
                or re.search(
                    r"(?:resource[_ -]?exhausted|quota|rate.?limit)",
                    str(exc),
                    flags=re.IGNORECASE,
                )
                is not None
            )
            if not quota_limited or attempt >= attempts:
                raise
            fallback = min(15.0 * (2 ** (attempt - 1)), 60.0)
            coordinator.register(exc.retry_after or fallback)
    raise MinerUError("Gemini 请求额度等待后仍未成功。")


def split_gemini_thought_delta(text: str, in_thought: bool) -> tuple[str, str, bool]:
    """Separate public Gemini ``<thought>`` stream content from final text."""
    visible_parts: list[str] = []
    thought_parts: list[str] = []
    index = 0
    while index < len(text):
        if in_thought:
            end_index = text.find("</thought>", index)
            if end_index == -1:
                thought_parts.append(text[index:])
                break
            thought_parts.append(text[index:end_index])
            index = end_index + len("</thought>")
            in_thought = False
            continue
        start_index = text.find("<thought>", index)
        if start_index == -1:
            visible_parts.append(text[index:])
            break
        visible_parts.append(text[index:start_index])
        index = start_index + len("<thought>")
        in_thought = True
    return "".join(visible_parts), "".join(thought_parts), in_thought


def gemini_stream_delta_is_thought(event: dict) -> bool:
    """Return whether Gemini explicitly labels this OpenAI-compatible delta as thought."""
    choices = event.get("choices") or []
    if not choices or not isinstance(choices[0], dict):
        return False
    delta = choices[0].get("delta") or {}
    if not isinstance(delta, dict):
        return False
    extra_content = delta.get("extra_content") or {}
    google = extra_content.get("google") if isinstance(extra_content, dict) else {}
    return bool(google.get("thought")) if isinstance(google, dict) else False


def extract_ai_response_text(result: dict) -> str:
    choices = result.get("choices") or []
    if choices:
        choice = choices[0] or {}
        message = choice.get("message") or {}
        content = message.get("content") or choice.get("text") or ""
        if isinstance(content, list):
            content = "".join(str(item.get("text") or item.get("content") or "") if isinstance(item, dict) else str(item) for item in content)
        if content:
            return str(content).strip()
    output_text = result.get("output_text")
    if output_text:
        return str(output_text).strip()
    output = result.get("output") or []
    parts: list[str] = []
    if isinstance(output, list):
        for item in output:
            for block in (item.get("content") or []) if isinstance(item, dict) else []:
                if isinstance(block, dict):
                    parts.append(str(block.get("text") or block.get("content") or ""))
    return "".join(parts).strip()


def extract_ai_stream_text(event: dict) -> str:
    choices = event.get("choices") or []
    if choices:
        choice = choices[0] or {}
        delta = choice.get("delta") or {}
        message = choice.get("message") or {}
        content = delta.get("content") or delta.get("text") or message.get("content") or choice.get("text") or ""
        if isinstance(content, list):
            content = "".join(str(item.get("text") or item.get("content") or "") if isinstance(item, dict) else str(item) for item in content)
        if content:
            return str(content)
    for key in ("response.output_text.delta", "output_text_delta", "text_delta"):
        if event.get("type") == key and event.get("delta"):
            return str(event.get("delta"))
    if event.get("type") in {"message.delta", "content.delta"} and event.get("delta"):
        delta = event.get("delta")
        if isinstance(delta, dict):
            return str(delta.get("text") or delta.get("content") or "")
        return str(delta)
    return ""


def stringify_reasoning_value(value) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value
    if isinstance(value, dict):
        parts = []
        for key in ("content", "text", "summary", "reasoning_content", "reasoning", "thinking", "thought"):
            if key in value:
                parts.append(stringify_reasoning_value(value.get(key)))
        return "".join(parts)
    if isinstance(value, list):
        return "".join(stringify_reasoning_value(item) for item in value)
    return str(value)


def extract_ai_stream_reasoning(event: dict) -> str:
    candidates = []
    choices = event.get("choices") or []
    for choice in choices:
        if not isinstance(choice, dict):
            continue
        delta = choice.get("delta") or {}
        message = choice.get("message") or {}
        for container in (delta, message, choice):
            if not isinstance(container, dict):
                continue
            for key in (
                "reasoning_content",
                "reasoning",
                "thinking",
                "thought",
                "reasoning_text",
                "reasoning_summary",
                "reasoning_details",
                "analysis",
            ):
                if key in container:
                    candidates.append(container.get(key))
    for key in ("reasoning_delta", "thinking_delta", "analysis_delta"):
        if key in event:
            candidates.append(event.get(key))
    return "".join(stringify_reasoning_value(value) for value in candidates)


def stable_translation_marker(markdown: str) -> str:
    digest = hashlib.sha256(markdown.encode("utf-8")).hexdigest()
    number = int(digest[:12], 16) % 100000000
    return f"<<<{number:08d}>>>"


def marker_detected(text: str, marker: str) -> bool:
    if marker in text:
        return True
    digits = "".join(ch for ch in marker if ch.isdigit())
    if not digits:
        return False
    tail = text[-1200:]
    separators = r"[\s_\-.,:;|/\\]*"
    digit_pattern = separators.join(re.escape(ch) for ch in digits)
    return re.search(r"(?<!\d)" + digit_pattern, tail, flags=re.UNICODE) is not None


def completion_token_instruction(marker: str) -> str:
    """
    构造全文连续翻译的完成令牌说明。

    为什么需要单独函数：
    1. 完整参考语料加入后，上下文会显著变长，模型更容易遗忘终止编号；
    2. 一些模型会把“完成令牌”泛化成 TRANSLATION COMPLETE 之类的英文短语；
    3. 因此必须反复强调：真实令牌就是 marker 本身，不是自然语言说明。
    """
    return (
        "Completion token requirement:\n"
        f"- The exact completion token is: {marker}\n"
        "- This token is a machine-detection sentinel, not a phrase to translate.\n"
        "- Do NOT replace it with 'TRANSLATION COMPLETE', 'translation complete', '完成', '结束', or any other words.\n"
        "- After the entire translation is complete, output the exact token on a new final line.\n"
        "- The final line must contain only the token itself, with no prefix, suffix, label, explanation, punctuation, or code fence."
    )


def markdown_blocks(text: str) -> list[str]:
    lines = text.splitlines(keepends=True)
    blocks: list[str] = []
    current: list[str] = []
    in_fence = False
    in_math = False
    in_table = False

    def flush():
        nonlocal current
        if current:
            blocks.append("".join(current))
            current = []

    for line in lines:
        stripped = line.strip()
        starts_fence = stripped.startswith("```") or stripped.startswith("~~~")
        starts_math = stripped in {"$$", "\\[", "\\]"} or stripped.startswith("$$")
        is_table_line = stripped.startswith("|") and stripped.endswith("|")

        if starts_fence:
            current.append(line)
            in_fence = not in_fence
            if not in_fence:
                flush()
            continue

        if in_fence:
            current.append(line)
            continue

        if starts_math and stripped.startswith("$$"):
            current.append(line)
            in_math = not in_math if stripped == "$$" else in_math
            if not in_math and stripped == "$$":
                flush()
            continue

        if in_math:
            current.append(line)
            continue

        if is_table_line:
            if not in_table and current:
                flush()
            in_table = True
            current.append(line)
            continue
        if in_table:
            flush()
            in_table = False

        if stripped == "":
            current.append(line)
            flush()
            continue

        if line.startswith("#") and current:
            flush()
        current.append(line)

    flush()
    return blocks


def split_markdown_for_translation(text: str, max_chars: int = 135000) -> list[str]:
    blocks = markdown_blocks(text)
    chunks: list[str] = []
    current = ""
    current_len = 0
    for block in blocks:
        if current and current_len + len(block) > max_chars:
            chunks.append(current)
            current = ""
            current_len = 0
        if len(block) > max_chars:
            chunks.append(block)
            continue
        current += block
        current_len += len(block)
    if current:
        chunks.append(current)
    return chunks


def target_language_instruction(target_language: str) -> str:
    language = (target_language or "简体中文").strip()
    if language in {"简体中文", "繁体中文"}:
        return f"Use {language}. For technical terms, prefer standard academic Chinese translations and keep necessary English abbreviations."
    return f"Use {language}. Do not translate into Chinese unless the source text itself contains Chinese that should be translated into {language}."


def translation_language_suffix(target_language: str) -> str:
    mapping = {
        "简体中文": "zh",
        "繁体中文": "zh-tw",
        "英文": "en",
        "英语": "en",
        "English": "en",
        "日文": "ja",
        "日语": "ja",
        "韩文": "ko",
        "韩语": "ko",
    }
    language = (target_language or "简体中文").strip()
    return mapping.get(language, re.sub(r"[^A-Za-z0-9_-]+", "_", language).strip("_").lower() or "translated")


def translation_output_path(markdown_path: Path, target_language: str) -> Path:
    suffix = translation_language_suffix(target_language)
    return markdown_path.with_name("full.zh.md" if suffix == "zh" else f"full.{suffix}.md")


def translation_variant_suffix(reference_paths: list[str] | None, target_language: str, model: str, provider_id: str = "") -> str:
    if machine_translate.is_machine_translation_provider(provider_id):
        if (provider_id or "").strip().lower() == machine_translate.MTRAN_SERVER_PROVIDER:
            return ".local-machine"
        return ".free-machine"
    if not reference_paths:
        return ""
    return f".ref-{reference_context_cache_key(reference_paths, target_language, model)[:12]}"


def translation_output_path_for_job(markdown_path: Path, job_config: "TranslationJobConfig") -> Path:
    suffix = translation_language_suffix(job_config.target_language)
    variant = translation_variant_suffix(
        job_config.reference_paths,
        job_config.target_language,
        job_config.ai_config.model,
        job_config.ai_config.provider_id,
    )
    return markdown_path.with_name(f"full.{suffix}{variant}.md")


LAYOUT_TRANSLATION_ARTIFACT_PROTOCOL = "layout-translation-artifact-v2-image-footnote"


def translation_job_identity(job_config: "TranslationJobConfig") -> dict:
    """生成可持久化的翻译配置身份，用于精确判断结果能否复用。"""
    provider_id = str(job_config.ai_config.provider_id or "").strip().lower()
    mode = "chunked" if str(job_config.mode or "").strip().lower() in {"chunked", "chunks"} else "full_context"
    reference_key = ""
    if job_config.reference_paths:
        reference_key = reference_context_cache_key(
            job_config.reference_paths,
            job_config.target_language,
            job_config.ai_config.model,
        )
    identity = {
        "version": 1,
        "provider_id": provider_id,
        "base_url": str(job_config.ai_config.base_url or "").strip().rstrip("/"),
        "model": str(job_config.ai_config.model or "").strip(),
        "request_body_mode": str(job_config.ai_config.request_body_mode or "standard").strip().lower(),
        "thinking_mode": str(job_config.ai_config.thinking_mode or "disabled").strip().lower(),
        "reasoning_effort": str(job_config.ai_config.reasoning_effort or "").strip().lower(),
        "source_language": str(job_config.source_language or "英文").strip() or "英文",
        "target_language": str(job_config.target_language or "简体中文").strip() or "简体中文",
        "mode": mode,
        "reference_key": reference_key,
    }
    custom_instruction_hash = translation_custom_instruction_hash(job_config.ai_config)
    if custom_instruction_hash:
        identity["custom_instruction_hash"] = custom_instruction_hash
    return identity


def translation_job_metadata_path(output_path: Path) -> Path:
    return output_path.with_name(f"{output_path.name}.meta.json")


def write_translation_job_metadata(output_path: Path, job_config: "TranslationJobConfig") -> Path:
    metadata_path = translation_job_metadata_path(output_path)
    payload = {
        "identity": translation_job_identity(job_config),
        "output_file": output_path.name,
        "updated_at": time.strftime("%Y-%m-%dT%H:%M:%S"),
    }
    write_text_atomic(metadata_path, json.dumps(payload, ensure_ascii=False, indent=2))
    return metadata_path


def translation_output_matches_job(markdown_path: Path, job_config: "TranslationJobConfig") -> bool:
    output_path = translation_output_path_for_job(markdown_path, job_config)
    if not output_path.exists():
        return False
    metadata_path = translation_job_metadata_path(output_path)
    try:
        payload = json.loads(metadata_path.read_text(encoding="utf-8", errors="replace"))
    except (OSError, json.JSONDecodeError):
        # 旧版本没有配置元数据，安全起见仅重译一次并补齐元数据。
        return False
    return payload.get("identity") == translation_job_identity(job_config)


def layout_translation_matches_job(markdown_path: Path, job_config: "TranslationJobConfig") -> bool:
    html_path = layout_translation_preview_html_path(markdown_path)
    bundle_path = layout_translation_bundle_path(markdown_path)
    if not html_path.exists() or not bundle_path.exists():
        return False
    try:
        payload = json.loads(bundle_path.read_text(encoding="utf-8", errors="replace"))
    except (OSError, json.JSONDecodeError):
        return False
    return (
        payload.get("_translation_job_identity") == translation_job_identity(job_config)
        and payload.get("_layout_translation_artifact_protocol") == LAYOUT_TRANSLATION_ARTIFACT_PROTOCOL
    )


def translation_work_dir_for_job(markdown_path: Path, job_config: "TranslationJobConfig") -> Path:
    variant = translation_variant_suffix(
        job_config.reference_paths,
        job_config.target_language,
        job_config.ai_config.model,
        job_config.ai_config.provider_id,
    )
    custom_instruction_hash = translation_custom_instruction_hash(job_config.ai_config)
    if custom_instruction_hash:
        variant += f".instruction-{custom_instruction_hash[:12]}"
    return markdown_path.parent / f"translation_chunks{variant}"


def remove_paths_quietly(paths: list[Path], log=None, skip_existing: set[Path] | None = None) -> None:
    skip_existing = skip_existing or set()
    for path in paths:
        try:
            resolved = path.resolve()
            if resolved in skip_existing:
                continue
            if not path.exists():
                continue
            if path.is_dir():
                shutil.rmtree(path)
            else:
                path.unlink()
            if log:
                log(f"已清理未完成任务的临时文件：{path}")
        except Exception as exc:
            if log:
                log(f"清理临时文件失败：{path} ({exc})")


def publish_layout_translation_artifacts(
    html_tmp_path: Path,
    html_path: Path,
    bundle_tmp_path: Path,
    bundle_path: Path,
) -> None:
    """Publish a new layout HTML/bundle pair, restoring the old pair on error.

    The two files cannot be replaced atomically as one filesystem operation.
    Keep short-lived same-directory backups so a disk error between the two
    replacements never leaves a prior published HTML paired with a new bundle
    (or the reverse).  This function never touches PDF / Word exports.
    """
    pairs = [(html_tmp_path, html_path), (bundle_tmp_path, bundle_path)]
    backups: dict[Path, Path | None] = {}
    published: set[Path] = set()
    try:
        for _tmp_path, destination in pairs:
            if destination.exists():
                backup = destination.with_name(f".{destination.name}.{uuid.uuid4().hex}.rollback")
                shutil.copy2(destination, backup)
                backups[destination] = backup
            else:
                backups[destination] = None
        for tmp_path, destination in pairs:
            tmp_path.replace(destination)
            published.add(destination)
    except Exception:
        for _tmp_path, destination in reversed(pairs):
            backup = backups.get(destination)
            try:
                if backup and backup.exists():
                    backup.replace(destination)
                elif destination in published and destination.exists():
                    destination.unlink()
            except OSError:
                pass
        raise
    finally:
        for backup in backups.values():
            if backup:
                try:
                    if backup.exists():
                        backup.unlink()
                except OSError:
                    pass


def snapshot_existing_paths(root: Path) -> set[Path]:
    if not root.exists():
        return set()
    try:
        return {path.resolve() for path in root.rglob("*")}
    except Exception:
        return set()


def remove_new_paths_under(root: Path, before_snapshot: set[Path], log=None) -> None:
    if not root.exists():
        return
    try:
        candidates = [path for path in root.rglob("*") if path.resolve() not in before_snapshot]
        candidates.sort(key=lambda path: len(path.parts), reverse=True)
        remove_paths_quietly(candidates, log)
    except Exception as exc:
        if log:
            log(f"清理临时缓存失败：{root} ({exc})")


def clear_translation_artifacts(
    markdown_path: Path,
    target_language: str,
    mode: str,
    reference_paths: list[str] | None = None,
    model: str = "",
    provider_id: str = "",
    custom_instruction: str = "",
) -> None:
    """删除指定目标语言的旧译文与翻译中间状态，确保真正从头重翻。"""
    output_variant = translation_variant_suffix(reference_paths, target_language, model, provider_id)
    work_variant = output_variant
    if not machine_translate.is_machine_translation_provider(provider_id):
        custom_instruction_hash = translation_custom_instruction_hash(custom_instruction)
        if custom_instruction_hash:
            work_variant += f".instruction-{custom_instruction_hash[:12]}"
    work_dir = markdown_path.parent / f"translation_chunks{work_variant}"
    language_suffix = translation_language_suffix(target_language)
    targets: list[Path] = [markdown_path.with_name(f"full.{language_suffix}{output_variant}.md")]
    if work_dir.exists():
        targets.extend(
            [
                work_dir / f"translation_context.{language_suffix}.md",
                work_dir / f"manifest.{language_suffix}.json",
                work_dir / f"full_context.{language_suffix}.translated.md",
                work_dir / f"full_context.{language_suffix}.state.json",
                work_dir / f"full_context.{language_suffix}.messages.json",
            ]
        )
        targets.extend(sorted(work_dir.glob(f"full_context.{language_suffix}.part_*.md")))
        targets.extend(sorted(work_dir.glob(f"chunk_*.{language_suffix}.translated.md*")))
        if mode == "chunks":
            targets.extend(sorted(work_dir.glob("chunk_*.source.md")))
    for path in targets:
        try:
            if path.exists():
                path.unlink()
        except OSError:
            pass


def clear_layout_translation_artifacts(
    markdown_path: Path,
    target_language: str,
    reference_paths: list[str] | None = None,
    model: str = "",
    provider_id: str = "",
    preserve_published_preview: bool = False,
) -> None:
    """Clear layout-translation artifacts before a forced retranslation.

    A retranslation must never reuse the block-translation cache.  It may,
    however, keep the currently published HTML/PDF/Word-state artifacts until
    a replacement is fully generated.  This lets the reader retain a known
    good version if the new model request fails or is cancelled.
    """
    language_suffix = translation_language_suffix(target_language)
    layout_html_path = layout_translation_preview_html_path(markdown_path)
    layout_pdf_path = layout_html_path.with_name(f"{layout_html_path.stem}.final-layout.pdf")
    targets: list[Path] = []
    if not preserve_published_preview:
        targets.extend(
            [
                layout_html_path,
                layout_pdf_path,
                layout_pdf_path.with_suffix(".pdf.meta.json"),
                layout_html_path.with_name(f"{layout_html_path.stem}.final-layout-state.json"),
            ]
        )
    if machine_translate.is_machine_translation_provider(provider_id):
        machine_suffix = "local-machine" if (provider_id or "").strip().lower() == machine_translate.MTRAN_SERVER_PROVIDER else "free-machine"
        targets.append(markdown_path.parent / f"layout_translation_blocks.{language_suffix}.{machine_suffix}.json")
    else:
        reference_suffix = reference_context_cache_key(reference_paths or [], target_language, model)[:12] if reference_paths else ""
        if reference_suffix:
            targets.append(markdown_path.parent / f"layout_translation_blocks.{language_suffix}.ref-{reference_suffix}.json")
        else:
            targets.append(markdown_path.parent / f"layout_translation_blocks.{language_suffix}.json")
    for path in targets:
        try:
            if path.exists():
                path.unlink()
        except OSError:
            pass


def translation_candidates(folder: Path) -> list[Path]:
    return sorted(folder.glob("full.*.md"), key=lambda path: path.stat().st_mtime, reverse=True)


def latest_translation_path(folder: Path) -> Path | None:
    for path in translation_candidates(folder):
        if path.name not in {"full.md", "full.cleaned.md"} and path.exists():
            return path
    return None


def translation_artifact_paths(markdown_path: Path) -> tuple[Path | None, Path | None]:
    """统一返回流式译文和排版译文，作为列表、按钮和阅读器的共同状态来源。"""
    markdown_path = Path(markdown_path)
    stream_path = latest_translation_path(markdown_path.parent)
    layout_candidate = layout_translation_preview_html_path(markdown_path)
    layout_path = layout_candidate if layout_candidate.exists() else None
    return stream_path, layout_path


def translation_status_badge(markdown_path: Path) -> str:
    """生成文档列表中的译文状态，避免只生成排版译文时被误判为“无译文”。"""
    stream_path, layout_path = translation_artifact_paths(markdown_path)
    if stream_path and layout_path:
        return "（有译文：流式+排版）"
    if stream_path:
        return "（有译文：流式）"
    if layout_path:
        return "（有译文：排版）"
    return ""


REFERENCE_CONTEXT_RECOMMENDED_FILES = 6

# Send the reference corpus intact so the translator can use its terminology and style.
# Include the complete reference documents in the translation context. The
# application does not truncate them; a large corpus may exceed the model limit.
REFERENCE_TOTAL_WARNING_CHARS = 300000
REFERENCE_CONTEXT_CACHE_VERSION = "v3_full_direct_corpus"


def is_supported_reference_file(path: Path) -> bool:
    """参考文件既支持 MinerU 可解析格式，也支持本地 Markdown/文本。"""
    suffix = path.suffix.lower()
    return suffix in SUPPORTED_INPUT_EXTENSIONS or suffix in {".md", ".markdown", ".txt"}


def reference_file_cache_key(path: Path) -> str:
    """根据路径、大小和修改时间生成参考文件解析缓存键。"""
    stat = path.stat()
    payload = f"{path.resolve()}|{stat.st_size}|{stat.st_mtime_ns}|{REFERENCE_CONTEXT_CACHE_VERSION}"
    return hashlib.sha256(payload.encode("utf-8", errors="replace")).hexdigest()[:24]


def parse_reference_file_to_markdown(path: Path, cache_root: Path, log) -> str:
    """
    将参考文件解析为 Markdown。

    设计说明：
    1. Markdown/文本直接读取；
    2. Word/HTML 等优先用本地 Pandoc 转 Markdown；
    3. PDF、图片、Office 等 MinerU 支持格式在本地转换不可用时走 MinerU 精准解析；
    4. 所有结果写入缓存，避免同一参考文献反复解析。
    """
    path = path.expanduser().resolve()
    if not path.exists() or not path.is_file():
        raise MinerUError(f"参考文件不存在: {path}")
    if not is_supported_reference_file(path):
        raise MinerUError(f"暂不支持作为参考文件的格式: {path.suffix}")

    cache_root.mkdir(parents=True, exist_ok=True)
    digest = reference_file_cache_key(path)
    cache_path = cache_root / f"{digest}.reference.md"
    if cache_path.exists():
        log(f"参考语料已命中缓存：{path.name}")
        return cache_path.read_text(encoding="utf-8", errors="replace")

    suffix = path.suffix.lower()
    if suffix in {".md", ".markdown", ".txt"}:
        text = path.read_text(encoding="utf-8", errors="replace")
        cache_path.write_text(text, encoding="utf-8")
        return text

    pandoc = find_pandoc()
    if pandoc and suffix in PANDOC_OFFICE_SUFFIXES:
        temp_path = cache_root / f"{digest}.pandoc.md"
        try:
            log(f"正在提取参考语料内容：{path.name}")
            subprocess.run(
                [str(pandoc), str(path), "-t", "markdown", "--wrap=none", "-o", str(temp_path)],
                cwd=str(path.parent),
                check=True,
                capture_output=True,
                text=True,
                timeout=120,
                **hidden_subprocess_kwargs(),
            )
            if temp_path.exists():
                text = temp_path.read_text(encoding="utf-8", errors="replace")
                cache_path.write_text(text, encoding="utf-8")
                return text
        except Exception as exc:
            log(f"本地格式提取未成功，尝试使用 MinerU 解析参考文件：{path.name}")
        finally:
            try:
                temp_path.unlink()
            except OSError:
                pass

    if suffix not in SUPPORTED_INPUT_EXTENSIONS:
        raise MinerUError(f"无法解析参考文件格式: {path.suffix}")

    token = load_mineru_token()
    output_dir = cache_root / f"{digest}.mineru"
    output_dir.mkdir(parents=True, exist_ok=True)
    parsed_path = output_dir / "full.reference.md"
    if parsed_path.exists():
        text = parsed_path.read_text(encoding="utf-8", errors="replace")
        cache_path.write_text(text, encoding="utf-8")
        return text

    log(f"正在解析参考语料：{path.name}")
    options = ParseOptions(
        model_version="MinerU-HTML" if suffix in {".html", ".htm"} else DEFAULT_MODEL_VERSION,
        enable_table=True,
        enable_formula=True,
        is_ocr=False,
    )
    batch_id, upload_url, _data_id = submit_precise_file(path, options, token)
    http_put_file(upload_url, path, log=log)
    result_item = poll_precise_result(batch_id, options, token, log)
    markdown, _zip_url, extract_dir = extract_markdown_from_zip(result_item, output_dir, log)
    cleaned, _image_records = simplify_markdown_images(markdown, output_dir, [extract_dir])
    parsed_path.write_text(cleaned, encoding="utf-8")
    cache_path.write_text(cleaned, encoding="utf-8")
    return cleaned


def reference_context_cache_key(paths: list[str], target_language: str, model: str) -> str:
    """
    生成参考语料包缓存键。

    说明：
    1. 现在参考语料不再先让模型提炼，所以理论上不依赖 model；
    2. 但保留 model 参数以兼容既有调用签名，避免牵连更多代码；
    3. 文件路径、大小、修改时间或目标语言变化后，会自动重新生成缓存。
    """
    parts = [REFERENCE_CONTEXT_CACHE_VERSION, target_language]
    for raw_path in paths:
        path = Path(raw_path).expanduser().resolve()
        try:
            stat = path.stat()
            parts.append(f"{path}|{stat.st_size}|{stat.st_mtime_ns}")
        except OSError:
            parts.append(str(path))
    return hashlib.sha256("\n".join(parts).encode("utf-8", errors="replace")).hexdigest()[:24]


def compact_reference_markdown(text: str) -> str:
    """
    对参考语料做轻量清洗，但不做语义提炼。

    设计原则：
    - 保留原始句式、搭配、语体和段落节奏；
    - 只移除过度影响上下文效率的图片、HTML 细节和连续空白；
    - 不改写、不总结、不让模型提前“二手加工”语感。
    """
    text = text or ""
    text = re.sub(r"!\[[^\]]*\]\([^)]+\)", "", text)
    text = re.sub(r"<img\b[^>]*>", "", text, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r"<details\b[^>]*>.*?</details>", "", text, flags=re.IGNORECASE | re.DOTALL)
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()


def full_reference_document(text: str) -> str:
    """
    返回完整参考文献正文。

    设计原则：
    1. 不抽样、不截断，避免破坏完整语篇、句子节奏和上下文推进；
    2. 只做轻量 Markdown 噪声清理，不做语义压缩；
    3. 让模型在高维语义空间中直接阅读完整参考文献，形成更自然的领域语感与文体感。
    """
    return compact_reference_markdown(text)


def build_reference_translation_context(
    reference_paths: list[str],
    config: AITranslateConfig,
    cache_root: Path,
    log,
    target_language: str,
    should_stop=None,
) -> str:
    """
    解析参考文件，并直接构造“完整参考语料包”加入翻译上下文。

    批量翻译中的多个工作线程会共享同一个语料缓存，因此整个缓存生成过程
    必须按缓存键串行化；命中缓存后仅执行一次磁盘读取。
    """
    unique_paths: list[str] = []
    seen = set()
    for raw_path in reference_paths or []:
        text = str(raw_path or "").strip()
        if not text:
            continue
        normalized = str(Path(text).expanduser().resolve())
        if normalized in seen:
            continue
        seen.add(normalized)
        unique_paths.append(normalized)

    if not unique_paths:
        return ""

    if len(unique_paths) > REFERENCE_CONTEXT_RECOMMENDED_FILES:
        log(
            f"已选择 {len(unique_paths)} 个参考文件；建议通常控制在 "
            f"{REFERENCE_CONTEXT_RECOMMENDED_FILES} 个以内。程序会继续解析并完整加入全部参考文件，"
            "但参考语料过长时可能超过所选模型的上下文窗口。"
        )

    context_dir = cache_root / "reference_context"
    context_dir.mkdir(parents=True, exist_ok=True)
    cache_key = reference_context_cache_key(unique_paths, target_language, config.model)
    corpus_path = context_dir / f"reference_corpus.{cache_key}.md"

    with reference_cache_lock(f"corpus:{cache_key}", should_stop=should_stop):
        if should_stop and should_stop():
            raise MinerUError("用户已停止翻译。")
        if corpus_path.exists():
            log("参考语料已就绪（已命中缓存）。")
            return corpus_path.read_text(encoding="utf-8", errors="replace")

        parsed_items: list[tuple[str, str]] = []
        parse_cache = context_dir / "parsed"
        for index, raw_path in enumerate(unique_paths, 1):
            if should_stop and should_stop():
                raise MinerUError("用户已停止翻译。")
            path = Path(raw_path)
            log(f"正在解析参考文件（{index}/{len(unique_paths)}）：{path.name}")
            parsed = parse_reference_file_to_markdown(path, parse_cache, log)
            if should_stop and should_stop():
                raise MinerUError("用户已停止翻译。")
            if parsed.strip():
                parsed_items.append((path.name, parsed))

        if not parsed_items:
            return ""

        sections: list[str] = []
        total_reference_chars = 0
        for index, (name, text) in enumerate(parsed_items, 1):
            # 完整加入参考文献，不抽样、不截断，最大限度保留篇章语境和表达节奏。
            document = full_reference_document(text)
            if not document.strip():
                continue
            total_reference_chars += len(document)
            sections.append(
                "\n\n"
                f"===== Full reference corpus {index}: {name} =====\n"
                f"{document}\n"
                f"===== End full reference corpus {index}: {name} ====="
            )

        if total_reference_chars > REFERENCE_TOTAL_WARNING_CHARS:
            log(
                f"完整参考语料约 {total_reference_chars} 字符，可能超过部分模型的上下文窗口。"
                "程序不会主动截断；如翻译接口报错或模型明显忽略前文，请换用更长上下文模型或减少参考文件。"
            )

        if not sections:
            return ""

        corpus = (
            f"【Direct Reference Corpus for Translation into {target_language}】\n"
            "The following materials are complete user-selected target-language/domain reference papers. "
            "They are provided directly, without prior AI summarization or sampling, so that the translator can absorb terminology, collocations, register, sentence rhythm, discourse flow, and academic expression habits from intact discourse.\n\n"
            "Use policy:\n"
            "1. Treat the reference corpus as soft stylistic and terminological evidence only.\n"
            "2. The source document to be translated has absolute priority over the reference corpus.\n"
            "3. Do not import facts, claims, citations, data, mechanisms, experimental conditions, or certainty levels from the references.\n"
            "4. Do not copy or patchwrite reference sentences; avoid plagiarism and excessive imitation.\n"
            "5. If the reference style conflicts with the source meaning, preserve the source meaning.\n"
            "6. Prefer broadly accepted academic usage visible across the corpus rather than idiosyncratic wording from one paper.\n"
            "7. The real translation source will appear after this corpus; translate that source, not the references.\n"
            + "".join(sections)
        )

        write_text_atomic(corpus_path, corpus)
        return corpus


def reference_context_instruction(reference_context: str) -> str:
    if not reference_context.strip():
        return ""
    return (
        "An optional direct target-journal/domain reference corpus is provided by the user. "
        "Read it as soft evidence for terminology, collocations, register, sentence rhythm, and discourse conventions. "
        "The source paper to be translated always has higher priority than the reference corpus. "
        "Never let the references override the source paper's meaning, logic, data, formulas, citations, uncertainty, or claims. "
        "Do not import facts from the references and do not imitate reference sentences verbatim."
    )


def build_translation_context(markdown: str, config: AITranslateConfig, cache_path: Path, log, target_language: str) -> str:
    if cache_path.exists():
        return cache_path.read_text(encoding="utf-8", errors="replace")
    sample = markdown[:24000]
    prompt = (
        f"Please read the following MinerU-parsed Markdown excerpt and prepare a concise translation guide for translating the paper into {target_language}.\n"
        f"{target_language_instruction(target_language)}\n"
        "The excerpt may contain parsing/OCR/Markdown/LaTeX defects. Identify the research field, style requirements, key terminology, and any formatting risks that should be handled carefully during translation.\n"
        "Include a terminology table mapping source terms to recommended target-language terms. Do not translate the full paper.\n\n"
        f"{sample}"
    )
    system_prompt = "You are a terminology reviewer for academic paper translation projects."
    custom_instruction_section = translation_custom_instruction_section(config)
    if custom_instruction_section:
        system_prompt += "\n\n" + custom_instruction_section
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": prompt},
    ]
    save_translation_request_audit(cache_path.parent, "流式-术语指南", config, messages, timeout=180)
    context = ai_chat_completion(
        config,
        messages,
        timeout=180,
        rate_limit_callback=log,
    )
    cache_path.write_text(context, encoding="utf-8")
    return context


def translate_markdown_text(
    markdown: str,
    config: AITranslateConfig,
    log,
    target_language: str = "简体中文",
    work_dir: Path | None = None,
    live_update=None,
    reasoning_update=None,
    reference_context: str = "",
    concurrency: int = 0,
) -> str:
    if work_dir is None:
        work_dir = WORKSPACE / ".translation_cache"
    if (work_dir / "full_context.enabled").exists():
        return translate_markdown_full_context(
            markdown,
            config,
            log,
            target_language,
            work_dir,
            live_update,
            reasoning_update,
            reference_context,
        )
    return translate_markdown_by_chunks(
        markdown,
        config,
        log,
        target_language,
        work_dir,
        live_update,
        reasoning_update,
        reference_context,
        concurrency,
    )


def _translate_markdown_by_chunks_serial_legacy(
    markdown: str,
    config: AITranslateConfig,
    log,
    target_language: str = "简体中文",
    work_dir: Path | None = None,
    live_update=None,
    reasoning_update=None,
    reference_context: str = "",
) -> str:
    chunks = split_markdown_for_translation(markdown)
    translated_chunks: list[str] = []
    work_dir = work_dir or WORKSPACE / ".translation_cache"
    work_dir.mkdir(parents=True, exist_ok=True)
    language_suffix = translation_language_suffix(target_language)
    context_path = work_dir / f"translation_context.{language_suffix}.md"
    context = build_translation_context(markdown, config, context_path, log, target_language)
    manifest = {
        "provider_id": config.provider_id,
        "base_url": config.base_url,
        "model": config.model,
        "target_language": target_language,
        "chunk_count": len(chunks),
        "reference_context_enabled": bool(reference_context.strip()),
        # 分块翻译模式下，每个 chunk 请求都会显式携带同一个 reference_context。
        # reference_context 本身来自共享缓存，避免每个分块或每篇批量文档重复解析参考文献。
        "reference_context_injected_per_chunk": bool(reference_context.strip()),
    }
    (work_dir / f"manifest.{language_suffix}.json").write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    system_prompt = (
        f"You are a professional academic translator and Markdown cleanup editor. Translate the provided MinerU-parsed Markdown into {target_language}. "
        f"{target_language_instruction(target_language)} "
        "The source Markdown may contain OCR/parsing defects, broken Markdown tables, imperfect LaTeX, misplaced spaces, duplicated line breaks, "
        "or formatting that does not fully reproduce the original paper. Use academic judgment to repair obvious Markdown/LaTeX/formatting defects while translating, "
        "but do not invent missing data, references, equations, captions, or conclusions. "
        "Preserve the document structure, headings, tables, code blocks, inline code, HTML tags, links, and image placeholders such as IMAGE_001. "
        "Keep anchor tags, ids, image references, and table structure intact whenever they are present. "
        "Be conservative about formatting: do not turn normal prose into headings, captions, centered text, bold text, lists, or tables unless the Markdown structure clearly indicates it. "
        "Act like a careful journal copy editor for layout cleanup: keep ordinary body paragraphs as body text with normal paragraph flow, repair obvious paragraph breaks, and separate captions from explanatory prose when MinerU has merged them. "
        "Figure/table captions should remain separate from body paragraphs; true captions such as 'Figure 1.', 'Fig. 1', 'Table 2.', or their translated equivalents should be caption-like blocks, while body paragraphs that merely mention figures or tables, such as 'Figure 1 shows ...' or 'Table 2 reports ...', must stay normal paragraphs. "
        "When Markdown/HTML export supports it, captions should be visually distinct from body text: centered where appropriate, slightly smaller than body text, and not followed by accidental centered body paragraphs. "
        "Preserve or normalize citation markers as academic citations; bracketed numeric references such as [1], [2-4], and [5,6] should remain citation tokens suitable for superscript styling in Word/PDF export rather than being translated into prose. "
        "Preserve equation-number references exactly: translate 'Eq. (16)' or 'Eqs. (3) and (4)' as normal equation references with the original parenthesized numbers, never as ~16!, ~3!, punctuation, or prose words. "
        "Hard output rule: this reader renders mathematics with MathJax. A formula is part of the paper's meaning, not formatting. Copy every mathematical expression verbatim, including its TeX body and delimiters such as \\(...\\), \\[...\\], $...$, and $$...$$. A naked TeX body such as R _ { 0 } is not renderable. Translate surrounding prose only; never rewrite, normalize, omit, or move a formula. "
        "Do not add explanations, comments, or code fences."
    )
    if is_epub_markdown(markdown):
        system_prompt += (
            " This is an EPUB source. EPUB Markdown attributes are structural data, not prose: "
            "preserve every `{#...}`, `{.class}`, `lang=`, `title=`, id, class, HTML comment, "
            "chapter marker, link target, image target, and attribute key/value byte-for-byte. "
            "Never translate attribute names or values, never remove square brackets around "
            "bracketed spans such as `[title]{#id lang=\"en\"}`, and never move an attribute "
            "block to ordinary text."
        )
    if reference_context.strip():
        system_prompt += " " + reference_context_instruction(reference_context)
    custom_instruction_section = translation_custom_instruction_section(config)
    if custom_instruction_section:
        system_prompt += "\n\n" + custom_instruction_section
    for index, chunk in enumerate(chunks, 1):
        chunk_src_path = work_dir / f"chunk_{index:04d}.source.md"
        chunk_out_path = work_dir / f"chunk_{index:04d}.{language_suffix}.translated.md"
        chunk_src_path.write_text(chunk, encoding="utf-8")
        if chunk_out_path.exists():
            log(f"分块 {index}/{len(chunks)} 已有翻译缓存，直接复用。")
            cached_chunk = chunk_out_path.read_text(encoding="utf-8", errors="replace")
            cached_chunk = repair_equation_reference_translation(chunk, cached_chunk)
            formula_issue = math_expression_integrity_issue(chunk, cached_chunk)
            if formula_issue:
                log(f"分块 {index}/{len(chunks)}：检测到公式存在微小差异，已保留已有译文。")
            translated_chunks.append(cached_chunk.rstrip() + "\n")
            if live_update:
                live_update("\n".join(translated_chunks).strip() + "\n")
            continue
        log(f"正在翻译第 {index}/{len(chunks)} 块（约 {len(chunk)} 字符）…")
        reference_section = (
            f"Optional direct target-journal/domain reference corpus loaded from cache. Read for style and terminology only; translate the source chunk after it:\n\n{reference_context}\n\n"
            if reference_context.strip()
            else ""
        )
        if reference_context.strip():
            log(f"第 {index}/{len(chunks)} 块已加载参考语料（约 {len(reference_context)} 字符）。")
        user_prompt = (
            f"{reference_section}"
            f"Translation guide for this paper:\n\n{context}\n\n"
            f"Translate the following MinerU-parsed Markdown source chunk into {target_language}. {target_language_instruction(target_language)} "
            "While translating, gently fix clear Markdown/table formatting problems caused by parsing, such as malformed table separators, "
            "obvious heading/caption formatting issues, and excessive line breaks. Preserve meaning and do not over-correct uncertain scientific content. "
            "Do not introduce new centering, bolding, heading levels, list markers, or caption formatting unless the source Markdown clearly marks that block as such. "
            "As a journal copy editor, split true figure/table captions away from surrounding explanatory paragraphs if MinerU merged them, and restore the surrounding explanation to normal paragraph formatting. "
            "Keep figure/table captions separate from body paragraphs; a paragraph that discusses a figure/table is still normal prose. Captions may be centered or caption-styled only when the block is clearly a caption. "
            "Keep numeric reference citations such as [1], [2-4], and [5,6] intact as citation markers suitable for superscript export styling. "
            "Hard output rule: the reader uses MathJax. Copy every source formula verbatim with its TeX body and delimiters; never emit a naked TeX body such as R _ { 0 }. Translate only the surrounding natural language. Keep equation-number references such as Eq. (16) and Eqs. (3) and (4) as equation references with the original parenthesized numbers; never turn them into ~16!, ~3!, punctuation, or prose words. "
            "Use the optional direct reference corpus only for terminology, collocation, register, rhythm, and target-field conventions; never distort the source meaning to fit it. "
            "Preserve image references, HTML tags, links, anchor ids, and all placeholders. Do not omit content and do not add explanations.\n\n"
            "===== BEGIN SOURCE CHUNK TO TRANSLATE =====\n"
            f"{chunk}\n"
            "===== END SOURCE CHUNK TO TRANSLATE ====="
        )
        current_chunk_parts: list[str] = []
        last_live_emit = 0.0

        def stream_current_chunk(delta: str):
            nonlocal last_live_emit
            current_chunk_parts.append(delta)
            now = time.monotonic()
            if live_update and now - last_live_emit >= 0.08:
                preview = "\n".join(translated_chunks + ["".join(current_chunk_parts)]).strip() + "\n"
                live_update(preview)
                last_live_emit = now

        messages = [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ]
        save_translation_request_audit(work_dir, f"流式-分块翻译-第{index}块", config, messages, timeout=180)
        translated = ai_chat_completion(
            config,
            messages,
            stream_callback=stream_current_chunk,
            reasoning_callback=reasoning_update,
        )
        translated = normalize_translated_inline_html(translated)
        translated = repair_equation_reference_translation(chunk, translated)
        formula_issue = math_expression_integrity_issue(chunk, translated)
        formula_retry_issue = math_expression_retry_issue(chunk, translated)
        if formula_issue and not formula_retry_issue:
            log(
                f"第 {index}/{len(chunks)} 块公式写法略有差异："
                f"{formula_issue}；数学结构完整，不执行整块重试。"
            )
        if formula_retry_issue:
            log(f"第 {index}/{len(chunks)} 块检测到公式结构差异，正在自动校正重试…")
            retry_prompt = (
                "补充修正刚才的完整译文。请重新输出完整译文，不要省略任何段落。"
                "尤其注意：阅读器使用 MathJax；每个源公式必须逐字保留其 TeX 和原始定界符，"
                "不得输出裸 TeX。\n\n"
                f"===== BEGIN SOURCE CHUNK TO TRANSLATE =====\n{chunk}\n===== END SOURCE CHUNK TO TRANSLATE ====="
            )
            retry_messages = [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
                {"role": "user", "content": retry_prompt},
            ]
            save_translation_request_audit(work_dir, f"流式-分块翻译-第{index}块-重试", config, retry_messages, timeout=180)
            translated = ai_chat_completion(
                config,
                retry_messages,
                stream_callback=stream_current_chunk,
                reasoning_callback=reasoning_update,
            )
            translated = normalize_translated_inline_html(translated)
            translated = repair_equation_reference_translation(chunk, translated)
            formula_issue = math_expression_integrity_issue(chunk, translated)
            retry_issue = math_expression_retry_issue(chunk, translated)
            if retry_issue:
                log(f"第 {index}/{len(chunks)} 块校正后仍有差异，已保留模型译文供人工核对。")
            elif formula_issue:
                log(f"第 {index}/{len(chunks)} 块重试后公式微调完成。")
        chunk_out_path.write_text(translated, encoding="utf-8")
        translated_chunks.append(translated.rstrip() + "\n")
        if live_update:
            live_update("\n".join(translated_chunks).strip() + "\n")
    # 分块翻译通常不会主动使用结束 marker，但不同模型仍可能在尾部附带“End marker”
    # “结束标记”等痕迹，这里统一做一次温和清理。
    return clean_translation_tail_artifacts("\n".join(translated_chunks))


def translate_markdown_by_chunks(
    markdown: str,
    config: AITranslateConfig,
    log,
    target_language: str = "简体中文",
    work_dir: Path | None = None,
    live_update=None,
    reasoning_update=None,
    reference_context: str = "",
    concurrency: int = 0,
) -> str:
    """Translate independent Markdown chunks concurrently and commit them in source order."""
    concurrency = normalize_translation_request_concurrency(
        config.provider_id,
        concurrency,
    )
    chunks = split_markdown_for_translation(markdown)
    work_dir = work_dir or WORKSPACE / ".translation_cache"
    work_dir.mkdir(parents=True, exist_ok=True)
    language_suffix = translation_language_suffix(target_language)
    context_path = work_dir / f"translation_context.{language_suffix}.md"
    context = build_translation_context(markdown, config, context_path, log, target_language)
    identity_parts = [
        STREAM_CHUNK_PROTOCOL,
        markdown,
        str(config.provider_id),
        str(config.base_url),
        str(config.model),
        target_language,
        hashlib.sha256(reference_context.encode("utf-8")).hexdigest(),
    ]
    custom_instruction_hash = translation_custom_instruction_hash(config)
    if custom_instruction_hash:
        identity_parts.append(custom_instruction_hash)
    document_identity = hashlib.sha256("\u241f".join(identity_parts).encode("utf-8")).hexdigest()

    def atomic_write_text(path: Path, text: str) -> None:
        temporary = path.with_name(f".{path.name}.{uuid.uuid4().hex}.tmp")
        try:
            temporary.write_text(text, encoding="utf-8")
            temporary.replace(path)
        finally:
            temporary.unlink(missing_ok=True)

    def atomic_write_json(path: Path, payload: dict) -> None:
        atomic_write_text(path, json.dumps(payload, ensure_ascii=False, indent=2))

    manifest = {
        "protocol": STREAM_CHUNK_PROTOCOL,
        "identity": document_identity,
        "provider_id": config.provider_id,
        "base_url": config.base_url,
        "model": config.model,
        "target_language": target_language,
        "chunk_count": len(chunks),
        "concurrency": concurrency,
        "max_rounds_per_chunk": STREAM_CONTINUATION_MAX_ROUNDS,
        "reference_context_enabled": bool(reference_context.strip()),
        "reference_context_injected_per_chunk": bool(reference_context.strip()),
    }
    atomic_write_json(work_dir / f"manifest.{language_suffix}.json", manifest)
    system_prompt = (
        f"You are a professional academic translator and Markdown cleanup editor. Translate the provided MinerU-parsed Markdown into {target_language}. "
        f"{target_language_instruction(target_language)} "
        "The source may contain OCR, Markdown, table, LaTeX, spacing, caption, or line-break defects. "
        "Repair only clear formatting defects while translating; never invent scientific content. "
        "Preserve headings, tables, code, HTML, links, anchors, image placeholders, citations, equation numbers, and every formula with its original TeX delimiters. "
        "Keep captions separate from explanatory prose and do not introduce unsupported styling. "
        "Output only translated Markdown and the exact completion token when the whole source chunk is complete."
    )
    if is_epub_markdown(markdown):
        system_prompt += (
            " This is an EPUB source. Preserve EPUB structure exactly: never translate or "
            "remove Markdown/HTML attribute keys or values, never drop square brackets from "
            "`[title]{#id lang=\"en\"}`, and keep chapter comments, links, image targets, ids, "
            "classes, and language attributes byte-for-byte."
        )
    if reference_context.strip():
        system_prompt += " " + reference_context_instruction(reference_context)
    custom_instruction_section = translation_custom_instruction_section(config)
    if custom_instruction_section:
        system_prompt += "\n\n" + custom_instruction_section

    coordinator_lock = threading.Lock()
    committed: list[str] = []
    next_commit_index = 0

    def translate_one(index: int, chunk: str) -> tuple[int, str, bool]:
        nonlocal next_commit_index
        chunk_number = index + 1
        chunk_key = hashlib.sha256(chunk.encode("utf-8")).hexdigest()
        cache_identity = hashlib.sha256(
            f"{document_identity}\u241f{chunk_number}\u241f{chunk_key}".encode("utf-8")
        ).hexdigest()
        source_path = work_dir / f"chunk_{chunk_number:04d}.source.md"
        output_path = work_dir / f"chunk_{chunk_number:04d}.{language_suffix}.translated.md"
        meta_path = output_path.with_suffix(output_path.suffix + ".json")
        state_path = output_path.with_suffix(output_path.suffix + ".state.json")
        atomic_write_text(source_path, chunk)

        cache_meta = None
        if meta_path.exists():
            try:
                cache_meta = json.loads(meta_path.read_text(encoding="utf-8"))
            except (OSError, ValueError, TypeError):
                cache_meta = None
        if (
            output_path.exists()
            and isinstance(cache_meta, dict)
            and cache_meta.get("protocol") == STREAM_CHUNK_PROTOCOL
            and cache_meta.get("identity") == cache_identity
            and cache_meta.get("complete") is True
        ):
            translated = output_path.read_text(encoding="utf-8", errors="replace")
            translated = repair_equation_reference_translation(chunk, translated).strip()
            formula_issue = math_expression_integrity_issue(chunk, translated)
            if formula_issue:
                log(f"分块 {chunk_number}/{len(chunks)} 缓存公式微调：{formula_issue}；已保留已有译文。")
            return index, translated, True

        marker = stable_translation_marker(
            f"{cache_identity}\n{chunk_number}\n{chunk_key}\n{chunk}"
        )
        reference_section = (
            "Optional direct target-journal/domain reference corpus loaded from cache. "
            "Read it for style and terminology only; translate the source chunk after it:\n\n"
            f"{reference_context}\n\n"
            if reference_context.strip()
            else ""
        )
        user_prompt = (
            f"{reference_section}"
            f"Translation guide for this paper:\n\n{context}\n\n"
            f"Translate the following MinerU-parsed Markdown source chunk completely into {target_language}. "
            f"{target_language_instruction(target_language)} "
            "Preserve all content, Markdown structures, formulas and delimiters, image references, HTML, links, anchors, placeholders, citations, and equation numbers. "
            "Repair only clear parsing and formatting defects. Do not add explanations. "
            "If one response is insufficient, translate as far as possible; the application will continue this same chunk conversation.\n\n"
            f"{completion_token_instruction(marker)}\n\n"
            "===== BEGIN SOURCE CHUNK TO TRANSLATE =====\n"
            f"{chunk}\n"
            "===== END SOURCE CHUNK TO TRANSLATE =====\n\n"
            f"Final reminder:\n{completion_token_instruction(marker)}"
        )
        messages = [
            {"role": "system", "content": system_prompt + " " + completion_token_instruction(marker)},
            {"role": "user", "content": user_prompt},
        ]
        parts: list[str] = []
        start_round = 1
        empty_rounds = 0
        repeated_rounds = 0
        if state_path.exists():
            try:
                state = json.loads(state_path.read_text(encoding="utf-8"))
            except (OSError, ValueError, TypeError):
                state = None
            if (
                isinstance(state, dict)
                and state.get("protocol") == STREAM_CHUNK_PROTOCOL
                and state.get("identity") == cache_identity
                and state.get("marker")
                and isinstance(state.get("messages"), list)
                and isinstance(state.get("parts"), list)
                and state.get("complete") is not True
                and 0 < int(state.get("round") or 0) < STREAM_CONTINUATION_MAX_ROUNDS
            ):
                marker = str(state["marker"])
                messages = [
                    {"role": str(item["role"]), "content": str(item["content"])}
                    for item in state["messages"]
                    if isinstance(item, dict)
                    and item.get("role") in {"system", "user", "assistant"}
                    and isinstance(item.get("content"), str)
                ]
                parts = [str(part) for part in state["parts"] if str(part).strip()]
                start_round = int(state["round"]) + 1
                empty_rounds = int(state.get("empty_rounds") or 0)
                repeated_rounds = int(state.get("repeated_rounds") or 0)
                if messages and messages[-1]["role"] == "assistant":
                    messages.append(
                        {
                            "role": "user",
                            "content": (
                                f"Continue translating from the exact previous interruption point into {target_language}. "
                                "Do not repeat translated content. Output only the remaining translated Markdown. "
                                f"When this source chunk is complete, output the exact token on a separate final line.\n"
                                f"{completion_token_instruction(marker)}"
                            ),
                        }
                    )
                log(f"分块 {chunk_number}/{len(chunks)} 从第 {start_round} 轮继续翻译。")

        complete = False
        for round_number in range(start_round, STREAM_CONTINUATION_MAX_ROUNDS + 1):
            log(
                f"正在翻译分块 {chunk_number}/{len(chunks)}（第 {round_number} 轮，"
                f"约 {len(chunk)} 字符）…"
            )
            streamed_parts: list[str] = []
            last_live_emit = 0.0

            def stream_current_round(delta: str):
                nonlocal last_live_emit
                streamed_parts.append(delta)
                now = time.monotonic()
                if not live_update or now - last_live_emit < 0.08:
                    return
                with coordinator_lock:
                    if index != next_commit_index:
                        return
                    preview = "\n\n".join(
                        committed + ["\n".join(parts), "".join(streamed_parts)]
                    ).strip() + "\n"
                live_update(preview)
                last_live_emit = now

            save_translation_request_audit(
                work_dir,
                f"流式-分块翻译-第{chunk_number}块-第{round_number}轮",
                config,
                messages,
                timeout=300,
            )
            content = ai_chat_completion(
                config,
                messages,
                timeout=300,
                stream_callback=stream_current_round,
                reasoning_callback=reasoning_update,
                rate_limit_callback=log,
            )
            normalized_content = str(content or "").strip()
            empty_rounds = 0 if normalized_content else empty_rounds + 1
            repeated_rounds = (
                repeated_rounds + 1
                if normalized_content and normalized_content == (parts[-1].strip() if parts else "")
                else 0
            )
            if empty_rounds >= 3 or repeated_rounds >= 3:
                raise MinerUError(
                    f"翻译分块 {chunk_number} 连续多轮没有有效进展，已保留续写缓存。"
                )
            if normalized_content:
                parts.append(str(content))
            messages.append({"role": "assistant", "content": str(content or "")})
            complete = marker_detected("\n".join(parts), marker)
            if not complete:
                messages.append(
                    {
                        "role": "user",
                        "content": (
                            f"Continue translating from the exact previous interruption point into {target_language}. "
                            "Do not repeat any translated content. Preserve Markdown, formulas, placeholders, links, and anchors. "
                            "Output only the remaining translation. "
                            f"When this source chunk is complete, output the exact token on a separate final line.\n"
                            f"{completion_token_instruction(marker)}"
                        ),
                    }
                )
            atomic_write_json(
                state_path,
                {
                    "protocol": STREAM_CHUNK_PROTOCOL,
                    "identity": cache_identity,
                    "marker": marker,
                    "round": round_number,
                    "complete": complete,
                    "empty_rounds": empty_rounds,
                    "repeated_rounds": repeated_rounds,
                    "parts": parts,
                    "messages": messages,
                },
            )
            if complete:
                break

        if not complete:
            raise MinerUError(
                f"翻译分块 {chunk_number} 达到 {STREAM_CONTINUATION_MAX_ROUNDS} 轮"
                f"仍未检测到完成标记：{marker}"
            )
        translated = clean_translation_tail_artifacts("\n".join(parts), marker)
        translated = normalize_translated_inline_html(translated)
        translated = repair_equation_reference_translation(chunk, translated).strip()
        formula_issue = math_expression_integrity_issue(chunk, translated)
        formula_retry_issue = math_expression_retry_issue(chunk, translated)
        if formula_issue and not formula_retry_issue:
            log(
                f"分块 {chunk_number}/{len(chunks)} 公式写法略有差异："
                f"{formula_issue}；数学结构完整，不执行整块重试。"
            )
        if formula_retry_issue:
            log(
                f"分块 {chunk_number}/{len(chunks)} 检测到公式结构差异，正在执行自动校正重试…"
            )
            retry_messages = [
                {
                    "role": "system",
                    "content": system_prompt + " " + completion_token_instruction(marker),
                },
                {
                    "role": "user",
                    "content": (
                        "Re-translate the complete source chunk below. The previous complete draft damaged or omitted mathematical markup. "
                        "Output the entire corrected translation, not a patch or explanation. "
                        "Copy every source formula verbatim with its TeX body and original delimiters. Do not omit any paragraph.\n\n"
                        f"{completion_token_instruction(marker)}\n\n"
                        "===== BEGIN SOURCE CHUNK TO TRANSLATE =====\n"
                        f"{chunk}\n"
                        "===== END SOURCE CHUNK TO TRANSLATE ====="
                    ),
                },
            ]
            save_translation_request_audit(
                work_dir,
                f"流式-分块严格重试-第{chunk_number}块",
                config,
                retry_messages,
                timeout=300,
            )
            retry_text = ai_chat_completion(
                config,
                retry_messages,
                timeout=300,
                reasoning_callback=reasoning_update,
                rate_limit_callback=log,
            )
            if marker_detected(retry_text, marker):
                corrected = clean_translation_tail_artifacts(retry_text, marker)
                corrected = normalize_translated_inline_html(corrected)
                corrected = repair_equation_reference_translation(chunk, corrected).strip()
                retry_issue = math_expression_retry_issue(chunk, corrected)
                if retry_issue:
                    log(f"校正后仍有微小差异，已保留重试前的完整译文供核对。")
                else:
                    translated = corrected
                    formula_issue = math_expression_integrity_issue(chunk, corrected)
                    if formula_issue:
                        log(f"重试后公式格式微调完成。")
            else:
                log("校正重试未返回完整标记，已保留重试前的完整译文。")
        atomic_write_text(output_path, translated + "\n")
        atomic_write_json(
            meta_path,
            {
                "protocol": STREAM_CHUNK_PROTOCOL,
                "identity": cache_identity,
                "complete": True,
                "marker": marker,
                "rounds": len(parts),
                "formula_issue": formula_issue,
            },
        )
        state_path.unlink(missing_ok=True)
        return index, translated, False

    future_results: dict[int, tuple[str, bool]] = {}
    failures: list[tuple[int, BaseException]] = []
    with ThreadPoolExecutor(
        max_workers=min(concurrency, max(1, len(chunks))),
        thread_name_prefix="stream-translation",
    ) as executor:
        futures = {
            executor.submit(translate_one, index, chunk): index
            for index, chunk in enumerate(chunks)
        }
        for future in as_completed(futures):
            index = futures[future]
            try:
                result_index, translated, cached = future.result()
                future_results[result_index] = (translated, cached)
            except BaseException as exc:
                failures.append((index, exc))
                continue
            previews: list[str] = []
            with coordinator_lock:
                while next_commit_index in future_results:
                    translated, _cached = future_results.pop(next_commit_index)
                    committed.append(translated.rstrip())
                    next_commit_index += 1
                    previews.append("\n\n".join(committed).strip() + "\n")
            if live_update:
                for preview in previews:
                    live_update(preview)

    if failures:
        failures.sort(key=lambda item: item[0])
        raise failures[0][1]
    final = clean_translation_tail_artifacts("\n\n".join(committed))
    if live_update:
        live_update(final)
    return final


def translate_markdown_full_context(
    markdown: str,
    config: AITranslateConfig,
    log,
    target_language: str,
    work_dir: Path,
    live_update=None,
    reasoning_update=None,
    reference_context: str = "",
) -> str:
    work_dir.mkdir(parents=True, exist_ok=True)
    marker = stable_translation_marker(markdown)
    language_suffix = translation_language_suffix(target_language)
    final_path = work_dir / f"full_context.{language_suffix}.translated.md"
    state_path = work_dir / f"full_context.{language_suffix}.state.json"
    transcript_path = work_dir / f"full_context.{language_suffix}.messages.json"
    partial_paths = sorted(work_dir.glob(f"full_context.{language_suffix}.part_*.md"))
    translated_parts = [path.read_text(encoding="utf-8", errors="replace") for path in partial_paths]
    if translated_parts and live_update:
        live_update(clean_translation_tail_artifacts("\n".join(translated_parts), marker))
    if translated_parts and marker_detected(translated_parts[-1], marker):
        cached_result = strip_translation_marker("\n".join(translated_parts), marker)
        formula_issue = math_expression_integrity_issue(markdown, cached_result)
        if formula_issue:
            log(f"全文翻译缓存公式微调：{formula_issue}；已保留已有译文。")
        return cached_result

    token_rules = completion_token_instruction(marker)
    system_prompt = (
        f"You are a professional academic paper translator and Markdown cleanup editor. Translate the full MinerU-parsed Markdown paper into {target_language}. "
        f"{target_language_instruction(target_language)} "
        "First understand the research field, core concepts, terminology chain, and argument structure. "
        "Keep terminology consistent and use formal, accurate, fluent academic style. "
        "The surrounding application is designed to continue automatically with follow-up requests if your reply ends before the full paper is finished, so you should keep translating as far as possible instead of stopping early because of length concerns. "
        "Never mention token limits, context windows, truncation, or capacity constraints. "
        "Never ask the user to split the paper into sections or send it again in smaller chunks. "
        "If the source is long, continue translating directly and output only the translated Markdown, not a refusal, warning, or meta explanation. "
        "MinerU parsing may introduce OCR errors, broken Markdown tables, imperfect LaTeX, misplaced spaces, duplicated line breaks, or incomplete formatting. "
        "Repair clear Markdown/LaTeX/formatting defects while translating, as a careful human editor would, but do not infer or invent missing scientific content. "
        "Preserve Markdown structure, headings, tables, code blocks, inline code, HTML tags, links, anchor ids, and image placeholders. "
        "Be conservative about layout: do not turn ordinary prose into centered text, bold text, headings, captions, lists, or tables unless the original Markdown block clearly requires it. "
        "Act like a journal copy editor when MinerU has clearly damaged layout: repair paragraph breaks, separate captions from adjacent explanatory prose, restore body paragraphs to normal paragraph formatting, and avoid letting caption alignment leak into following paragraphs. "
        "Keep figure/table captions and body paragraphs distinct; true captions such as 'Figure 1.', 'Fig. 1', 'Table 2.', or their translated equivalents should remain caption-like blocks, while sentences like 'Figure 1 shows ...' or 'Table 2 reports ...' are normal body text, not captions. "
        "Where Markdown/HTML export supports it, true captions may be centered and slightly smaller than body text, but normal explanatory paragraphs should not be centered. "
        "Preserve numeric reference citations such as [1], [2-4], and [5,6] as citation markers suitable for superscript styling in Word/PDF export. "
        "Preserve equation-number references exactly: translate 'Eq. (16)' or 'Eqs. (3) and (4)' as normal equation references with the original parenthesized numbers, never as ~16!, ~3!, punctuation, or prose words. "
        "Hard output rule: this reader renders mathematics with MathJax. A formula is part of the paper's meaning, not formatting. Copy every mathematical expression verbatim, including its TeX body and delimiters such as \\(...\\), \\[...\\], $...$, and $$...$$. A naked TeX body such as R _ { 0 } is not renderable. Translate surrounding prose only; never rewrite, normalize, omit, or move a formula. If a symbol, citation, or numeric value is ambiguous, keep the original token rather than guessing. "
        "Do not omit content, add explanations, or wrap the whole output in a code fence. "
        f"{token_rules}"
    )
    if is_epub_markdown(markdown):
        system_prompt += (
            " This is an EPUB source. EPUB Markdown attributes are structural data, not prose: "
            "preserve every `{#...}`, `{.class}`, `lang=`, `title=`, id, class, HTML comment, "
            "chapter marker, link target, image target, and attribute key/value byte-for-byte. "
            "Never translate attribute names or values, never remove square brackets around "
            "bracketed spans such as `[title]{#id lang=\"en\"}`, and never move an attribute "
            "block to ordinary text."
        )
    if reference_context.strip():
        system_prompt += " " + reference_context_instruction(reference_context)
    custom_instruction_section = translation_custom_instruction_section(config)
    if custom_instruction_section:
        system_prompt += "\n\n" + custom_instruction_section
    reference_section = (
        f"Optional direct target-journal/domain reference corpus. Read for style and terminology only; translate the source paper after it:\n\n{reference_context}\n\n"
        if reference_context.strip()
        else ""
    )
    user_prompt = (
        f"Translate the following MinerU-parsed Markdown paper completely into {target_language}.\n\n"
        "Rules:\n"
        "1. Understand the research field and terminology before translating.\n"
        "2. Preserve all Markdown tables, formulas, image references, HTML anchor ids, and placeholders such as IMAGE_001.\n"
        "3. While translating, fix obvious parsing-related Markdown/formatting defects, including broken table syntax, excessive line breaks, and malformed captions. The reader uses MathJax: preserve every formula verbatim with its original TeX delimiters; never output naked TeX as ordinary text.\n"
        "4. Act like a journal copy editor for clear MinerU layout damage: split captions away from merged explanatory paragraphs, restore ordinary paragraphs to normal body formatting, and prevent caption centering from leaking into following body text.\n"
        "5. Be conservative about visual formatting: do not introduce new centering, bolding, heading levels, list markers, or caption formatting unless the original Markdown block clearly requires it.\n"
        "6. Keep figure/table captions and body paragraphs distinct; sentences that discuss figures or tables, such as 'Figure 1 shows ...' or 'Table 2 reports ...', must remain normal paragraphs.\n"
        "7. True figure/table captions may be caption-like blocks, centered when appropriate, and visually smaller than body text where the output format allows; surrounding explanatory paragraphs must stay normal body text.\n"
        "8. Preserve numeric reference citations such as [1], [2-4], and [5,6] as citation markers suitable for superscript styling in Word/PDF export.\n"
        "9. Preserve equation-number references such as Eq. (16) and Eqs. (3) and (4) as equation references with the original parenthesized numbers; never turn them into ~16!, ~3!, punctuation, or prose words.\n"
        "10. Do not over-correct: if a formula, number, citation, symbol, or sentence fragment is ambiguous, keep the original token or translate literally rather than guessing.\n"
        "11. Do not invent missing content, do not delete difficult content, and do not add translator notes or explanations.\n"
        "12. If an optional direct reference corpus is provided, use it only for terminology, collocation, register, rhythm, and target-field conventions; never distort the source meaning to fit it.\n"
        "13. The surrounding application will continue automatically if your reply ends before the full paper is finished, so keep translating as far as possible instead of stopping early because of length concerns.\n"
        "14. Never mention token limits, context windows, truncation, or capacity constraints, and never ask for chapter-by-chapter resubmission.\n"
        "15. The completion token is not a natural-language phrase. Never replace it with TRANSLATION COMPLETE or any similar words.\n\n"
        f"{token_rules}\n\n"
        f"{reference_section}"
        "Markdown to translate begins after this line. Translate only this Markdown source, not the reference corpus above.\n\n"
        "===== BEGIN SOURCE MARKDOWN TO TRANSLATE =====\n"
        f"{markdown}\n"
        "===== END SOURCE MARKDOWN TO TRANSLATE =====\n\n"
        "Final reminder:\n"
        f"{token_rules}\n"
    )
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt},
    ]
    for part in translated_parts:
        messages.append({"role": "assistant", "content": part})
        if marker not in part:
            messages.append(
                {
                    "role": "user",
                    "content": (
                        f"Continue translating from the previous interruption point into {target_language}. "
                        "Do not repeat translated content. Continue preserving Markdown, formulas, image placeholders, links, and anchor ids. "
                        "Keep gently repairing only obvious MinerU parsing/Markdown/LaTeX defects, without inventing missing scientific content. "
                        "The surrounding application will continue automatically if this reply ends before the full paper is finished, so keep translating as far as possible instead of stopping early because of length concerns. "
                        "Never mention token limits, context windows, truncation, or capacity constraints. "
                        "Never ask the user to split the paper into sections or resend parts. "
                        "Continue directly with the remaining translation and output only translated Markdown plus the final completion token when finished. "
                        "When the full text is finally completed, output the exact machine-detection completion token on a separate final line. "
                        "Do not replace it with TRANSLATION COMPLETE or any other words. "
                        f"{completion_token_instruction(marker)}"
                    ),
                }
            )

    max_rounds = STREAM_CONTINUATION_MAX_ROUNDS
    start_round = len(translated_parts) + 1
    for round_index in range(start_round, max_rounds + 1):
        log(f"正在进行全文连续翻译（第 {round_index} 轮）…")
        current_round_parts: list[str] = []
        last_live_emit = 0.0

        def stream_current_round(delta: str):
            nonlocal last_live_emit
            current_round_parts.append(delta)
            now = time.monotonic()
            if live_update and now - last_live_emit >= 0.08:
                live_update(clean_translation_tail_artifacts("\n".join(translated_parts + ["".join(current_round_parts)]), marker))
                last_live_emit = now

        save_translation_request_audit(work_dir, f"流式-全文连续翻译-第{round_index}轮", config, messages, timeout=300)
        content = ai_chat_completion(
            config,
            messages,
            timeout=300,
            stream_callback=stream_current_round,
            reasoning_callback=reasoning_update,
            rate_limit_callback=log,
        )
        content = normalize_translated_inline_html(content)
        content = repair_equation_reference_translation(markdown, content)
        part_path = work_dir / f"full_context.{language_suffix}.part_{round_index:04d}.md"
        part_path.write_text(content, encoding="utf-8")
        translated_parts.append(content)
        if live_update:
            live_update(clean_translation_tail_artifacts("\n".join(translated_parts), marker))
        messages.append({"role": "assistant", "content": content})
        transcript_path.write_text(json.dumps(messages, ensure_ascii=False, indent=2), encoding="utf-8")
        state_path.write_text(
            json.dumps(
                {
                    "provider_id": config.provider_id,
                    "base_url": config.base_url,
                    "model": config.model,
                    "target_language": target_language,
                    "marker": marker,
                    "round": round_index,
                    "reference_context_enabled": bool(reference_context.strip()),
                    "complete": marker_detected(content, marker),
                },
                ensure_ascii=False,
                indent=2,
            ),
            encoding="utf-8",
        )
        if marker_detected(content, marker):
            # 先删除 marker，再清理模型可能泄漏的“结束标记”“End marker”等尾部痕迹。
            result = strip_translation_marker("\n".join(translated_parts), marker)
            result = clean_translation_tail_artifacts(result, marker)
            result = normalize_translated_inline_html(result)
            result = repair_equation_reference_translation(markdown, result)
            formula_issue = math_expression_integrity_issue(markdown, result)
            if formula_issue:
                log(f"全文翻译公式微调：{formula_issue}；已保留模型译文供人工核对。")
            final_path.write_text(result, encoding="utf-8")
            return result
        messages.append(
                {
                    "role": "user",
                    "content": (
                        f"Continue translating from the previous interruption point into {target_language}. "
                        "Do not repeat translated content. Continue preserving Markdown, formulas, image placeholders, links, and anchor ids. "
                        "Keep gently repairing only obvious MinerU parsing/Markdown/LaTeX defects, without inventing missing scientific content. "
                        "The surrounding application will continue automatically if this reply ends before the full paper is finished, so keep translating as far as possible instead of stopping early because of length concerns. "
                        "Never mention token limits, context windows, truncation, or capacity constraints. "
                        "Never ask the user to split the paper into sections or resend parts. "
                        "Continue directly with the remaining translation and output only translated Markdown plus the final completion token when finished. "
                    "When the full text is finally completed, output the exact machine-detection completion token on a separate final line. "
                    "Do not replace it with TRANSLATION COMPLETE or any other words. "
                    f"{completion_token_instruction(marker)}"
                ),
            }
        )

    raise MinerUError(f"连续翻译达到 {max_rounds} 轮仍未检测到截止标记: {marker}")


def strip_translation_marker(text: str, marker: str) -> str:
    """
    删除翻译结果末尾的截止标记。

    说明：
    1. 优先删除完整 marker。
    2. 如果模型把 marker 中的数字打散，也尝试基于数字模式清理。
    3. 本函数只负责删除 marker 本体，不处理“结束标记：”“End marker:”等提示语残留，
       这类残留由 clean_translation_tail_artifacts 统一清理。
    """
    exact_start = text.rfind(marker)
    if exact_start >= 0:
        line_start = text.rfind("\n", 0, exact_start) + 1
        prefix_on_line = text[line_start:exact_start]
        start = (
            line_start
            if (not prefix_on_line.strip() or re.fullmatch(r"[\W_]+", prefix_on_line.strip(), flags=re.UNICODE))
            else exact_start
        )
        return text[:start].rstrip() + "\n"

    digits = "".join(ch for ch in marker if ch.isdigit())
    if digits:
        separators = r"[\s_\-.,:;|/\\]*"
        digit_pattern = separators.join(re.escape(ch) for ch in digits)
        match = list(re.finditer(r"(?<!\d)" + digit_pattern, text, flags=re.UNICODE))
        if match:
            start = match[-1].start()
            line_start = text.rfind("\n", 0, start) + 1
            prefix_on_line = text[line_start:start]
            if not prefix_on_line.strip() or re.fullmatch(r"[\W_]+", prefix_on_line.strip(), flags=re.UNICODE):
                start = line_start
            text = text[:start]

    return text.rstrip() + "\n"


def clean_translation_tail_artifacts(text: str, marker: str = "") -> str:
    """
    清理翻译结果末尾可能残留的“结束标记提示语”或模型泄漏痕迹。

    设计目标：
    1. 不影响正文中正常出现的“标记”“marker”等单词；
    2. 仅针对文末最后几行做清理；
    3. 兼容不同模型可能输出的中英文尾注痕迹。
    """
    normalized = (text or "").rstrip()

    # 先删除 marker 本体，避免后续尾部清理时受其干扰。
    if marker:
        normalized = strip_translation_marker(normalized, marker).rstrip()

    # 仅检查文末最后若干行，避免误删正文。
    lines = normalized.splitlines()
    if not lines:
        return "\n"

    tail_window = 6
    head_lines = lines[:-tail_window] if len(lines) > tail_window else []
    tail_lines = lines[-tail_window:] if len(lines) > tail_window else lines[:]

    # 常见尾部提示语模式：
    # - 结束标记：
    # - 截止标记：
    # - End marker:
    # - Completion token:
    # - marker:
    # - 标记：<<<12345678>>>
    tail_label_pattern = re.compile(
        r"""
        ^\s*
        (?:
            结束标记
            |截止标记
            |完成标记
            |结束符
            |标记
            |end\s*marker
            |completion\s*token
            |stop\s*marker
            |final\s*marker
            |marker
            |token
        )
        \s*
        [：:：\-—]?
        \s*
        (?:
            <<<.*?>>>
            |[\(\[]?\s*\d[\d\s_\-.,:;|/\\]*\s*[\)\]]?
        )?
        \s*$
        """,
        re.IGNORECASE | re.VERBOSE,
    )

    # 纯 marker 括号残片或空提示残片
    dangling_tail_pattern = re.compile(
        r"""
        ^\s*
        (?:
            <<<.*?
            |>>>
            |[（(]?\s*(?:end\s*marker|marker|token|结束标记|截止标记|标记)\s*[)）]?
        )
        \s*$
        """,
        re.IGNORECASE | re.VERBOSE,
    )

    # 从末尾向前清理连续的尾部痕迹行
    while tail_lines:
        last_line = tail_lines[-1].strip()
        if not last_line:
            tail_lines.pop()
            continue
        if tail_label_pattern.fullmatch(last_line) or dangling_tail_pattern.fullmatch(last_line):
            tail_lines.pop()
            continue
        break

    cleaned_lines = head_lines + tail_lines
    return "\n".join(cleaned_lines).rstrip() + "\n"


def query_quota(token: str) -> dict | None:
    try:
        result = http_json("GET", f"{API_V4_BASE_URL}/quota", token=token, timeout=30)
    except MinerUError:
        return None
    if result.get("code") == 0 and isinstance(result.get("data"), dict):
        return result["data"]
    return None


def http_put_file(upload_url: str, file_path: Path, timeout: int = 300, attempts: int = 4, log=None) -> None:
    parsed = urllib.parse.urlparse(upload_url)
    body = file_path.read_bytes()
    path = parsed.path or "/"
    if parsed.query:
        path = f"{path}?{parsed.query}"

    connection_cls = http.client.HTTPSConnection if parsed.scheme == "https" else http.client.HTTPConnection
    last_error: Exception | None = None

    for attempt in range(1, attempts + 1):
        connection = connection_cls(parsed.netloc, timeout=timeout)
        try:
            if log:
                log("正在上传文件…" if attempt == 1 else f"正在重新上传文件（第 {attempt}/{attempts} 次尝试）…")

            connection.request(
                "PUT",
                path,
                body=body,
                headers={
                    "User-Agent": USER_AGENT,
                    "Content-Length": str(len(body)),
                    "Connection": "close",
                },
            )
            response = connection.getresponse()
            response_body = response.read().decode("utf-8", errors="replace")

            if response.status in (200, 201, 204):
                return

            error = MinerUError(f"文件上传失败 HTTP {response.status}: {response_body}")
            if response.status in {408, 429} or 500 <= response.status < 600:
                last_error = error
                if attempt < attempts:
                    if log:
                        log(f"上传暂时失败，将重试: {error}")
                    time.sleep(min(2 ** attempt, 20))
                    continue

            raise error
        except (OSError, TimeoutError) as exc:
            last_error = exc
            if attempt < attempts:
                if log:
                    log(f"上传连接中断，将重试: {exc}")
                time.sleep(min(2 ** attempt, 20))
                continue
            raise MinerUError(f"文件上传失败: {exc}") from exc
        finally:
            connection.close()

    raise MinerUError(f"文件上传失败，多次重试仍未成功: {last_error}")


def http_bytes(url: str, timeout: int = 300) -> bytes:
    request = urllib.request.Request(url, headers={"User-Agent": USER_AGENT}, method="GET")
    try:
        with urllib.request.urlopen(request, timeout=timeout) as response:
            return response.read()
    except urllib.error.HTTPError as exc:
        detail = exc.read().decode("utf-8", errors="replace")
        raise MinerUError(f"下载结果失败 HTTP {exc.code}: {detail}") from exc
    except urllib.error.URLError as exc:
        raise MinerUError(f"下载结果失败: {exc.reason}") from exc


def http_bytes_with_retries(url: str, log, attempts: int = 5, timeout: int = 180) -> bytes:
    last_error: Exception | None = None
    for attempt in range(1, attempts + 1):
        try:
            log("正在下载解析结果…" if attempt == 1 else f"正在重新下载解析结果（第 {attempt}/{attempts} 次尝试）…")
            return http_bytes(url, timeout=timeout)
        except MinerUError as exc:
            last_error = exc
            log(f"下载暂时失败: {exc}")
            if attempt < attempts:
                time.sleep(min(5 * attempt, 30))
    raise MinerUError(f"多次下载结果压缩包失败，最后错误: {last_error}")


def build_mineru_data_id(file_path: Path, max_length: int = 128) -> str:
    """
    构造发送给 MinerU 的 data_id，并严格限制总长度。

    设计说明：
    1. MinerU 接口要求 files.data_id 长度不能超过 128。
    2. 原始文件名可能很长，不能直接使用 file_path.stem。
    3. 这里保留一部分可读文件名，再拼接短随机后缀，兼顾可读性与唯一性。
    """
    # 只保留相对稳定的字符，避免极端文件名带来兼容性问题。
    raw_stem = (file_path.stem or "").strip()
    safe_stem = re.sub(r"\s+", "-", raw_stem)
    safe_stem = re.sub(r"[^A-Za-z0-9._\-\u4e00-\u9fff]+", "-", safe_stem, flags=re.UNICODE)
    safe_stem = re.sub(r"-{2,}", "-", safe_stem).strip("-._")

    if not safe_stem:
        safe_stem = "document"

    # MinerU 对 data_id 总长度有限制，因此后缀长度也必须计入总长预算。
    unique_suffix = uuid.uuid4().hex[:8]
    suffix = f"-{unique_suffix}"

    # 至少给主体保留 1 个字符，避免出现空前缀。
    max_prefix_length = max(1, max_length - len(suffix))
    if len(safe_stem) > max_prefix_length:
        safe_stem = safe_stem[:max_prefix_length].rstrip("-._")

    if not safe_stem:
        safe_stem = "d"

    data_id = f"{safe_stem}{suffix}"

    # 二次兜底，确保绝不超长。
    if len(data_id) > max_length:
        fallback_prefix_length = max(1, max_length - len(suffix))
        data_id = f"{safe_stem[:fallback_prefix_length]}{suffix}"

    return data_id


def submit_precise_file(file_path: Path, options: ParseOptions, token: str) -> tuple[str, str, str]:
    model_version = "MinerU-HTML" if file_path.suffix.lower() in {".html", ".htm"} else options.model_version
    data_id = build_mineru_data_id(file_path)
    payload = {
        "enable_formula": options.enable_formula,
        "enable_table": options.enable_table,
        "model_version": model_version,
        "files": [
            {
                "name": file_path.name,
                "is_ocr": options.is_ocr,
                "data_id": data_id,
            }
        ],
    }
    result = http_json("POST", f"{API_V4_BASE_URL}/file-urls/batch", payload, token=token)
    if result.get("code") != 0:
        raise MinerUError(f"创建精准解析任务失败: {result.get('msg', result)}")

    data = result.get("data") or {}
    batch_id = data.get("batch_id")
    file_urls = data.get("file_urls") or []
    if isinstance(file_urls, str):
        upload_url = file_urls
    elif isinstance(file_urls, list) and file_urls:
        first = file_urls[0]
        upload_url = first.get("url") if isinstance(first, dict) else str(first)
    else:
        upload_url = ""
    if not batch_id or not upload_url:
        raise MinerUError(f"创建精准解析任务响应缺少 batch_id/file_urls: {result}")
    return batch_id, upload_url, data_id


def poll_precise_result(batch_id: str, options: ParseOptions, token: str, log) -> dict:
    started = time.time()
    transient_errors = 0
    while time.time() - started < options.timeout_seconds:
        try:
            result = http_json("GET", f"{API_V4_BASE_URL}/extract-results/batch/{batch_id}", token=token, timeout=120)
        except MinerUError as exc:
            transient_errors += 1
            elapsed = int(time.time() - started)
            log(f"[{elapsed}s] 查询结果暂时失败，第 {transient_errors} 次重试: {exc}")
            time.sleep(min(options.poll_interval_seconds * transient_errors, 30))
            continue

        transient_errors = 0
        if result.get("code") != 0:
            raise MinerUError(f"查询精准解析任务失败: {result.get('msg', result)}")

        data = result.get("data") or {}
        items = data.get("extract_result") or data.get("extract_results") or data.get("files") or []
        if isinstance(items, dict):
            items = [items]

        elapsed = int(time.time() - started)
        if not items:
            log(f"[{elapsed}s] 等待云端解析返回结果…")
            time.sleep(options.poll_interval_seconds)
            continue

        item = items[0]
        state = str(item.get("state") or item.get("status") or "").lower()
        zip_url = item.get("full_zip_url") or item.get("zip_url") or item.get("result_url")
        progress = item.get("extract_progress")
        if zip_url and (not state or state in {"done", "finished", "success"}):
            log(f"[{elapsed}s] 云端精准解析完成，正在下载解析结果…")
            return item
        if state in {"failed", "fail", "error"}:
            raise MinerUError(f"精准解析失败: {item.get('err_msg') or item.get('msg') or item}")

        if isinstance(progress, dict):
            log(f"[{elapsed}s] 解析中: {progress.get('extracted_pages', '?')}/{progress.get('total_pages', '?')} 页")
        else:
            log(f"[{elapsed}s] 云端解析状态: {state or '处理中'}")
        time.sleep(options.poll_interval_seconds)

    raise MinerUError(f"精准解析轮询超时，batch_id={batch_id}")


def extract_markdown_from_zip(result_item: dict, output_dir: Path, log) -> tuple[str, str, Path]:
    zip_url = result_item.get("full_zip_url") or result_item.get("zip_url") or result_item.get("result_url")
    if not zip_url:
        raise MinerUError(f"结果中没有 full_zip_url: {result_item}")

    zip_path = output_dir / "mineru_result.zip"
    extract_dir = output_dir / "mineru_result"
    if extract_dir.exists():
        shutil.rmtree(extract_dir)
    extract_dir.mkdir(parents=True, exist_ok=True)

    zip_path.write_bytes(http_bytes_with_retries(zip_url, log))
    with zipfile.ZipFile(zip_path) as archive:
        archive.extractall(extract_dir)

    # 压缩包只用于解压，成功解压后删除，避免工作目录长期残留大体积 zip。
    try:
        if zip_path.exists():
            zip_path.unlink()
    except Exception:
        pass

    md_candidates = sorted(extract_dir.rglob("full.md")) or sorted(extract_dir.rglob("*.md"))
    if not md_candidates:
        raise MinerUError(f"结果压缩包中没有找到 Markdown 文件: {zip_path}")
    return md_candidates[0].read_text(encoding="utf-8", errors="replace"), zip_url, extract_dir


def extension_from_data_uri(header: str) -> str:
    mime = header.split(";", 1)[0].replace("data:", "").strip().lower()
    return mimetypes.guess_extension(mime) or ".bin"


def extension_from_target(target: str) -> str:
    parsed = urllib.parse.urlparse(target)
    suffix = Path(urllib.parse.unquote(parsed.path)).suffix
    return suffix if suffix and len(suffix) <= 8 else ".png"


def simplify_markdown_images(
    markdown: str,
    output_dir: Path,
    source_dirs: list[Path] | None = None,
    image_index_start: int = 0,
) -> tuple[str, list[dict]]:
    images_dir = output_dir / "images"
    images_dir.mkdir(parents=True, exist_ok=True)
    source_dirs = source_dirs or []
    records: list[dict] = []
    counter = max(0, int(image_index_start))

    def next_image_id() -> str:
        nonlocal counter
        counter += 1
        return f"IMAGE_{counter:03d}"

    def save_data_uri(image_id: str, target: str) -> tuple[str, str]:
        header, b64_data = target.split(",", 1)
        ext = extension_from_data_uri(header)
        save_path = images_dir / f"{image_id.lower()}{ext}"
        save_path.write_bytes(base64.b64decode(b64_data, validate=False))
        return f"images/{save_path.name}", str(save_path)

    def find_source_image(target: str) -> Path | None:
        unwrapped = target.strip()
        if unwrapped.startswith("<") and unwrapped.endswith(">"):
            unwrapped = unwrapped[1:-1]
        parsed = urllib.parse.urlparse(unwrapped)
        if parsed.scheme and parsed.scheme != "file":
            return None
        raw_path = urllib.parse.unquote(parsed.path if parsed.scheme == "file" else unwrapped)
        candidates = [Path(raw_path)]
        for source_dir in source_dirs:
            candidates.append(source_dir / raw_path)
            candidates.append(source_dir / Path(raw_path).name)
            candidates.append(source_dir / "images" / Path(raw_path).name)
        found = next((candidate for candidate in candidates if candidate.is_file()), None)
        if found:
            return found
        basename = Path(raw_path).name
        if basename:
            for source_dir in source_dirs:
                found = next((candidate for candidate in source_dir.rglob(basename) if candidate.is_file()), None)
                if found:
                    return found
        return None

    def copy_source_image(image_id: str, source: Path) -> tuple[str, str]:
        suffix = source.suffix or ".jpg"
        save_path = images_dir / f"{image_id.lower()}{suffix}"
        shutil.copy2(source, save_path)
        return f"images/{save_path.name}", str(save_path)

    def replace_markdown_image(match: re.Match) -> str:
        alt = match.group("alt")
        target = match.group("target").strip()
        title = (match.group("title") or "").strip()
        image_id = next_image_id()
        saved_file = None
        try:
            if target.startswith("data:image/") and "," in target:
                clean_target, saved_file = save_data_uri(image_id, target)
            else:
                source_image = find_source_image(target)
                if source_image:
                    clean_target, saved_file = copy_source_image(image_id, source_image)
                else:
                    clean_target = f"images/{image_id.lower()}{extension_from_target(target)}"
        except Exception as exc:
            clean_target = f"images/{image_id.lower()}{extension_from_target(target)}"
            records.append({"id": image_id, "warning": f"图片保存失败: {exc}"})

        records.append(
            {
                "id": image_id,
                "alt": alt,
                "original_target": target,
                "title": title,
                "clean_target": clean_target,
                "saved_file": saved_file,
            }
        )
        return f"![{image_id}]({clean_target})"

    md_image_pattern = re.compile(
        r"!\[(?P<alt>[^\]]*)\]\(\s*(?P<target><[^>]+>|[^\s)]+)(?P<title>\s+['\"][^'\"]*['\"])?\s*\)",
        re.DOTALL,
    )
    cleaned = md_image_pattern.sub(replace_markdown_image, markdown)

    def replace_html_img(match: re.Match) -> str:
        attrs = match.group("attrs")
        src = re.search(r"""src\s*=\s*["'](?P<src>.*?)["']""", attrs, re.IGNORECASE | re.DOTALL)
        alt = re.search(r"""alt\s*=\s*["'](?P<alt>.*?)["']""", attrs, re.IGNORECASE | re.DOTALL)
        if not src:
            return match.group(0)
        fake = f"![{alt.group('alt') if alt else ''}]({src.group('src')})"
        return md_image_pattern.sub(replace_markdown_image, fake)

    return re.sub(r"<img\b(?P<attrs>[^>]*)>", replace_html_img, cleaned, flags=re.IGNORECASE | re.DOTALL), records


def first_recursive_file(folder: Path, patterns: tuple[str, ...]) -> Path | None:
    for pattern in patterns:
        candidate = next(folder.rglob(pattern), None)
        if candidate is not None:
            return candidate
    return None


def rebase_page_indices(value, page_offset: int):
    """Copy JSON while shifting explicit zero-based page index fields."""
    if isinstance(value, list):
        return [rebase_page_indices(item, page_offset) for item in value]
    if not isinstance(value, dict):
        return value
    rebased = {}
    for key, item in value.items():
        if key in {"page_idx", "page_index"} and isinstance(item, int) and not isinstance(item, bool):
            rebased[key] = item + page_offset
        else:
            rebased[key] = rebase_page_indices(item, page_offset)
    return rebased


def rewrite_mineru_asset_references(value, basename_map: dict[str, str]):
    """Point MinerU JSON image references at the canonical merged image dir."""
    if isinstance(value, list):
        return [rewrite_mineru_asset_references(item, basename_map) for item in value]
    if isinstance(value, dict):
        return {
            key: rewrite_mineru_asset_references(item, basename_map)
            for key, item in value.items()
        }
    if not isinstance(value, str) or not basename_map:
        return value
    normalized = urllib.parse.unquote(value.replace("\\", "/"))
    basename = Path(normalized).name
    replacement = basename_map.get(basename)
    if replacement and (
        normalized == basename
        or "/images/" in f"/{normalized.lstrip('/')}"
        or Path(normalized).suffix.lower() in IMAGE_SUFFIXES
    ):
        return replacement
    # Some structured outputs embed image paths inside short HTML fragments.
    rewritten = value
    for original, target in basename_map.items():
        rewritten = re.sub(
            rf"(?<![A-Za-z0-9_.-])(?:images[/\\])?{re.escape(original)}(?![A-Za-z0-9_.-])",
            f"images/{target}",
            rewritten,
        )
    return rewritten


def _load_json_if_present(path: Path | None):
    if path is None:
        return None
    try:
        return json.loads(path.read_text(encoding="utf-8", errors="replace"))
    except (OSError, json.JSONDecodeError):
        return None


def write_json_streaming(path: Path, payload) -> None:
    """Write large generated JSON without first allocating one giant string."""
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary = path.with_name(f".{path.name}.{uuid.uuid4().hex}.tmp")
    try:
        with temporary.open("w", encoding="utf-8", newline="\n") as handle:
            json.dump(payload, handle, ensure_ascii=False, indent=2)
            handle.write("\n")
            handle.flush()
            os.fsync(handle.fileno())
        os.replace(temporary, path)
    finally:
        try:
            temporary.unlink(missing_ok=True)
        except OSError:
            pass


def merge_mineru_part_results(part_results: list[dict], output_dir: Path, log) -> tuple[str, str, list[dict], Path]:
    """Merge independently parsed MinerU files into one canonical document."""
    if not part_results:
        raise MinerUError("没有可合并的 MinerU 分片结果。")

    extract_root = output_dir / "mineru_result"
    if extract_root.exists():
        shutil.rmtree(extract_root)
    parts_root = extract_root / "parts"
    parts_root.mkdir(parents=True, exist_ok=True)
    layout_images_dir = extract_root / "images"
    layout_images_dir.mkdir(parents=True, exist_ok=True)

    raw_chunks: list[str] = []
    clean_chunks: list[str] = []
    all_image_records: list[dict] = []
    merged_pages: list = []
    merged_model_pages: list = []
    merged_content_v2_pages: list = []
    merged_content_legacy: list = []
    layout_template: dict | None = None
    image_index = 0
    page_offset = 0

    for part_index, part in enumerate(part_results, start=1):
        part_name = f"part_{part_index:03d}"
        copied_extract_dir = parts_root / part_name
        shutil.copytree(Path(part["extract_dir"]), copied_extract_dir)
        markdown = str(part["markdown"])
        raw_chunks.append(markdown.strip())
        cleaned, image_records = simplify_markdown_images(
            markdown,
            output_dir,
            [copied_extract_dir],
            image_index_start=image_index,
        )
        ids = {
            str(record.get("id") or "")
            for record in image_records
            if isinstance(record, dict) and record.get("id")
        }
        image_index += len(ids)
        clean_chunks.append(cleaned.strip())
        for record in image_records:
            if isinstance(record, dict):
                record["source_part"] = part_index
                record["source_page_range"] = [part["start_page"], part["end_page"]]
        all_image_records.extend(image_records)

        basename_map = {
            Path(str(record.get("original_target") or "").replace("\\", "/")).name:
                Path(str(record.get("clean_target") or "")).name
            for record in image_records
            if isinstance(record, dict)
            and record.get("original_target")
            and record.get("clean_target")
        }
        for record in image_records:
            if not isinstance(record, dict):
                continue
            original_name = Path(str(record.get("original_target") or "").replace("\\", "/")).name
            target_name = basename_map.get(original_name)
            saved_file = Path(str(record.get("saved_file") or ""))
            if target_name and saved_file.is_file():
                shutil.copy2(saved_file, layout_images_dir / target_name)

        # layout.json may reference equation/table crops that full.md does not
        # expose as Markdown images. Preserve and namespace every remaining
        # MinerU asset so independent tasks can never overwrite one another.
        for source_asset in copied_extract_dir.rglob("*"):
            if not source_asset.is_file() or source_asset.suffix.lower() not in IMAGE_SUFFIXES:
                continue
            original_name = source_asset.name
            if original_name in basename_map:
                continue
            target_name = f"p{part_index:03d}_{original_name}"
            duplicate_index = 2
            while (layout_images_dir / target_name).exists():
                target_name = f"p{part_index:03d}_{source_asset.stem}_{duplicate_index}{source_asset.suffix}"
                duplicate_index += 1
            shutil.copy2(source_asset, layout_images_dir / target_name)
            basename_map[original_name] = target_name

        layout_path = first_recursive_file(copied_extract_dir, ("layout.json", "*_middle.json"))
        layout_payload = _load_json_if_present(layout_path)
        if isinstance(layout_payload, dict):
            if layout_template is None:
                # Only pdf_info is replaced below. A shallow top-level copy is
                # sufficient and avoids duplicating the first part's page tree.
                layout_template = dict(layout_payload)
            pages = layout_payload.get("pdf_info")
            if isinstance(pages, list):
                pages = rebase_page_indices(pages, page_offset)
                pages = rewrite_mineru_asset_references(pages, basename_map)
                for local_index, page in enumerate(pages):
                    if isinstance(page, dict):
                        page["page_idx"] = page_offset + local_index
                merged_pages.extend(pages)

        model_path = first_recursive_file(copied_extract_dir, ("*_model.json", "model.json"))
        model_payload = _load_json_if_present(model_path)
        if isinstance(model_payload, list):
            merged_model_pages.extend(
                rewrite_mineru_asset_references(
                    rebase_page_indices(model_payload, page_offset),
                    basename_map,
                )
            )

        content_v2_path = first_recursive_file(
            copied_extract_dir,
            ("*_content_list_v2.json", "content_list_v2.json"),
        )
        content_v2_payload = _load_json_if_present(content_v2_path)
        if isinstance(content_v2_payload, list):
            merged_content_v2_pages.extend(
                rewrite_mineru_asset_references(
                    rebase_page_indices(content_v2_payload, page_offset),
                    basename_map,
                )
            )
        legacy_candidates = [
            path
            for pattern in ("*_content_list.json", "content_list.json")
            for path in copied_extract_dir.rglob(pattern)
            if "content_list_v2" not in path.name
        ]
        content_legacy_payload = _load_json_if_present(legacy_candidates[0] if legacy_candidates else None)
        if isinstance(content_legacy_payload, list):
            merged_content_legacy.extend(
                rewrite_mineru_asset_references(
                    rebase_page_indices(content_legacy_payload, page_offset),
                    basename_map,
                )
            )

        page_offset += int(part["end_page"]) - int(part["start_page"]) + 1

    if layout_template is not None and merged_pages:
        layout_template["pdf_info"] = merged_pages
        write_json_streaming(extract_root / "layout.json", layout_template)
    if merged_model_pages:
        write_json_streaming(extract_root / "merged_model.json", merged_model_pages)
    if merged_content_v2_pages:
        write_json_streaming(extract_root / "merged_content_list_v2.json", merged_content_v2_pages)
    if merged_content_legacy:
        write_json_streaming(extract_root / "merged_content_list.json", merged_content_legacy)

    raw_markdown = "\n\n".join(chunk for chunk in raw_chunks if chunk).rstrip() + "\n"
    clean_markdown = "\n\n".join(chunk for chunk in clean_chunks if chunk).rstrip() + "\n"
    if len(part_results) > 1:
        log(
            f"已按原始页序合并 {len(part_results)} 个解析分片；"
            f"全局版面页索引已重建为 0-{max(0, page_offset - 1)}。"
        )
    return raw_markdown, clean_markdown, all_image_records, extract_root


def mineru_result_organization_message(part_count: int) -> str:
    """Return a progress message that reflects whether the upload was split."""
    if part_count > 1:
        return "正在合并 MinerU 分片结果并重建全局页码与资源编号..."
    return "正在整理 MinerU 单文件解析结果与资源文件..."


def find_pandoc() -> Path | None:
    return preview_tools.find_pandoc(WORKSPACE)


_WORD_REFINE_CHECK_LOCK = threading.Lock()
_WORD_REFINE_CHECK_CACHE: tuple[bool, str] | None = None


def check_word_refine_available() -> tuple[bool, str]:
    global _WORD_REFINE_CHECK_CACHE
    with _WORD_REFINE_CHECK_LOCK:
        if _WORD_REFINE_CHECK_CACHE is not None:
            return _WORD_REFINE_CHECK_CACHE

    if os.name != "nt":
        result = (False, "仅 Windows 下支持 Word 精校。")
    else:
        script = r"""
$ErrorActionPreference = 'Stop'
$word = $null
try {
    $word = New-Object -ComObject Word.Application
    $word.Visible = $false
    $word.DisplayAlerts = 0
    Write-Output 'OK'
} finally {
    if ($word -ne $null) { $word.Quit() }
}
"""
        try:
            probe = subprocess.run(
                ["powershell", "-NoProfile", "-ExecutionPolicy", "Bypass", "-Command", script],
                capture_output=True,
                text=True,
                timeout=15,
                **hidden_subprocess_kwargs(),
            )
            if probe.returncode == 0:
                result = (True, "需要本机安装 Word。会额外尝试把表格和正文里的公式升级成 Word 原生公式对象，速度更慢。")
            else:
                detail = (probe.stderr or probe.stdout or "").strip()
                result = (False, f"当前环境不可用：{detail or '未检测到可用的 Word COM。'}")
        except Exception as exc:
            result = (False, f"当前环境不可用：{exc}")

    with _WORD_REFINE_CHECK_LOCK:
        _WORD_REFINE_CHECK_CACHE = result
    return result


def render_preview_html(markdown_path: Path, log=None, style: ExportStyleSettings | None = None) -> Path | None:
    return preview_tools.render_preview_html(markdown_path, WORKSPACE, log, style=style)


def render_export_html(markdown_path: Path, log=None, style: ExportStyleSettings | None = None) -> Path | None:
    return preview_tools.render_export_html(markdown_path, WORKSPACE, log, style=style)


def polish_preview_html(html_path: Path, style: ExportStyleSettings | None = None) -> None:
    preview_tools.polish_html(html_path, style)


def configure_web_view(view) -> None:
    if not view:
        return
    settings = view.settings()
    settings.setAttribute(QWebEngineSettings.WebAttribute.LocalContentCanAccessFileUrls, True)
    settings.setAttribute(QWebEngineSettings.WebAttribute.LocalContentCanAccessRemoteUrls, True)
    for attr_name in ("PdfViewerEnabled", "PluginsEnabled"):
        attr = getattr(QWebEngineSettings.WebAttribute, attr_name, None)
        if attr is not None:
            try:
                settings.setAttribute(attr, True)
            except Exception:
                pass


def qurl_from_payload(payload: str) -> QUrl:
    url = QUrl(str(payload))
    if url.isLocalFile():
        return QUrl.fromLocalFile(url.toLocalFile())
    return url


def apply_layout_debug_mode_to_web_view(web_view, enabled: bool) -> None:
    if not web_view:
        return
    script = f"""
    (() => {{
      const debug = {str(bool(enabled)).lower()};
      const apply = () => {{
        if (!document.body) return;
      document.body.classList.toggle('layout-debug', debug);
      document.body.classList.toggle('layout-production', !debug);
      let style = document.getElementById('mineru-layout-debug-runtime-style');
      if (!style) {{
        style = document.createElement('style');
        style.id = 'mineru-layout-debug-runtime-style';
        style.textContent = `
          /* 正文迭代检查逻辑的样式和标签由预览页自身维护。 */
        `;
        document.head.appendChild(style);
      }}
      if (window.__mineruRunLayoutFill) window.__mineruRunLayoutFill();
      for (const node of document.querySelectorAll('.layout-flow-stream, .layout-block')) {{
        node.classList.remove('caption-fit-debug');
      }}
      if (!debug) {{
        for (const node of document.querySelectorAll('.layout-flow-stream, .layout-block')) {{
          node.classList.remove('fit-limiter');
          node.classList.remove('fit-blocker');
          node.classList.remove('body-iteration-collision');
          node.dataset.fitLabel = '';
          node.dataset.fitDebug = '';
          node.title = '';
        }}
      }}
      if (window.__mineruFitLayoutPages) window.__mineruFitLayoutPages();
      }};
      if (!document.body) {{
        document.addEventListener('DOMContentLoaded', apply, {{ once: true }});
        return;
      }}
      apply();
    }})();
    """
    try:
        web_view.page().runJavaScript(script)
    except Exception:
        pass


class ElidedLabel(QLabel):
    """自动省略过长文本的标签，避免长论文标题撑大界面。"""

    def __init__(self, text: str = "", parent=None, elide_mode=Qt.TextElideMode.ElideRight):
        super().__init__("", parent)
        self._full_text = str(text or "")
        self._elide_mode = elide_mode

        # Keep the full label text from expanding the parent layout.
        self.setMinimumWidth(0)
        self.setWordWrap(False)
        self.setSizePolicy(QSizePolicy.Policy.Ignored, QSizePolicy.Policy.Preferred)
        self.setToolTip(self._full_text)
        self._refresh_elided_text()

    def setText(self, text: str):
        """保存完整文本，但界面只显示按当前宽度省略后的文本。"""
        self._full_text = str(text or "")
        self.setToolTip(self._full_text)
        self._refresh_elided_text()

    def fullText(self) -> str:
        """返回未省略的完整文本。"""
        return self._full_text

    def resizeEvent(self, event):
        super().resizeEvent(event)
        self._refresh_elided_text()

    def _refresh_elided_text(self):
        """按标签当前可用宽度生成省略文本。"""
        available_width = self.contentsRect().width()

        # Use a conservative width until layout is ready to avoid visible reflow.
        if available_width <= 0:
            available_width = 160

        display_text = self.fontMetrics().elidedText(
            self._full_text,
            self._elide_mode,
            available_width,
        )
        if QLabel.text(self) != display_text:
            QLabel.setText(self, display_text)


def apply_reader_font_to_web_view(web_view, font_pt: int):
    if not web_view:
        return
    safe_pt = max(9, int(font_pt or 12))
    line_height_pt = journal_reading_line_height_pt(safe_pt)
    script = (
        "document.documentElement.style.setProperty('--reader-font-size', "
        + json.dumps(f"{safe_pt}pt", ensure_ascii=False)
        + ");"
        + "document.documentElement.style.setProperty('--reader-line-height', "
        + json.dumps(f"{line_height_pt}pt", ensure_ascii=False)
        + ");"
    )
    try:
        web_view.page().runJavaScript(script)
    except Exception:
        pass


def ensure_web_view_mathjax_typeset(web_view) -> None:
    """Retry MathJax startup for a locally loaded Pandoc preview.

    QWebEngine may report the local HTML load as finished while its deferred
    MathJax CDN script is still starting.  A fresh reader window usually gives
    that script enough time, whereas a reused main-pane WebEngine can retain
    the raw TeX DOM.  After MathJax changes an equation from its TeX placeholder
    to SVG, the layout reader also needs one geometry pass; otherwise the first
    document shown after startup can keep a stale equation-number gutter until
    the user switches away and back.  The bounded retry makes both views use
    the same result without putting expensive equation fitting on every resize.
    """
    if not web_view:
        return
    script = r"""
    (() => {
      const typesetWhenReady = (remaining) => {
        const mathjax = window.MathJax;
        if (mathjax && mathjax.startup && mathjax.startup.promise) {
          const refitLayoutEquations = () => {
            if (window.__mineruFitLayoutEquations) window.__mineruFitLayoutEquations();
          };
          mathjax.startup.promise
            .then(() => mathjax.typesetPromise ? mathjax.typesetPromise() : null)
            .then(() => {
              // MathJax resolves before Chromium necessarily paints its SVG.
              // Two frames provide stable font metrics; the short follow-up
              // covers the first visible QWebEngine viewport resize.
              requestAnimationFrame(() => requestAnimationFrame(refitLayoutEquations));
              window.setTimeout(refitLayoutEquations, 180);
            })
            .catch(() => {});
          return;
        }
        if (remaining > 0) window.setTimeout(() => typesetWhenReady(remaining - 1), 200);
      };
      typesetWhenReady(25);
    })();
    """
    try:
        web_view.page().runJavaScript(script)
    except Exception:
        pass


def apply_reader_font_to_text_browser(viewer: QTextBrowser, font_pt: int, style: ExportStyleSettings | None = None):
    if not viewer:
        return
    safe_pt = max(9, int(font_pt or 12))
    style = resolve_export_style(style)
    reader_font_family = bundled_reader_qt_font_family()
    line_spacing_pt = journal_reading_line_height_pt(
        safe_pt,
        base_line_spacing_pt=clamp_int(style.line_spacing_pt, 15, 30, 20),
    )
    viewer.setStyleSheet(
        (
            'QTextBrowser#readerPane {{'
            'font-family: "{reader_font_family}", "{body_font_latin}", "{body_font_cjk}", serif;'
            'font-size: {font_pt}pt;'
            'line-height: {line_spacing_pt}pt;'
            '}}'
        ).format(
            reader_font_family=reader_font_family,
            body_font_latin=style.body_font_latin,
            body_font_cjk=style.body_font_cjk,
            font_pt=safe_pt,
            line_spacing_pt=line_spacing_pt,
        )
    )


def journal_reading_line_height_pt(font_pt: int, base_line_spacing_pt: int = 20) -> int:
    """
    Approximate compact journal-style leading while staying readable on screen.

    Real journal body text is usually tighter than general web/article reading,
    but large on-screen font sizes still need extra room to avoid collisions.
    We keep the mapping simple and deterministic so both preview backends stay
    visually aligned and easy to reason about.
    """
    safe_pt = max(9, int(font_pt or 12))
    base_line_spacing_pt = max(12, int(base_line_spacing_pt or 20))
    ratio = 1.42 if safe_pt <= 14 else 1.45
    return max(base_line_spacing_pt, int(round(safe_pt * ratio)))


def make_export_markdown(markdown_path: Path, style: ExportStyleSettings | None = None) -> Path:
    markdown_path = markdown_path.resolve()
    raw = markdown_path.read_text(encoding="utf-8", errors="replace")
    normalized = preview_tools.normalize_markdown_for_export(
        raw,
        export_style_markdown_image_width(style),
        preview_tools.layout_image_width_percentages(markdown_path),
    )
    out_path = markdown_path.with_name(f".{markdown_path.stem}.export.md")
    out_path.write_text(normalized, encoding="utf-8")
    return out_path


def stored_original_path(output_dir: Path, source_path: Path) -> Path:
    return output_dir / source_path.name


def find_stored_original(folder: Path, meta: dict | None = None) -> Path | None:
    # Document-chat bubbles use the parsed Markdown path and do not carry the
    # metadata loaded by MainWindow.
    # 因此在这里自行读取任务元数据，优先锁定上传时保存的原始副本，不能依赖
    # 文件夹遍历顺序（其中可能混有 preview / layout 等生成的 HTML）。
    if meta is None:
        meta_path = folder / "mineru_task.json"
        try:
            loaded = json.loads(meta_path.read_text(encoding="utf-8", errors="replace"))
            meta = loaded if isinstance(loaded, dict) else {}
        except (OSError, json.JSONDecodeError):
            meta = {}

    meta = meta or {}
    stored_path = Path(str(meta.get("source_file") or "")).expanduser()
    if stored_path.is_file():
        return stored_path

    source_name = Path(str(meta.get("source_file") or meta.get("source_pdf") or "")).name
    candidates: list[Path] = []
    if source_name:
        candidates.append(folder / source_name)
    candidates.extend(
        path for path in folder.iterdir()
        if (
            path.is_file()
            and path.suffix.lower() in SUPPORTED_INPUT_EXTENSIONS
            and not path.name.lower().startswith(("preview.", "layout.", "debug."))
        )
    )
    candidates.extend(folder.glob("mineru_result/*_origin.pdf"))
    return next((path for path in candidates if path.exists()), None)


def basic_export_warnings(path: Path) -> list[str]:
    warnings: list[str] = []
    if not path.exists():
        return ["导出文件不存在。"]
    size = path.stat().st_size
    if size < 5_000:
        warnings.append("导出文件体积异常偏小，请打开确认内容是否完整。")
    suffix = path.suffix.lower()
    if suffix == ".docx":
        try:
            with zipfile.ZipFile(path) as archive:
                names = set(archive.namelist())
                if "word/document.xml" not in names:
                    warnings.append("Word 文档结构不完整，缺少 word/document.xml。")
                if not any(name.startswith("word/") for name in names):
                    warnings.append("Word 文档缺少 word 目录内容。")
        except Exception as exc:
            warnings.append(f"Word 文档无法作为 docx 打开: {exc}")
    elif suffix == ".pdf":
        try:
            with path.open("rb") as file:
                header = file.read(5)
            if header != b"%PDF-":
                warnings.append("PDF 文件头异常。")
        except Exception as exc:
            warnings.append(f"PDF 文件读取失败: {exc}")
    return warnings


OOXML_NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
}
for prefix, uri in OOXML_NS.items():
    ET.register_namespace(prefix, uri)


def w_tag(name: str) -> str:
    return f"{{{OOXML_NS['w']}}}{name}"


def wp_tag(name: str) -> str:
    return f"{{{OOXML_NS['wp']}}}{name}"


def a_tag(name: str) -> str:
    return f"{{{OOXML_NS['a']}}}{name}"


def ensure_child(parent: ET.Element, tag: str, before_tags: set[str] | None = None) -> ET.Element:
    child = parent.find(tag)
    if child is not None:
        return child
    child = ET.Element(tag)
    if before_tags:
        for index, existing in enumerate(list(parent)):
            if existing.tag in before_tags:
                parent.insert(index, child)
                return child
    parent.append(child)
    return child


def set_w_val(element: ET.Element, value: str):
    element.set(w_tag("val"), value)


def heading_level_for_paragraph(paragraph: ET.Element) -> int:
    style = paragraph.find("./w:pPr/w:pStyle", OOXML_NS)
    value = style.get(w_tag("val"), "") if style is not None else ""
    match = re.search(r"(\d+)$", value)
    if match:
        try:
            return max(1, min(6, int(match.group(1))))
        except Exception:
            return 1
    if value.startswith("标题"):
        return 1
    return 0


def heading_size_half_points(style: ExportStyleSettings | None = None, level: int = 1) -> str:
    style = resolve_export_style(style)
    if level <= 1:
        return style_half_points(style.heading1_pt)
    if level == 2:
        return style_half_points(style.heading2_pt)
    return style_half_points(style.heading3_pt)


def set_run_style(
    run: ET.Element,
    size_half_points: str | None = None,
    bold: bool = False,
    *,
    style: ExportStyleSettings | None = None,
    heading: bool = False,
):
    style = resolve_export_style(style)
    rpr = ensure_child(run, w_tag("rPr"))
    fonts = ensure_child(rpr, w_tag("rFonts"))
    fonts.set(w_tag("ascii"), style.heading_font_latin if heading else style.body_font_latin)
    fonts.set(w_tag("hAnsi"), style.heading_font_latin if heading else style.body_font_latin)
    fonts.set(w_tag("eastAsia"), style.heading_font_cjk if heading else style.body_font_cjk)
    color = ensure_child(rpr, w_tag("color"))
    color.set(w_tag("val"), "000000")
    if size_half_points is None:
        size_half_points = heading_size_half_points(style, 1) if heading else style_half_points(style.body_font_pt)
    size = ensure_child(rpr, w_tag("sz"))
    size.set(w_tag("val"), size_half_points)
    size_cs = ensure_child(rpr, w_tag("szCs"))
    size_cs.set(w_tag("val"), size_half_points)
    if bold:
        ensure_child(rpr, w_tag("b"))
        ensure_child(rpr, w_tag("bCs"))


def append_word_text_segments(
    paragraph: ET.Element,
    text: str,
    *,
    bold: bool = False,
    size_half_points: str | None = None,
    style: ExportStyleSettings | None = None,
    heading: bool = False,
):
    segments = parse_texish_segments(text)
    if not segments:
        segments = [InlineTextSegment(text)]

    for segment in segments:
        parts = segment.text.split("\n")
        for part_index, part in enumerate(parts):
            if part_index > 0:
                break_run = ET.SubElement(paragraph, w_tag("r"))
                set_run_style(break_run, size_half_points, bold=bold, style=style, heading=heading)
                ET.SubElement(break_run, w_tag("br"))
            if not part:
                continue
            run = ET.SubElement(paragraph, w_tag("r"))
            set_run_style(run, size_half_points, bold=bold, style=style, heading=heading)
            if segment.vertical_align:
                rpr = ensure_child(run, w_tag("rPr"))
                vert_align = ensure_child(rpr, w_tag("vertAlign"))
                vert_align.set(w_tag("val"), segment.vertical_align)
            text_node = ET.SubElement(run, w_tag("t"))
            if part.startswith(" ") or part.endswith(" "):
                text_node.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            text_node.text = part


def paragraph_text(paragraph: ET.Element) -> str:
    return "".join(node.text or "" for node in paragraph.findall(".//w:t", OOXML_NS)).strip()


# 图题/表题识别必须保守。
# 旧规则允许 “Figure 1 shows ...” / “Table 1 shows ...” 这类正文句被误判为 caption，
# 导致导出或预览时被居中、加粗。这里要求编号后出现明确的题注分隔符或行尾。
CAPTION_CJK_RE = re.compile(
    r"^(?:图|表)\s*[\d一二三四五六七八九十IVXivx]+(?=\s*[（(.:：．、\-—]|$)"
)
CAPTION_LATIN_RE = re.compile(
    r"^(?:Figure|Fig\.?|Table)\s*\d+[A-Za-z]?(?=\s*[\(\[.:：．、\-—]|$)",
    re.I,
)


def is_caption_lead_text(text: str) -> bool:
    stripped = (text or "").strip()
    if not stripped:
        return False
    return bool(CAPTION_CJK_RE.match(stripped) or CAPTION_LATIN_RE.match(stripped))


def is_caption_text(text: str) -> bool:
    return is_caption_lead_text(text)


def strip_html_text(text: str) -> str:
    if not text:
        return ""
    text = re.sub(r"<[^>]+>", "", text)
    text = html.unescape(text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def normalize_escaped_tex_delimiters(text: str) -> str:
    """Decode one extra JSON-escape layer only inside complete inline TeX."""

    def replace_formula(match: re.Match) -> str:
        body = re.sub(
            r"\\{2,}",
            lambda run: "\\" * (len(run.group(0)) // 2)
            if len(run.group(0)) % 2 == 0
            else run.group(0),
            match.group(1),
        )
        return rf"\({body}\)"

    return re.sub(r"\\\\\(([^\n]*?)\\\\\)", replace_formula, str(text or ""))


def normalize_translated_inline_html(text: str) -> str:
    result = normalize_escaped_tex_delimiters(text)
    result = re.sub(r"&lt;(\/?)(sup|sub)&gt;", r"<\1\2>", result, flags=re.I)
    result = re.sub(r"&amp;lt;(\/?)(sup|sub)&amp;gt;", r"<\1\2>", result, flags=re.I)
    return result


# Translation must not silently turn mathematics into ordinary prose.  This is
# deliberately a validator, not a formula-repairer: MinerU remains the source
# of truth and a model gets a retry rather than a guessed local rewrite.
MATH_EXPRESSION_RE = re.compile(
    r"\\\[[\s\S]*?\\\]"
    r"|\\\([\s\S]*?\\\)"
    r"|(?<!\\)\$\$[\s\S]*?(?<!\\)\$\$"
    r"|(?<!\\)\$(?![\s$])[^\n$]*(?<!\\)\$(?!\d)"
)


def math_expression_signatures(text: str) -> list[str]:
    """Return source-order math signatures, preserving their delimiters."""
    return [re.sub(r"\s+", " ", match.group(0)).strip() for match in MATH_EXPRESSION_RE.finditer(str(text or ""))]


def math_expression_integrity_issue(source_text: str, translated_text: str) -> str:
    """Describe lost/changed source mathematics without attempting a repair."""
    source = math_expression_signatures(source_text)
    if not source:
        return ""
    translated = math_expression_signatures(translated_text)
    if source == translated:
        return ""
    return (
        f"原文含 {len(source)} 个数学表达式，译文保留 {len(translated)} 个；"
        "行内/行间公式必须保留原 TeX 与定界符。"
    )


def math_expression_retry_issue(source_text: str, translated_text: str) -> str:
    """Return only structural math loss that justifies another paid request.

    Exact differences remain visible through ``math_expression_integrity_issue``.
    This narrower check tolerates extra formulas, harmless whitespace or
    punctuation changes, redundant grouping braces, presentation-only
    ``\\mathrm`` wrappers, OCR list markers, and a source formula split across
    adjacent TeX spans.
    """
    expected = math_expression_signatures(source_text)
    if not expected:
        return ""
    actual = math_expression_signatures(translated_text)

    def body(token: str) -> str:
        value = str(token or "")
        if (
            (value.startswith(r"\(") and value.endswith(r"\)"))
            or (value.startswith(r"\[") and value.endswith(r"\]"))
        ):
            return value[2:-2]
        if value.startswith("$$") and value.endswith("$$"):
            return value[2:-2]
        if value.startswith("$") and value.endswith("$"):
            return value[1:-1]
        return value

    def collapse_redundant_braces(value: str) -> str:
        output = str(value or "")
        previous = None
        while output != previous:
            previous = output
            output = re.sub(r"\{\s*\{([^{}]*)\}\s*\}", r"{\1}", output)
        return output

    def normalized_body(value: str) -> str:
        output = unicodedata.normalize("NFKC", str(value or ""))
        output = re.sub(r"\s+", " ", output).strip()
        output = re.sub(
            r"(?:\\[,;:!]\s*|[,;:]\s+|\s+)\((?:i{1,3}|iv|v|[a-c])\)\s*$",
            "",
            output,
            flags=re.I,
        )
        output = collapse_redundant_braces(output)
        previous = None
        while output != previous:
            previous = output
            output = re.sub(r"\\mathrm\s*\{([^{}]*)\}", r"\1", output)
            output = re.sub(r"\\mathrm\s+([A-Za-z])", r"\1", output)
            output = re.sub(r"\\mathrm(?=[A-Za-z])", "", output)
            output = collapse_redundant_braces(output)
        return re.sub(r"[\s,.;:，。；：、]+", "", output)

    expected_bodies = [normalized_body(body(token)) for token in expected]
    actual_bodies = [normalized_body(body(token)) for token in actual]
    actual_index = 0
    for expected_index, expected_body in enumerate(expected_bodies):
        combined = ""
        first_actual_index = actual_index
        while actual_index < len(actual_bodies):
            proposed = combined + actual_bodies[actual_index]
            if expected_body.startswith(proposed):
                combined = proposed
                actual_index += 1
                if combined == expected_body:
                    break
                continue
            # Additional formula spans remain review warnings, but do not by
            # themselves justify re-sending the complete translation chunk.
            actual_index += 1
        if combined != expected_body:
            candidate_index = min(first_actual_index, len(actual) - 1)
            actual_body = actual_bodies[candidate_index][:180] if candidate_index >= 0 else ""
            return (
                f"公式#{expected_index + 1}数学主体疑似变化（标准化源: {expected_body[:180]}；"
                f"标准化当前值: {actual_body}；当前公式序号: {candidate_index + 1}）"
            )
    return ""


EQUATION_REFERENCE_RE = re.compile(
    r"\b(?:Eq|Eqs|Equation|Equations)\.?\s+"
    r"(?:(?![.;。；]\s).){0,120}"
    r"[（(]\s*[A-Za-z]?\d+[A-Za-z]?\s*[)）]",
    re.I,
)
EQUATION_REFERENCE_NUMBER_RE = re.compile(r"[（(]\s*([A-Za-z]?\d+[A-Za-z]?)\s*[)）]")
GARBLED_PAREN_REFERENCE_RE = re.compile(r"(?<![A-Za-z0-9])(?:[~～〜]\s*)([A-Za-z]?\d+[A-Za-z]?|[a-z])\s*[!！](?![A-Za-z0-9])")


def source_equation_reference_numbers(text: str) -> list[str]:
    numbers: list[str] = []
    for match in EQUATION_REFERENCE_RE.finditer(str(text or "")):
        for number_match in EQUATION_REFERENCE_NUMBER_RE.finditer(match.group(0)):
            number = number_match.group(1).strip()
            if number and number not in numbers:
                numbers.append(number)
    return numbers


def repair_equation_reference_translation(source_text: str, translated_text: str) -> str:
    numbers = source_equation_reference_numbers(source_text)
    result = GARBLED_PAREN_REFERENCE_RE.sub(lambda match: f"({match.group(1)})", str(translated_text or ""))
    if not numbers:
        return result
    for number in numbers:
        escaped = re.escape(number)
        result = re.sub(rf"(?:[~～〜]\s*)?{escaped}\s*[!！]", f"式 ({number})", result)
        result = re.sub(rf"(?<![A-Za-z0-9])式\s*{escaped}(?![A-Za-z0-9])", f"式 ({number})", result)
        result = re.sub(rf"(?<![A-Za-z0-9])方程\s*{escaped}(?![A-Za-z0-9])", f"方程 ({number})", result)
    result = re.sub(r"(?:方程|公式)\s*[。.]\s*(式\s*\()", r"\1", result)
    result = re.sub(r"\)\s*(和|与|及)\s*式", r") \1式", result)
    return result


def looks_like_caption_continuation(text: str) -> bool:
    stripped = (text or "").strip()
    if not stripped:
        return False
    compact = re.sub(r"\s+", " ", stripped)
    if len(compact) > 120:
        return False
    if is_caption_lead_text(compact):
        return False
    if re.match(r"^(?:As shown|Figure|Fig\.?|Table|图|表)\b", compact, re.I):
        return False
    return bool(
        re.match(r"^(?:[\(\[]?[a-z0-9]|at\s+\d|and\s+|or\s+|with\s+|where\s+|when\s+|[（(]?[a-z0-9])", compact, re.I)
        or compact.endswith((".", "。"))
    )


def mark_caption_like_paragraphs(text: str) -> str:
    paragraph_pattern = re.compile(r"<p(?P<attrs>[^>]*)>(?P<body>.*?)</p>", re.I | re.S)
    image_only_pattern = re.compile(r"^\s*(?:<a\b[^>]*>\s*</a>\s*)*(?:<img\b[^>]*>\s*)+\s*$", re.I | re.S)
    table_only_pattern = re.compile(
        r"^\s*(?:<a\b[^>]*>\s*</a>\s*)*(?:<div\b[^>]*class=\"table-wrap\"[^>]*>.*?</div>|<table\b.*?</table>)\s*$",
        re.I | re.S,
    )
    anchor_only_pattern = re.compile(
        r"^\s*(?:<a\b[^>]*class=\"sync-anchor\"[^>]*>\s*</a>|<span\b[^>]*class=\"sync-anchor\"[^>]*>\s*</span>|<a\b[^>]*id=\"doc-block-[^\"]*\"[^>]*>\s*</a>)+\s*$",
        re.I | re.S,
    )
    parts: list[str] = []
    last_index = 0
    recent_visual_anchor = False
    previous_caption = False

    for match in paragraph_pattern.finditer(text):
        parts.append(text[last_index:match.start()])
        attrs = match.group("attrs") or ""
        body = match.group("body") or ""
        raw_text = strip_html_text(body)
        attrs_has_caption = re.search(r'\bclass\s*=\s*"[^"]*\bcaption-like\b', attrs, re.I) is not None
        is_anchor_only = bool(anchor_only_pattern.fullmatch(body))
        is_visual_anchor = bool(image_only_pattern.fullmatch(body) or table_only_pattern.fullmatch(body))
        mark_caption = attrs_has_caption

        if not mark_caption and raw_text:
            if is_caption_lead_text(raw_text) and (recent_visual_anchor or previous_caption):
                mark_caption = True
            elif previous_caption and looks_like_caption_continuation(raw_text):
                mark_caption = True

        if mark_caption and not attrs_has_caption:
            class_match = re.search(r'(\bclass\s*=\s*")([^"]*)(")', attrs, re.I)
            if class_match:
                attrs = attrs[:class_match.start(2)] + class_match.group(2).strip() + " caption-like" + attrs[class_match.end(2):]
            else:
                attrs += ' class="caption-like"'

        parts.append(f"<p{attrs}>{body}</p>")
        last_index = match.end()

        if is_anchor_only:
            continue
        recent_visual_anchor = is_visual_anchor
        previous_caption = mark_caption
        if raw_text and not mark_caption:
            recent_visual_anchor = False

    parts.append(text[last_index:])
    return "".join(parts)


def is_heading_paragraph(paragraph: ET.Element) -> bool:
    return heading_level_for_paragraph(paragraph) > 0


def is_formula_only_paragraph(paragraph: ET.Element) -> bool:
    has_omath = paragraph.find(".//m:oMath", OOXML_NS) is not None or paragraph.find(".//m:oMathPara", OOXML_NS) is not None
    if not has_omath:
        return False
    text = paragraph_text(paragraph)
    if not text:
        return True
    compact = re.sub(r"\s+", "", text)
    return bool(compact) and len(compact) <= 24


def set_paragraph_format(paragraph: ET.Element, style: ExportStyleSettings | None = None):
    style = resolve_export_style(style)
    ppr = ensure_child(paragraph, w_tag("pPr"), before_tags={w_tag("r")})
    spacing = ensure_child(ppr, w_tag("spacing"))
    spacing.set(w_tag("before"), "0")
    spacing.set(w_tag("after"), "0")

    text = paragraph_text(paragraph)
    heading_level = heading_level_for_paragraph(paragraph)
    heading = heading_level > 0
    has_drawing = paragraph.find(".//w:drawing", OOXML_NS) is not None
    has_omath = paragraph.find(".//m:oMath", OOXML_NS) is not None or paragraph.find(".//m:oMathPara", OOXML_NS) is not None
    formula_only = is_formula_only_paragraph(paragraph)
    is_caption = is_caption_text(text)
    if has_drawing or has_omath:
        spacing.set(w_tag("line"), style_line_twips(style))
        spacing.set(w_tag("lineRule"), "atLeast")
    else:
        spacing.set(w_tag("line"), style_line_twips(style))
        spacing.set(w_tag("lineRule"), "auto")
    if has_drawing or is_caption:
        jc = ensure_child(ppr, w_tag("jc"))
        jc.set(w_tag("val"), "center")
    if heading:
        jc = ensure_child(ppr, w_tag("jc"))
        jc.set(w_tag("val"), "left")
    if formula_only:
        jc = ensure_child(ppr, w_tag("jc"))
        jc.set(w_tag("val"), "center")
    elif not heading and not has_drawing and not is_caption:
        jc = ensure_child(ppr, w_tag("jc"))
        jc.set(w_tag("val"), "both")
    if not heading and not has_drawing and not is_caption and not formula_only:
        ind = ensure_child(ppr, w_tag("ind"))
        ind.set(w_tag("firstLine"), style_first_line_twips(style))
    for run in paragraph.findall("./w:r", OOXML_NS):
        if is_caption:
            size_half_points = style_half_points(style.caption_font_pt)
        elif heading:
            size_half_points = heading_size_half_points(style, heading_level)
        else:
            size_half_points = style_half_points(style.body_font_pt)
        set_run_style(run, size_half_points, bold=heading, style=style, heading=heading)


def set_table_format(table: ET.Element, style: ExportStyleSettings | None = None):
    style = resolve_export_style(style)

    def cell_span(cell: ET.Element) -> int:
        tc_pr = cell.find("./w:tcPr", OOXML_NS)
        if tc_pr is None:
            return 1
        grid_span = tc_pr.find("./w:gridSpan", OOXML_NS)
        if grid_span is None:
            return 1
        try:
            return max(1, int(grid_span.get(w_tag("val")) or 1))
        except Exception:
            return 1

    rows = table.findall("./w:tr", OOXML_NS)
    if not rows:
        return

    # 统计每一行的“有效列数”，如果各行一致，说明这是可安全固定宽度的表格。
    row_col_counts: list[int] = []
    for row in rows:
        col_count = 0
        for cell in row.findall("./w:tc", OOXML_NS):
            col_count += cell_span(cell)
        if col_count > 0:
            row_col_counts.append(col_count)

    tbl_pr = ensure_child(table, w_tag("tblPr"), before_tags={w_tag("tr")})

    # 统一居中，但不改变 Pandoc 生成的表格结构。
    jc = ensure_child(tbl_pr, w_tag("jc"))
    jc.set(w_tag("val"), "center")

    tbl_borders = ensure_child(tbl_pr, w_tag("tblBorders"))
    top_border = ensure_child(tbl_borders, w_tag("top"))
    top_border.set(w_tag("val"), "single")
    top_border.set(w_tag("sz"), "12")
    top_border.set(w_tag("space"), "0")
    top_border.set(w_tag("color"), "303840")
    bottom_border = ensure_child(tbl_borders, w_tag("bottom"))
    bottom_border.set(w_tag("val"), "single")
    bottom_border.set(w_tag("sz"), "12")
    bottom_border.set(w_tag("space"), "0")
    bottom_border.set(w_tag("color"), "303840")
    inside_h = ensure_child(tbl_borders, w_tag("insideH"))
    inside_h.set(w_tag("val"), "nil")
    inside_v = ensure_child(tbl_borders, w_tag("insideV"))
    inside_v.set(w_tag("val"), "nil")
    left_border = ensure_child(tbl_borders, w_tag("left"))
    left_border.set(w_tag("val"), "nil")
    right_border = ensure_child(tbl_borders, w_tag("right"))
    right_border.set(w_tag("val"), "nil")

    # 统一段落间距和字体，避免导出后单元格内文本过散。
    for paragraph in table.findall(".//w:p", OOXML_NS):
        ppr = ensure_child(paragraph, w_tag("pPr"), before_tags={w_tag("r")})
        spacing = ensure_child(ppr, w_tag("spacing"))
        spacing.set(w_tag("before"), "0")
        spacing.set(w_tag("after"), "0")
        spacing.set(w_tag("line"), str(max(200, int(round(clamp_int(style.line_spacing_pt, 15, 30, 20) * 16)))))
        spacing.set(w_tag("lineRule"), "auto")
        ind = ppr.find("./w:ind", OOXML_NS)
        if ind is not None:
            ppr.remove(ind)
        jc = ensure_child(ppr, w_tag("jc"))
        jc.set(w_tag("val"), "center")

        for run in paragraph.findall("./w:r", OOXML_NS):
            set_run_style(run, style_half_points(style.body_font_pt), style=style)

    for row_index, row in enumerate(rows):
        tr_pr = ensure_child(row, w_tag("trPr"), before_tags={w_tag("tc")})
        row_height = ensure_child(tr_pr, w_tag("trHeight"))
        row_height.set(w_tag("val"), "1")
        row_height.set(w_tag("hRule"), "atLeast")
        for cell in row.findall("./w:tc", OOXML_NS):
            tc_pr = ensure_child(cell, w_tag("tcPr"))
            tc_mar = ensure_child(tc_pr, w_tag("tcMar"))
            for side in ("top", "bottom", "left", "right"):
                margin = ensure_child(tc_mar, w_tag(side))
                margin.set(w_tag("w"), "36" if side in ("left", "right") else "18")
                margin.set(w_tag("type"), "dxa")
            if row_index == 0:
                tc_borders = ensure_child(tc_pr, w_tag("tcBorders"))
                bottom = ensure_child(tc_borders, w_tag("bottom"))
                bottom.set(w_tag("val"), "single")
                bottom.set(w_tag("sz"), "8")
                bottom.set(w_tag("space"), "0")
                bottom.set(w_tag("color"), "5D6872")

    # 仅对“列数稳定”的普通表格做轻量统一，不重写 Pandoc 已经生成好的列宽。
    if not row_col_counts:
        return

    column_count = max(row_col_counts)
    if column_count < 2 or len(set(row_col_counts)) != 1:
        # 列数不稳定或疑似复杂合并表格：只做样式，不强行改宽。
        return

    tbl_grid = table.find("./w:tblGrid", OOXML_NS)
    grid_widths: list[int] = []
    if tbl_grid is not None:
        for grid_col in tbl_grid.findall("./w:gridCol", OOXML_NS):
            try:
                width = int(grid_col.get(w_tag("w")) or 0)
            except Exception:
                width = 0
            if width > 0:
                grid_widths.append(width)

    # Keep stable column widths supplied by Pandoc instead of replacing them
    # with equal-width columns.
    if len(grid_widths) != column_count:
        return

    tbl_layout = ensure_child(tbl_pr, w_tag("tblLayout"))
    tbl_layout.set(w_tag("type"), "fixed")

    tbl_w = ensure_child(tbl_pr, w_tag("tblW"))
    tbl_w.set(w_tag("type"), "dxa")
    tbl_w.set(w_tag("w"), str(sum(grid_widths)))

    tbl_ind = ensure_child(tbl_pr, w_tag("tblInd"))
    tbl_ind.set(w_tag("type"), "dxa")
    tbl_ind.set(w_tag("w"), "0")

    # 给每个单元格补齐对应宽度，避免 Word 重新推导宽度。
    for row in rows:
        col_index = 0
        for cell in row.findall("./w:tc", OOXML_NS):
            span = cell_span(cell)
            span_end = min(column_count, col_index + span)
            cell_width = sum(grid_widths[col_index:span_end]) if col_index < column_count else grid_widths[-1]
            if cell_width <= 0:
                cell_width = grid_widths[min(col_index, len(grid_widths) - 1)]

            tc_pr = ensure_child(cell, w_tag("tcPr"))
            tc_w = ensure_child(tc_pr, w_tag("tcW"))
            tc_w.set(w_tag("type"), "dxa")
            tc_w.set(w_tag("w"), str(cell_width))

            col_index += span


def estimate_table_column_widths(rows: list[list[HtmlTableCell]], total_width: int = 9000) -> list[int]:
    occupancy: dict[tuple[int, int], tuple[int, int, HtmlTableCell]] = {}
    start_map: dict[tuple[int, int], HtmlTableCell] = {}
    total_rows = 0
    total_cols = 0
    for row_index, row in enumerate(rows):
        col_index = 0
        while (row_index, col_index) in occupancy:
            col_index += 1
        for cell in row:
            while (row_index, col_index) in occupancy:
                col_index += 1
            start_map[(row_index, col_index)] = cell
            for row_offset in range(cell.rowspan):
                for col_offset in range(cell.colspan):
                    occupancy[(row_index + row_offset, col_index + col_offset)] = (row_index, col_index, cell)
            total_cols = max(total_cols, col_index + cell.colspan)
            col_index += cell.colspan
        total_rows = max(total_rows, row_index + 1)
    total_rows = max(total_rows, max((pos[0] + 1 for pos in occupancy), default=0))
    if total_cols <= 0:
        return []

    weights = [1.0] * total_cols
    for (row_index, col_index), cell in start_map.items():
        text_weight = max(1.0, min(14.0, len(cell.text.replace("\n", " ")) / 3.5))
        portion = text_weight / max(1, cell.colspan)
        for offset in range(cell.colspan):
            weights[col_index + offset] = max(weights[col_index + offset], portion)

    total_weight = sum(weights) or float(total_cols)
    widths = [max(720, int(total_width * (weight / total_weight))) for weight in weights]
    delta = total_width - sum(widths)
    widths[-1] += delta
    return widths


def build_ooxml_table(rows: list[list[HtmlTableCell]]) -> ET.Element | None:
    if not rows:
        return None

    occupancy: dict[tuple[int, int], tuple[int, int, HtmlTableCell]] = {}
    start_map: dict[tuple[int, int], HtmlTableCell] = {}
    total_rows = 0
    total_cols = 0
    for row_index, row in enumerate(rows):
        col_index = 0
        for cell in row:
            while (row_index, col_index) in occupancy:
                col_index += 1
            start_map[(row_index, col_index)] = cell
            for row_offset in range(cell.rowspan):
                for col_offset in range(cell.colspan):
                    occupancy[(row_index + row_offset, col_index + col_offset)] = (row_index, col_index, cell)
            total_rows = max(total_rows, row_index + cell.rowspan)
            total_cols = max(total_cols, col_index + cell.colspan)
            col_index += cell.colspan

    if total_cols <= 0:
        return None

    column_widths = estimate_table_column_widths(rows)
    if len(column_widths) != total_cols:
        column_widths = [max(9000 // total_cols, 720)] * total_cols
        column_widths[-1] += 9000 - sum(column_widths)

    table = ET.Element(w_tag("tbl"))
    tbl_pr = ET.SubElement(table, w_tag("tblPr"))
    tbl_style = ET.SubElement(tbl_pr, w_tag("tblStyle"))
    tbl_style.set(w_tag("val"), "TableGrid")
    tbl_w = ET.SubElement(tbl_pr, w_tag("tblW"))
    tbl_w.set(w_tag("type"), "dxa")
    tbl_w.set(w_tag("w"), str(sum(column_widths)))
    tbl_layout = ET.SubElement(tbl_pr, w_tag("tblLayout"))
    tbl_layout.set(w_tag("type"), "fixed")
    tbl_jc = ET.SubElement(tbl_pr, w_tag("jc"))
    tbl_jc.set(w_tag("val"), "center")

    tbl_grid = ET.SubElement(table, w_tag("tblGrid"))
    for width in column_widths:
        grid_col = ET.SubElement(tbl_grid, w_tag("gridCol"))
        grid_col.set(w_tag("w"), str(width))

    for row_index in range(total_rows):
        row = ET.SubElement(table, w_tag("tr"))
        col_index = 0
        while col_index < total_cols:
            source = occupancy.get((row_index, col_index))
            if source is None:
                tc = ET.SubElement(row, w_tag("tc"))
                tc_pr = ET.SubElement(tc, w_tag("tcPr"))
                tc_w = ET.SubElement(tc_pr, w_tag("tcW"))
                tc_w.set(w_tag("type"), "dxa")
                tc_w.set(w_tag("w"), str(column_widths[col_index]))
                v_align = ET.SubElement(tc_pr, w_tag("vAlign"))
                v_align.set(w_tag("val"), "center")
                ET.SubElement(tc, w_tag("p"))
                col_index += 1
                continue

            start_row, start_col, cell = source
            if start_col < col_index:
                col_index += 1
                continue
            if start_row == row_index:
                text = cell.text
                restart_vmerge = cell.rowspan > 1
            else:
                if start_col != col_index:
                    col_index += 1
                    continue
                text = ""
                restart_vmerge = False

            cell_width = sum(column_widths[start_col:start_col + cell.colspan])
            tc = ET.SubElement(row, w_tag("tc"))
            tc_pr = ET.SubElement(tc, w_tag("tcPr"))
            tc_w = ET.SubElement(tc_pr, w_tag("tcW"))
            tc_w.set(w_tag("type"), "dxa")
            tc_w.set(w_tag("w"), str(cell_width))
            if cell.colspan > 1:
                grid_span = ET.SubElement(tc_pr, w_tag("gridSpan"))
                grid_span.set(w_tag("val"), str(cell.colspan))
            if cell.rowspan > 1:
                v_merge = ET.SubElement(tc_pr, w_tag("vMerge"))
                if restart_vmerge:
                    v_merge.set(w_tag("val"), "restart")
            v_align = ET.SubElement(tc_pr, w_tag("vAlign"))
            v_align.set(w_tag("val"), "center")

            paragraph = ET.SubElement(tc, w_tag("p"))
            ppr = ET.SubElement(paragraph, w_tag("pPr"))
            spacing = ET.SubElement(ppr, w_tag("spacing"))
            spacing.set(w_tag("before"), "0")
            spacing.set(w_tag("after"), "0")
            spacing.set(w_tag("line"), "240")
            spacing.set(w_tag("lineRule"), "auto")
            jc = ET.SubElement(ppr, w_tag("jc"))
            jc.set(w_tag("val"), "center")
            append_word_text_segments(paragraph, text, bold=cell.is_header)
            col_index = start_col + cell.colspan

    set_table_format(table)
    return table


def looks_like_formula_text(text: str) -> bool:
    stripped = (text or "").strip()
    if not stripped:
        return False
    return any(token in stripped for token in ("\\", "_", "^", "{", "}"))


def normalize_texish_for_word_math(text: str) -> str:
    if not text:
        return ""
    normalized = text.replace("\r", "").replace("\n", " ")
    normalized = re.sub(r"\s+", " ", normalized)
    normalized = re.sub(r"\s+}", "}", normalized)
    normalized = re.sub(r"{\s+", "{", normalized)
    normalized = re.sub(r"\s+([_^/])", r"\1", normalized)
    normalized = re.sub(r"([_^/])\s+", r"\1", normalized)
    return normalized.strip()


WORD_OMATH_SEED_RE = re.compile(r"\\[A-Za-z]+|[A-Za-z0-9\\{}().%/]*[_^][A-Za-z0-9\\{}().%/{}^_]*")
WORD_OMATH_CONTIGUOUS_CHARS = set(
    "ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789\\{}_^().%/+-=<>·×~"
)


def split_text_for_word_omath(text: str) -> list[WordTableRefineChunk]:
    text = normalize_texish_for_word_math(text)
    if not text:
        return [WordTableRefineChunk(kind="text", text="")]

    def extend_formula_span(source: str, start: int, end: int) -> tuple[int, int]:
        while start > 0 and source[start - 1] in WORD_OMATH_CONTIGUOUS_CHARS:
            start -= 1
        while end < len(source) and source[end] in WORD_OMATH_CONTIGUOUS_CHARS:
            end += 1
        while True:
            tail = source[end:]
            operator_match = re.match(r"\s*[+\-=/<>·×~]\s*[A-Za-z0-9\\{}_^().%/]+", tail)
            if operator_match:
                end += operator_match.end()
                continue
            current = source[start:end]
            slash_tail = re.match(r"\s+[A-Za-z0-9().%/]+", tail)
            if slash_tail and looks_like_formula_text(current) and "/" in slash_tail.group(0):
                end += slash_tail.end()
                continue
            break
        return start, end

    spans: list[tuple[int, int]] = []
    for match in WORD_OMATH_SEED_RE.finditer(text):
        start, end = extend_formula_span(text, match.start(), match.end())
        if spans and start <= spans[-1][1]:
            spans[-1] = (spans[-1][0], max(spans[-1][1], end))
        else:
            spans.append((start, end))

    if not spans:
        return [WordTableRefineChunk(kind="text", text=text)]

    chunks: list[WordTableRefineChunk] = []
    cursor = 0
    for start, end in spans:
        if cursor < start:
            chunks.append(WordTableRefineChunk(kind="text", text=text[cursor:start]))
        formula_text = text[start:end]
        chunks.append(WordTableRefineChunk(kind="formula", text=formula_text))
        cursor = end
    if cursor < len(text):
        chunks.append(WordTableRefineChunk(kind="text", text=text[cursor:]))

    merged: list[WordTableRefineChunk] = []
    for chunk in chunks:
        if not chunk.text:
            continue
        if merged and merged[-1].kind == chunk.kind:
            merged[-1].text += chunk.text
        else:
            merged.append(WordTableRefineChunk(kind=chunk.kind, text=chunk.text))
    return merged or [WordTableRefineChunk(kind="text", text=text)]


def texish_visible_text(text: str) -> str:
    return "".join(segment.text for segment in parse_texish_segments(text)) or text


def formula_span_dicts_for_word_omath(
    text: str,
    *,
    use_texish_visible_text: bool = False,
) -> list[dict[str, int | str]] | None:
    chunks = split_text_for_word_omath(text)
    spans: list[dict[str, int | str]] = []
    cursor = 0
    for chunk in chunks:
        visible_text = texish_visible_text(chunk.text) if use_texish_visible_text else chunk.text
        end = cursor + len(visible_text)
        if chunk.kind == "formula" and chunk.text.strip():
            spans.append(
                {
                    "kind": chunk.kind,
                    "text": chunk.text,
                    "visible_text": visible_text,
                    "normalized_text": normalize_texish_for_word_math(chunk.text),
                    "normalized_visible_text": normalize_texish_for_word_math(visible_text),
                    "start": cursor,
                    "end": end,
                }
            )
        cursor = end
    return spans or None


def materialize_word_table_formula_targets(rows: list[list[HtmlTableCell]]) -> list[WordTableFormulaTarget]:
    occupancy: dict[tuple[int, int], tuple[int, int, HtmlTableCell]] = {}
    total_rows = 0
    total_cols = 0
    for row_index, row in enumerate(rows):
        col_index = 0
        for cell in row:
            while (row_index, col_index) in occupancy:
                col_index += 1
            for row_offset in range(cell.rowspan):
                for col_offset in range(cell.colspan):
                    occupancy[(row_index + row_offset, col_index + col_offset)] = (row_index, col_index, cell)
            total_rows = max(total_rows, row_index + cell.rowspan)
            total_cols = max(total_cols, col_index + cell.colspan)
            col_index += cell.colspan

    targets: list[WordTableFormulaTarget] = []
    for row_index in range(total_rows):
        col_index = 0
        while col_index < total_cols:
            source = occupancy.get((row_index, col_index))
            if source is None:
                col_index += 1
                continue
            start_row, start_col, cell = source
            if start_col < col_index:
                col_index += 1
                continue
            if start_row == row_index:
                targets.append(WordTableFormulaTarget(text=cell.text))
            col_index = start_col + cell.colspan
    return targets


def build_word_table_refine_targets(rows: list[list[HtmlTableCell]]) -> list[list[dict[str, int | str]]]:
    targets: list[list[dict[str, int | str]]] = []
    for item in materialize_word_table_formula_targets(rows):
        targets.append(formula_span_dicts_for_word_omath(item.text, use_texish_visible_text=True) or [])
    return targets


def build_word_body_refine_targets(path: Path) -> list[list[dict[str, int | str]] | None]:
    if not path.exists():
        return []
    try:
        with zipfile.ZipFile(path, "r") as archive:
            document_xml = archive.read("word/document.xml")
    except Exception:
        return []

    try:
        root = ET.fromstring(document_xml)
    except ET.ParseError:
        return []

    body = root.find("./w:body", OOXML_NS)
    if body is None:
        return []

    targets: list[list[dict[str, int | str]] | None] = []
    for child in body:
        if child.tag != w_tag("p"):
            continue
        if child.find(".//w:drawing", OOXML_NS) is not None:
            targets.append(None)
            continue
        if child.find(".//m:oMath", OOXML_NS) is not None or child.find(".//m:oMathPara", OOXML_NS) is not None:
            targets.append(None)
            continue
        if child.find(".//w:fldSimple", OOXML_NS) is not None or child.find(".//w:instrText", OOXML_NS) is not None:
            targets.append(None)
            continue
        text = paragraph_text(child)
        if not text or is_heading_paragraph(child) or is_caption_text(text):
            targets.append(None)
            continue
        targets.append(formula_span_dicts_for_word_omath(text))
    return targets


def refine_docx_tables_with_word_omath(
    path: Path,
    placeholders: list[HtmlTablePlaceholder],
    log=None,
) -> tuple[bool, str]:
    path = path.resolve()
    if not path.exists():
        return False, "Word 导出文件不存在，无法精校。"

    table_targets = [build_word_table_refine_targets(item.rows) for item in placeholders]
    body_targets = build_word_body_refine_targets(path)
    has_table_formula = any(
        chunk.get("kind") == "formula"
        for table in table_targets
        for cell_chunks in table
        for chunk in cell_chunks
    )
    has_body_formula = any(
        chunk and any(part.get("kind") == "formula" for part in chunk)
        for chunk in body_targets
    )
    if not has_table_formula and not has_body_formula:
        return False, "文档中没有检测到需要升级为 Word 公式的内容。"

    mapping_path = path.with_suffix(path.suffix + ".omath-map.json")
    script_path = path.with_suffix(path.suffix + ".omath-refine.ps1")
    mapping = {
        "tables": table_targets,
        "paragraphs": body_targets,
    }
    mapping_path.write_text(json.dumps(mapping, ensure_ascii=False), encoding="utf-8")

    script = "\n".join([
        "$ErrorActionPreference = 'Stop'",
        "$OutputEncoding = [Console]::OutputEncoding = [Text.UTF8Encoding]::new()",
        "$docPath = $args[0]",
        "$mapPath = $args[1]",
        "$targets = Get-Content -LiteralPath $mapPath -Raw -Encoding UTF8 | ConvertFrom-Json",
        "$tableTargets = @()",
        "$paragraphTargets = @()",
        "if ($targets.tables -ne $null) { $tableTargets = @($targets.tables) }",
        "if ($targets.paragraphs -ne $null) { $paragraphTargets = @($targets.paragraphs) }",
        "$word = $null",
        "$doc = $null",
        "try {",
        "    $word = New-Object -ComObject Word.Application",
        "    $word.Visible = $false",
        "    $word.DisplayAlerts = 0",
        "    $doc = $word.Documents.Open($docPath, $false, $false)",
        "    $tableCount = [Math]::Min($doc.Tables.Count, $tableTargets.Count)",
        "    $tableUpdated = 0",
        "    $bodyUpdated = 0",
        "    for ($t = 1; $t -le $tableCount; $t++) {",
        "        $table = $doc.Tables.Item($t)",
        "        $cells = @($tableTargets[$t - 1])",
        "        $cellCount = [Math]::Min($table.Range.Cells.Count, $cells.Count)",
        "        for ($i = 1; $i -le $cellCount; $i++) {",
        "            $chunks = @($cells[$i - 1])",
        "            $hasFormula = $false",
        "            foreach ($chunk in $chunks) {",
        "                if ([string]$chunk.kind -eq 'formula' -and -not [string]::IsNullOrWhiteSpace([string]$chunk.text)) {",
        "                    $hasFormula = $true",
        "                    break",
        "                }",
        "            }",
        "            if (-not $hasFormula) { continue }",
        "            $cellRange = $table.Range.Cells.Item($i).Range.Duplicate",
        "            $cellRange.End = [Math]::Max($cellRange.Start, $cellRange.End - 1)",
        "            for ($chunkIndex = $chunks.Count - 1; $chunkIndex -ge 0; $chunkIndex--) {",
        "                $chunk = $chunks[$chunkIndex]",
        "                $chunkText = [string]$chunk.text",
        "                $visibleText = [string]$chunk.visible_text",
        "                $normalizedChunkText = [string]$chunk.normalized_text",
        "                $normalizedVisibleText = [string]$chunk.normalized_visible_text",
        "                if ([string]::IsNullOrWhiteSpace($chunkText) -or [string]::IsNullOrWhiteSpace($visibleText)) { continue }",
        "                $startOffset = [int]$chunk.start",
        "                $endOffset = [int]$chunk.end",
        "                if ($endOffset -le $startOffset) { continue }",
        "                $targetRange = $doc.Range($cellRange.Start + $startOffset, $cellRange.Start + $endOffset)",
        "                $targetText = [string]$targetRange.Text",
        "                $normalizedTargetText = (($targetText -replace '\\r','') -replace '\\n',' ') -replace '\\s+',' '",
        "                $normalizedTargetText = $normalizedTargetText.Trim()",
        "                if ($targetText -eq $visibleText -or $normalizedTargetText -eq $normalizedVisibleText -or $normalizedTargetText -eq $normalizedChunkText) {",
        "                    $targetRange.Text = $chunkText",
        "                    $formulaRange = $doc.Range($targetRange.Start, $targetRange.Start + $chunkText.Length)",
        "                    $formulaRange.OMaths.Add($formulaRange) | Out-Null",
        "                    $formulaRange.OMaths.BuildUp()",
        "                    $tableUpdated++",
        "                }",
        "            }",
        "        }",
        "    }",
        "    $bodyParagraphIndex = 0",
        "    foreach ($paragraph in $doc.Paragraphs) {",
        "        $range = $paragraph.Range",
        "        if ($range -eq $null) { continue }",
        "        if ($range.Information(12)) { continue }",
        "        if ($bodyParagraphIndex -ge $paragraphTargets.Count) { break }",
        "        $chunks = $paragraphTargets[$bodyParagraphIndex]",
        "        $bodyParagraphIndex++",
        "        if ($chunks -eq $null) { continue }",
        "        if ($range.OMaths.Count -gt 0) { continue }",
        "        $workingRange = $range.Duplicate",
        "        $workingRange.End = [Math]::Max($workingRange.Start, $workingRange.End - 1)",
        "        for ($chunkIndex = $chunks.Count - 1; $chunkIndex -ge 0; $chunkIndex--) {",
        "            $chunk = $chunks[$chunkIndex]",
        "            $chunkText = [string]$chunk.text",
        "            $visibleText = [string]$chunk.visible_text",
        "            $normalizedChunkText = [string]$chunk.normalized_text",
        "            $normalizedVisibleText = [string]$chunk.normalized_visible_text",
        "            if ([string]::IsNullOrWhiteSpace($chunkText)) { continue }",
        "            $startOffset = [int]$chunk.start",
        "            $endOffset = [int]$chunk.end",
        "            if ($endOffset -le $startOffset) { continue }",
        "            $matchRange = $doc.Range($workingRange.Start + $startOffset, $workingRange.Start + $endOffset)",
        "            $matchText = [string]$matchRange.Text",
        "            $normalizedMatchText = (($matchText -replace '\\r','') -replace '\\n',' ') -replace '\\s+',' '",
        "            $normalizedMatchText = $normalizedMatchText.Trim()",
        "            if ($matchText -eq $chunkText -or $matchText -eq $visibleText -or $normalizedMatchText -eq $normalizedChunkText -or $normalizedMatchText -eq $normalizedVisibleText) {",
        "                if ($matchText -ne $chunkText) { $matchRange.Text = $chunkText; $matchRange = $doc.Range($matchRange.Start, $matchRange.Start + $chunkText.Length) }",
        "                $matchRange.OMaths.Add($matchRange) | Out-Null",
        "                $matchRange.OMaths.BuildUp()",
        "                $bodyUpdated++",
        "            }",
        "        }",
        "    }",
        "    $doc.Save()",
        "    Write-Output ('TABLE_UPDATED=' + $tableUpdated + ';BODY_UPDATED=' + $bodyUpdated)",
        "} finally {",
        "    if ($doc -ne $null) { $doc.Close() }",
        "    if ($word -ne $null) { $word.Quit() }",
        "}",
    ])
    script_path.write_text(script, encoding="utf-8")
    cmd = [
        "powershell",
        "-NoProfile",
        "-ExecutionPolicy",
        "Bypass",
        "-File",
        str(script_path),
        str(path),
        str(mapping_path),
    ]
    try:
        result = subprocess.run(
            cmd,
            check=True,
            capture_output=True,
            text=True,
            timeout=360,
            cwd=str(path.parent),
            **hidden_subprocess_kwargs(),
        )
        message = (result.stdout or result.stderr or "").strip() or "Word 精校已完成。"
        if log:
            log(f"Word 格式优化：{message}")
        return True, message
    except subprocess.CalledProcessError as exc:
        detail = (exc.stderr or exc.stdout or str(exc)).strip()
        if log:
            log(f"Word 精校失败，保留当前导出结果: {detail}")
        return False, detail
    except Exception as exc:
        detail = str(exc).strip()
        if log:
            log(f"Word 精校失败，保留当前导出结果: {detail}")
        return False, detail
    finally:
        for temp_path in (mapping_path, script_path):
            try:
                temp_path.unlink()
            except OSError:
                pass


def replace_docx_placeholders_with_tables(path: Path, placeholders: list[HtmlTablePlaceholder]) -> None:
    if not path.exists() or not placeholders:
        return

    placeholder_map = {item.marker: item for item in placeholders}
    document_name = "word/document.xml"
    with zipfile.ZipFile(path, "r") as archive:
        if document_name not in archive.namelist():
            return
        document_xml = archive.read(document_name)

    root = ET.fromstring(document_xml)
    body = root.find("./w:body", OOXML_NS)
    if body is None:
        return

    children = list(body)
    for index, child in enumerate(children):
        if child.tag != w_tag("p"):
            continue
        marker = paragraph_text(child)
        placeholder = placeholder_map.get(marker)
        if not placeholder:
            continue
        table = build_ooxml_table(placeholder.rows)
        if table is None:
            continue
        body.remove(child)
        body.insert(index, table)

    rewrite_zip_members(
        path,
        {document_name: ET.tostring(root, encoding="utf-8", xml_declaration=True)},
    )


def rewrite_zip_members(path: Path, replacements: dict[str, bytes]) -> None:
    """Rewrite selected ZIP members without loading the complete archive."""
    tmp_path = path.with_suffix(path.suffix + ".tmp")
    try:
        with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(
            tmp_path,
            "w",
            compression=zipfile.ZIP_DEFLATED,
        ) as target:
            for info in source.infolist():
                replacement = replacements.get(info.filename)
                if replacement is not None:
                    target.writestr(info, replacement)
                    continue
                with source.open(info, "r") as input_handle, target.open(info, "w") as output_handle:
                    shutil.copyfileobj(input_handle, output_handle, length=1024 * 1024)
        tmp_path.replace(path)
    finally:
        try:
            tmp_path.unlink(missing_ok=True)
        except OSError:
            pass


def set_docx_image_width(root: ET.Element, target_cx: int | None = None, style: ExportStyleSettings | None = None):
    """Constrain oversized Word images without enlarging source-sized ones.

    Pandoc has already converted each source-page percentage from Markdown to
    an OOXML extent. Replacing every extent with one configured width destroys
    those per-image ratios, so post-processing now only guards against overflow.
    """
    if target_cx is None:
        target_cx = export_style_text_width_emu()
    for extent in root.findall(".//wp:extent", OOXML_NS):
        try:
            old_cx = int(extent.get("cx") or 0)
            old_cy = int(extent.get("cy") or 0)
        except ValueError:
            continue
        if old_cx <= 0 or old_cy <= 0:
            continue
        if old_cx <= target_cx:
            continue
        new_cy = max(1, int(old_cy * (target_cx / old_cx)))
        extent.set("cx", str(target_cx))
        extent.set("cy", str(new_cy))
        inline = extent.getparent() if hasattr(extent, "getparent") else None
    for graphic_extent in root.findall(".//a:ext", OOXML_NS):
        try:
            old_cx = int(graphic_extent.get("cx") or 0)
            old_cy = int(graphic_extent.get("cy") or 0)
        except ValueError:
            continue
        if old_cx <= 0 or old_cy <= 0:
            continue
        if old_cx <= target_cx:
            continue
        new_cy = max(1, int(old_cy * (target_cx / old_cx)))
        graphic_extent.set("cx", str(target_cx))
        graphic_extent.set("cy", str(new_cy))


def postprocess_exported_docx(path: Path, style: ExportStyleSettings | None = None):
    """
    导出 Word 后的后处理。

    说明：
    1. 优先使用更轻量、对表格结构更安全的 OOXML 处理。
    2. python-docx 方案对复杂表格、合并单元格、嵌套表格更容易引发表格错乱。
    3. 如果 OOXML 后处理失败，再回退到 python-docx。
    """
    try:
        postprocess_exported_docx_ooxml(path, style)
        return
    except Exception:
        pass

    postprocess_exported_docx_with_python_docx(path, style)


def postprocess_exported_docx_with_python_docx(path: Path, style: ExportStyleSettings | None = None):
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    style = resolve_export_style(style)
    document = Document(str(path))

    for section in document.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)

    def set_run_font(run, *, heading: bool = False):
        run.font.name = style.heading_font_latin if heading else style.body_font_latin
        run.font.size = Pt(style.heading1_pt if heading else style.body_font_pt)
        run.font.bold = bool(heading)
        run.font.color.rgb = RGBColor(0, 0, 0)
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), style.heading_font_latin if heading else style.body_font_latin)
        rfonts.set(qn("w:hAnsi"), style.heading_font_latin if heading else style.body_font_latin)
        rfonts.set(qn("w:eastAsia"), style.heading_font_cjk if heading else style.body_font_cjk)

    def set_paragraph_spacing(paragraph, *, center: bool = False, heading: bool = False):
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = Pt(style.line_spacing_pt)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
        if not center and not heading:
            paragraph.paragraph_format.first_line_indent = Cm(style.first_line_indent_cm)
        else:
            paragraph.paragraph_format.first_line_indent = Cm(0)
        has_math = bool(paragraph._element.xpath(".//m:oMath | .//m:oMathPara"))
        compact = re.sub(r"\s+", "", paragraph.text or "")
        formula_only = has_math and (not compact or len(compact) <= 24)
        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER
            if center or formula_only
            else WD_ALIGN_PARAGRAPH.LEFT
            if heading
            else WD_ALIGN_PARAGRAPH.JUSTIFY
        )
        for run in paragraph.runs:
            set_run_font(run, heading=heading)

    for paragraph in document.paragraphs:
        style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
        style_display_name = (paragraph.style.name or "") if paragraph.style else ""
        is_heading = bool(paragraph.style) and (
            style_name.startswith("heading") or style_display_name.startswith("标题")
        )
        is_caption = is_caption_text(paragraph.text.strip())
        has_picture = bool(paragraph._element.xpath(".//w:drawing"))
        set_paragraph_spacing(paragraph, center=is_caption or has_picture, heading=is_heading)

    usable_width = document.sections[0].page_width - document.sections[0].left_margin - document.sections[0].right_margin
    for inline_shape in document.inline_shapes:
        if inline_shape.width and inline_shape.width > usable_width:
            ratio = inline_shape.height / inline_shape.width if inline_shape.height else 1
            inline_shape.width = usable_width
            inline_shape.height = int(usable_width * ratio)

    for table in document.tables:
        # 表格结构尽量保持原样，避免复杂表格、合并单元格被强制重排后发生错乱。
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        try:
            table.style = "Table Grid"
        except Exception:
            pass

        # 仅做轻量格式统一，不改写单元格宽度和表格总宽度。
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    # 保持表格内部文本可读，但不要强制改写列宽。
                    paragraph.paragraph_format.first_line_indent = Cm(0)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_paragraph_spacing(paragraph, center=True)

    document.save(str(path))


def postprocess_exported_docx_ooxml(path: Path, style: ExportStyleSettings | None = None):
    if not path.exists():
        return
    style = resolve_export_style(style)
    document_name = "word/document.xml"
    styles_name = "word/styles.xml"
    with zipfile.ZipFile(path, "r") as archive:
        names = set(archive.namelist())
        if document_name not in names:
            return
        document_xml = archive.read(document_name)
        styles_xml = archive.read(styles_name) if styles_name in names else None

    root = ET.fromstring(document_xml)
    for paragraph in root.findall(".//w:p", OOXML_NS):
        set_paragraph_format(paragraph, style)
    for table in root.findall(".//w:tbl", OOXML_NS):
        set_table_format(table, style)
    set_docx_image_width(root, style=style)
    replacements = {
        document_name: ET.tostring(root, encoding="utf-8", xml_declaration=True),
    }

    if styles_xml is not None:
        styles_root = ET.fromstring(styles_xml)
        for style in styles_root.findall(".//w:style", OOXML_NS):
            style_id = (style.get(w_tag("styleId")) or "").lower()
            rpr = ensure_child(style, w_tag("rPr"))
            fonts = ensure_child(rpr, w_tag("rFonts"))
            color = ensure_child(rpr, w_tag("color"))
            color.set(w_tag("val"), "000000")
            if style_id.startswith("heading") or style_id.startswith("title"):
                fonts.set(w_tag("ascii"), style.heading_font_latin)
                fonts.set(w_tag("hAnsi"), style.heading_font_latin)
                fonts.set(w_tag("eastAsia"), style.heading_font_cjk)
                ensure_child(rpr, w_tag("b"))
                size = ensure_child(rpr, w_tag("sz"))
                size.set(w_tag("val"), heading_size_half_points(style, 1))
                size_cs = ensure_child(rpr, w_tag("szCs"))
                size_cs.set(w_tag("val"), heading_size_half_points(style, 1))
            else:
                fonts.set(w_tag("ascii"), style.body_font_latin)
                fonts.set(w_tag("hAnsi"), style.body_font_latin)
                fonts.set(w_tag("eastAsia"), style.body_font_cjk)
                size = ensure_child(rpr, w_tag("sz"))
                size.set(w_tag("val"), style_half_points(style.body_font_pt))
                size_cs = ensure_child(rpr, w_tag("szCs"))
                size_cs.set(w_tag("val"), style_half_points(style.body_font_pt))
            ppr = ensure_child(style, w_tag("pPr"))
            spacing = ensure_child(ppr, w_tag("spacing"))
            spacing.set(w_tag("before"), "0")
            spacing.set(w_tag("after"), "0")
            spacing.set(w_tag("line"), style_line_twips(style))
            spacing.set(w_tag("lineRule"), "exact")
        replacements[styles_name] = ET.tostring(styles_root, encoding="utf-8", xml_declaration=True)

    rewrite_zip_members(path, replacements)


def scan_parsed_docs(root: Path) -> list[ParsedDoc]:
    docs: list[ParsedDoc] = []
    for md_path in root.rglob("full.cleaned.md"):
        if ".venv" in md_path.parts or "mineru_result" in md_path.parts:
            continue
        folder = md_path.parent
        if not is_generated_output_dir(folder):
            continue
        source_pdf = ""
        meta_path = folder / "mineru_task.json"
        if meta_path.exists():
            try:
                meta = json.loads(meta_path.read_text(encoding="utf-8", errors="replace"))
                source_pdf = str(meta.get("source_display_name") or meta.get("source_pdf") or "")
            except json.JSONDecodeError:
                pass
        title = Path(source_pdf).stem if source_pdf else folder.name
        title = f"{title}{translation_status_badge(md_path)}"
        docs.append(ParsedDoc(title=title, folder=folder, markdown_path=md_path, source_pdf=source_pdf))
    for md_path in root.rglob("full.*.md"):
        if ".venv" in md_path.parts or "mineru_result" in md_path.parts:
            continue
        if md_path.name in {"full.md", "full.cleaned.md"}:
            continue
        if not is_generated_output_dir(md_path.parent):
            continue
        if (md_path.parent / "full.cleaned.md").exists():
            continue
        docs.append(ParsedDoc(title=f"{md_path.parent.name} (译文)", folder=md_path.parent, markdown_path=md_path))
    return sorted(docs, key=lambda item: item.markdown_path.stat().st_mtime, reverse=True)


class MinerUWorker(QThread):
    log_signal = Signal(str)
    progress_signal = Signal(int)
    finished_signal = Signal(bool, str, str)

    def __init__(self, pdf_path: str, output_dir: str, options: ParseOptions):
        super().__init__()
        self.pdf_path = Path(pdf_path)
        self.output_dir = Path(output_dir)
        self.options = options

    def log(self, message: str) -> None:
        self.log_signal.emit(message)

    def run(self) -> None:
        try:
            self.output_dir.mkdir(parents=True, exist_ok=True)
            self.progress_signal.emit(5)
            if self.isInterruptionRequested():
                raise MinerUError("用户已停止解析。")

            size_mb = self.pdf_path.stat().st_size / 1024 / 1024
            original_copy = stored_original_path(self.output_dir, self.pdf_path)
            mark_generated_output_dir(self.output_dir, self.pdf_path, self.options)
            if self.pdf_path.resolve() != original_copy.resolve():
                shutil.copy2(self.pdf_path, original_copy)
            self.log(f"已选择解析文件：{self.pdf_path.name} ({size_mb:.2f} MB)")
            self.log("正在验证 MinerU 访问令牌…")
            token = load_mineru_token()

            meta_path = self.output_dir / "mineru_task.json"
            part_work_root = Path(tempfile.mkdtemp(prefix="mineru_part_results_"))
            part_results: list[dict] = []
            part_metadata: list[dict] = []
            try:
                with temporary_mineru_upload_parts(self.pdf_path, log=self.log) as upload_parts:
                    total_parts = len(upload_parts)
                    for part_index, upload_part in enumerate(upload_parts, start=1):
                        if self.isInterruptionRequested():
                            raise MinerUError("用户已停止解析。")
                        prefix = f"分片 {part_index}/{total_parts}" if total_parts > 1 else "文档"
                        self.log(
                            f"{prefix}：第 {upload_part.start_page}-{upload_part.end_page} 页，"
                            f"正在创建精准解析任务（模型: {self.options.model_version}）…"
                        )
                        with temporary_mineru_upload_file(upload_part.path, self.log) as (
                            upload_path,
                            upload_name,
                            used_short_copy,
                        ):
                            batch_id, upload_url, data_id = submit_precise_file(
                                upload_path,
                                self.options,
                                token,
                            )
                            self.log(f"{prefix} 已创建任务（ID: {batch_id}）")
                            if self.isInterruptionRequested():
                                raise MinerUError("用户已停止解析。")

                            self.log(f"{prefix}：正在上传文件至云端…")
                            http_put_file(upload_url, upload_path, log=self.log)
                            self.progress_signal.emit(
                                min(75, 10 + int(((part_index - 0.55) / total_parts) * 65))
                            )
                            if self.isInterruptionRequested():
                                raise MinerUError("用户已停止解析。")

                            self.log(f"{prefix}：等待云端解析完成…")
                            result_item = poll_precise_result(
                                batch_id,
                                self.options,
                                token,
                                self.log,
                            )
                            zip_url = (
                                result_item.get("full_zip_url")
                                or result_item.get("zip_url")
                                or result_item.get("result_url")
                            )
                            upload_meta = {
                                "mineru_upload_name": upload_name,
                                "mineru_upload_file": str(upload_path),
                                "mineru_data_id": data_id,
                                "used_short_upload_copy": used_short_copy,
                            }

                        part_output_dir = part_work_root / f"part_{part_index:03d}"
                        part_output_dir.mkdir(parents=True, exist_ok=True)
                        markdown, zip_url, extract_dir = extract_markdown_from_zip(
                            result_item,
                            part_output_dir,
                            self.log,
                        )
                        part_results.append(
                            {
                                "index": part_index,
                                "start_page": upload_part.start_page,
                                "end_page": upload_part.end_page,
                                "markdown": markdown,
                                "extract_dir": extract_dir,
                            }
                        )
                        part_metadata.append(
                            {
                                "index": part_index,
                                "start_page": upload_part.start_page,
                                "end_page": upload_part.end_page,
                                "page_count": upload_part.page_count,
                                "batch_id": batch_id,
                                "zip_url": zip_url,
                                **upload_meta,
                                "result_item": result_item,
                            }
                        )
                        self.progress_signal.emit(
                            min(78, 10 + int((part_index / total_parts) * 68))
                        )

                self.log(mineru_result_organization_message(total_parts))
                markdown, cleaned, image_records, extract_dir = merge_mineru_part_results(
                    part_results,
                    self.output_dir,
                    self.log,
                )
            finally:
                shutil.rmtree(part_work_root, ignore_errors=True)

            self.progress_signal.emit(80)

            raw_path = self.output_dir / "full.md"
            clean_path = self.output_dir / "full.cleaned.md"
            map_path = self.output_dir / "image_map.json"
            raw_path.write_text(markdown, encoding="utf-8")
            clean_path.write_text(cleaned, encoding="utf-8")
            map_path.write_text(json.dumps(image_records, ensure_ascii=False, indent=2), encoding="utf-8")
            primary_part = part_metadata[0]
            meta_path.write_text(
                json.dumps(
                    {
                        "batch_id": primary_part["batch_id"],
                        "model_version": self.options.model_version,
                        "zip_url": primary_part["zip_url"],
                        "source_pdf": str(self.pdf_path),
                        "source_file": str(original_copy),
                        "source_display_name": self.pdf_path.name,
                        "extract_dir": str(extract_dir),
                        "auto_split": len(part_metadata) > 1,
                        "page_limit": MINERU_UPLOAD_PAGE_LIMIT,
                        "total_pages": sum(int(part["page_count"]) for part in part_metadata),
                        "parts": part_metadata,
                        "result_item": primary_part["result_item"],
                    },
                    ensure_ascii=False,
                    indent=2,
                ),
                encoding="utf-8",
            )

            render_preview_html(clean_path, self.log)
            self.progress_signal.emit(100)
            self.finished_signal.emit(True, f"处理完成: {clean_path}", str(clean_path))
        except Exception as exc:
            if self.output_dir.exists() and not (self.output_dir / "full.cleaned.md").exists():
                try:
                    shutil.rmtree(self.output_dir)
                    self.log(f"已清理失败输出目录: {self.output_dir}")
                except Exception as cleanup_exc:
                    self.log(f"清理失败输出目录失败: {cleanup_exc}")
            self.finished_signal.emit(False, str(exc), "")


class PreviewRenderWorker(QThread):
    finished_signal = Signal(int, str, str, str)

    def __init__(
        self,
        generation: int,
        source_path: str,
        parsed_markdown: str,
        mode_name: str,
        prefer_layout: bool = False,
    ):
        super().__init__()
        self.generation = generation
        self.source_path = Path(source_path) if source_path else None
        self.parsed_markdown = Path(parsed_markdown) if parsed_markdown else None
        self.mode = PreviewMode(mode_name)
        self.prefer_layout = prefer_layout
        self.preview_provider = preview_tools.SourcePreviewProvider(WORKSPACE)

    def run(self):
        try:
            if not self.source_path:
                self.finished_signal.emit(self.generation, "html", simple_file_html(Path("missing")), "")
                return
            kind, payload = self.preview_provider.source_url_or_html(
                self.source_path,
                self.parsed_markdown,
                self.mode,
                prefer_layout=self.prefer_layout,
            )
            self.finished_signal.emit(self.generation, kind, str(payload), "")
        except Exception as exc:
            self.finished_signal.emit(self.generation, "html", simple_file_html(Path("missing")), str(exc))


class LayoutTranslateWorker(QThread):
    log_signal = Signal(str)
    reasoning_signal = Signal(str)
    preview_signal = Signal(str)
    finished_signal = Signal(bool, str, str)
    edge_download_signal = Signal(object)

    def __init__(
        self,
        markdown_path: str,
        ai_config: AITranslateConfig,
        target_language: str,
        source_language: str = "英文",
        local_machine_parallelism: int = machine_translate.MTRAN_SERVER_DEFAULT_PARALLELISM,
        request_concurrency: int = 0,
        reference_paths: list[str] | None = None,
        translation_mode: str = "full_context",
        force: bool = False,
    ):
        super().__init__()
        self.markdown_path = Path(markdown_path)
        self.ai_config = ai_config
        self.target_language = target_language
        self.source_language = source_language
        self.local_machine_parallelism = machine_translate.normalize_parallelism(
            local_machine_parallelism
        )
        self.request_concurrency = normalize_translation_request_concurrency(
            self.ai_config.provider_id,
            request_concurrency,
        )
        self.reference_paths = list(reference_paths or [])
        self.translation_mode = (
            "chunked"
            if str(translation_mode or "").strip().lower() in {"chunked", "chunks"}
            else "full_context"
        )
        self.force = force

    def log(self, message: str):
        self.log_signal.emit(message)

    def request_edge_model_download(self, source_language: str, target_language: str) -> bool:
        request = {"source": source_language, "target": target_language, "approved": False, "event": threading.Event()}
        self.edge_download_signal.emit(request)
        request["event"].wait()
        return bool(request["approved"])

    def preview(self, message: str):
        if self.isInterruptionRequested():
            raise MinerUError("用户已停止翻译。")
        self.preview_signal.emit(message)

    def reasoning(self, text: str):
        if text:
            self.reasoning_signal.emit(text)

    def run(self):
        cleanup_targets: list[Path] = []
        preexisting_cleanup_targets: set[Path] = set()
        work_dir_snapshot: set[Path] = set()
        work_dir: Path | None = None
        try:
            _translation_cancel_state.should_stop = self.isInterruptionRequested
            source_config = self.ai_config
            config = AITranslateConfig(
                provider_id=source_config.provider_id,
                api_key=source_config.api_key,
                base_url=source_config.base_url,
                model=source_config.model,
                prompt_cache_key=make_translation_cache_key(
                    translation_cache_document_identity(str(self.markdown_path.resolve()), source_config),
                    source_config.model,
                    self.target_language,
                ),
                request_body_mode=source_config.request_body_mode,
                show_reasoning=source_config.show_reasoning,
                thinking_mode=source_config.thinking_mode,
                reasoning_effort=source_config.reasoning_effort,
                # The layout worker rebuilds the request config to attach a
                # document-specific cache key.  Keep this execution-mode
                # flag as well; otherwise the UI can correctly announce the
                # fast mode while the worker silently falls back to the
                # ordinary full-context request.
                deepseek_fast_layout_translation=source_config.deepseek_fast_layout_translation,
                custom_translation_instruction=source_config.custom_translation_instruction,
            )
            # PyInstaller stores Python modules in its archive, so the
            # translation helper must be imported by module name instead of
            # loading a source .py file from the installation directory.
            import layout_translate_preview as module

            bundle = load_layout_preview_bundle(self.markdown_path)
            if not bundle:
                raise MinerUError("当前文档缺少 MinerU layout.json，无法进行排版翻译。请重新解析该文档。")
            source_layout_path = render_layout_preview_html(self.markdown_path, log=self.log)
            if not source_layout_path:
                raise MinerUError("无法生成排版预览，排版翻译已停止。")

            translated_bundle = copy.deepcopy(bundle)
            records = module.iter_translatable_blocks(translated_bundle["page_info"])
            if not records:
                raise MinerUError("没有找到可翻译的版面文本块。")
            if self.isInterruptionRequested():
                raise MinerUError("用户已停止翻译。")
            if getattr(config, "deepseek_fast_layout_translation", False):
                self.log(
                    "DeepSeek 高速并发翻译后台配置已确认：将以无思考、完整 Markdown 前缀和缓存探针执行。"
                )
            self.log(f"已提取 {len(records)} 个版面文本块。")
            if not is_free_machine_translation_config(config):
                self.log("独立公式保持原样，行内公式与上下文段落协同翻译。")
            if is_free_machine_translation_config(config):
                self.preview(f"正在生成排版译文...\n\n- 已提取版面文本块: {len(records)}\n- 正在使用免费机翻逐块翻译。")
            else:
                self.preview(f"正在生成排版译文...\n\n- 已提取版面文本块: {len(records)}\n- 正在请求翻译模型。")

            language_suffix = translation_language_suffix(self.target_language)
            if is_free_machine_translation_config(config):
                machine_suffix = "local-machine" if config.provider_id == machine_translate.MTRAN_SERVER_PROVIDER else "free-machine"
                cache_path = self.markdown_path.parent / f"layout_translation_blocks.{language_suffix}.{machine_suffix}.json"
            elif self.reference_paths:
                reference_suffix = reference_context_cache_key(self.reference_paths, self.target_language, config.model)[:12]
                cache_path = self.markdown_path.parent / f"layout_translation_blocks.{language_suffix}.ref-{reference_suffix}.json"
            else:
                cache_path = self.markdown_path.parent / f"layout_translation_blocks.{language_suffix}.json"
            cleanup_targets.append(cache_path)
            out_path = layout_translation_preview_html_path(self.markdown_path)
            cleanup_targets.append(out_path)
            bundle_path = layout_translation_bundle_path(self.markdown_path)
            cleanup_targets.append(bundle_path)
            preexisting_cleanup_targets = {path.resolve() for path in cleanup_targets if path.exists()}
            if self.force and cache_path.exists():
                cache_path.unlink()
                preexisting_cleanup_targets.discard(cache_path.resolve())
            reference_context = ""
            if is_free_machine_translation_config(config):
                if self.reference_paths:
                    self.log("免费机翻不使用参考文件上下文，已忽略参考文件。")
                translations = machine_translate.translate_record_texts(
                    records,
                    self.target_language,
                    provider_id=config.provider_id,
                    source_language=self.source_language,
                    base_url=config.base_url,
                    api_key=config.api_key,
                    parallelism=self.local_machine_parallelism,
                    log=self.log,
                    live_update=self.preview,
                    should_stop=self.isInterruptionRequested,
                    edge_download_consent=self.request_edge_model_download,
                )
                translations = module.repair_record_translations(records, translations)
                if self.isInterruptionRequested():
                    raise MinerUError("用户已停止翻译。")
                payload = {
                    "target_language": self.target_language,
                    "source_language": self.source_language,
                    "model": config.model,
                    "local_machine_parallelism": self.local_machine_parallelism,
                    "translations": [{"id": record.block_id, "text": translations.get(record.block_id, record.text)} for record in records],
                    "machine_translation": True,
                }
                cache_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
            elif self.reference_paths:
                reference_cache_root = current_work_dir() / "_reference_corpus_cache"
                reference_context = build_reference_translation_context(
                    self.reference_paths,
                    config,
                    reference_cache_root,
                    self.log,
                    self.target_language,
                    self.isInterruptionRequested,
                )
            if not is_free_machine_translation_config(config):
                self.log("正在等待模型翻译完成；深度思考期间请耐心等待，可随时手动停止。")
                translations = module.translate_records(
                    records,
                    config,
                    self.target_language,
                    cache_path,
                    max_chars=module.LAYOUT_GROUP_MAX_CHARS,
                    max_blocks=module.LAYOUT_GROUP_MAX_BLOCKS,
                    concurrency=self.request_concurrency,
                    translation_mode=self.translation_mode,
                    reference_context=reference_context,
                    formula_context=None,
                    log=self.log,
                    preview_callback=self.preview,
                    reasoning_callback=self.reasoning,
                    should_stop=self.isInterruptionRequested,
                    full_markdown_context=(
                        self.markdown_path.read_text(encoding="utf-8", errors="replace")
                        if getattr(config, "deepseek_fast_layout_translation", False)
                        else ""
                    ),
                )
            if self.isInterruptionRequested():
                raise MinerUError("用户已停止翻译。")
            module.apply_translations(records, translations)
            # 将本次翻译配置写入排版 bundle。批量候选判断必须比较配置身份，
            # 不能只看固定的“current”HTML 是否存在。
            translated_bundle["_translation_job_identity"] = translation_job_identity(
                TranslationJobConfig(
                    ai_config=self.ai_config,
                    source_language=self.source_language,
                    target_language=self.target_language,
                    mode=self.translation_mode,
                    reference_paths=list(self.reference_paths),
                    local_machine_parallelism=self.local_machine_parallelism,
                    request_concurrency=self.request_concurrency,
                )
            )
            # Keep old bundles from being reused when the block taxonomy or
            # fitting protocol changes.  In particular, image footnotes now
            # need a fresh translation and a dynamic fit pass.
            translated_bundle["_layout_translation_artifact_protocol"] = LAYOUT_TRANSLATION_ARTIFACT_PROTOCOL
            tmp_out_path = out_path.with_name(f".{out_path.name}.tmp")
            tmp_bundle_path = bundle_path.with_name(f".{bundle_path.name}.tmp")
            cleanup_targets.append(tmp_out_path)
            cleanup_targets.append(tmp_bundle_path)
            module.render_translated_layout(
                self.markdown_path,
                translated_bundle,
                source_layout_path,
                tmp_out_path,
                debug_overlay=False,
                reset_fit_cache=True,
                bundle_out_path=tmp_bundle_path,
            )
            publish_layout_translation_artifacts(
                tmp_out_path,
                out_path,
                tmp_bundle_path,
                bundle_path,
            )
            self.finished_signal.emit(True, f"排版翻译已完成：{out_path}", str(out_path))
        except Exception as exc:
            # A valid partial layout cache is a resumable checkpoint, not a
            # failed publication artifact. Preserve it across cancellation or
            # a provider error; temporary HTML/bundle outputs are still removed.
            if "cache_path" in locals() and cache_path.exists():
                preexisting_cleanup_targets.add(cache_path.resolve())
            remove_paths_quietly(cleanup_targets, self.log, preexisting_cleanup_targets)
            self.finished_signal.emit(False, str(exc), "")


class TranslateWorker(QThread):
    log_signal = Signal(str)
    reasoning_signal = Signal(str)
    preview_signal = Signal(str)
    # `preview_signal` remains for compatibility with callers outside the main
    # window.  The UI uses this delta form so a long translation does not copy
    # its complete prefix across the Qt queued connection on every update.
    preview_delta_signal = Signal(str, bool)
    finished_signal = Signal(bool, str, str)
    edge_download_signal = Signal(object)

    def __init__(self, markdown_path: str, job_config: TranslationJobConfig):
        super().__init__()
        self.markdown_path = Path(markdown_path)
        self.job_config = job_config
        self._last_preview_emit = 0.0
        self._last_preview_markdown = ""
        self._has_preview_markdown = False

    def log(self, message: str):
        self.log_signal.emit(message)

    def request_edge_model_download(self, source_language: str, target_language: str) -> bool:
        request = {"source": source_language, "target": target_language, "approved": False, "event": threading.Event()}
        self.edge_download_signal.emit(request)
        request["event"].wait()
        return bool(request["approved"])

    def preview(self, markdown: str, force: bool = False):
        if self.isInterruptionRequested():
            raise MinerUError("用户已停止翻译。")
        now = time.monotonic()
        if force or now - self._last_preview_emit >= 0.08:
            markdown = str(markdown or "")
            reset = (not self._has_preview_markdown) or not markdown.startswith(self._last_preview_markdown)
            delta = markdown if reset else markdown[len(self._last_preview_markdown):]
            if delta or reset:
                self.preview_delta_signal.emit(delta, reset)
            # Preserve the established signal for batch/external callers. The
            # interactive reader is connected only to preview_delta_signal.
            self.preview_signal.emit(markdown)
            self._last_preview_markdown = markdown
            self._has_preview_markdown = True
            self._last_preview_emit = now

    def reasoning(self, text: str):
        if text:
            self.reasoning_signal.emit(text)

    def run(self):
        cleanup_targets: list[Path] = []
        preexisting_cleanup_targets: set[Path] = set()
        work_dir = None
        work_dir_snapshot: set[Path] = set()
        try:
            _translation_cancel_state.should_stop = self.isInterruptionRequested
            if self.isInterruptionRequested():
                raise MinerUError("用户已停止翻译。")
            if not self.markdown_path.exists():
                raise MinerUError("当前文档不存在，无法翻译。")
            source_config = self.job_config.ai_config
            config = AITranslateConfig(
                provider_id=source_config.provider_id,
                api_key=source_config.api_key,
                base_url=source_config.base_url,
                model=source_config.model,
                prompt_cache_key=make_translation_cache_key(
                    translation_cache_document_identity(str(self.markdown_path.resolve()), source_config),
                    source_config.model,
                    self.job_config.target_language,
                ),
                request_body_mode=source_config.request_body_mode,
                show_reasoning=source_config.show_reasoning,
                thinking_mode=source_config.thinking_mode,
                reasoning_effort=source_config.reasoning_effort,
                custom_translation_instruction=source_config.custom_translation_instruction,
            )
            if is_free_machine_translation_config(config):
                self.log(f"已选择翻译服务：{translation_provider_label(config.provider_id)}")
            else:
                self.log(f"已选择翻译模型：{config.model}")
            self.preview("正在连接翻译服务，译文会在这里实时显示...\n", force=True)
            markdown = self.markdown_path.read_text(encoding="utf-8", errors="replace")
            work_dir = translation_work_dir_for_job(self.markdown_path, self.job_config)
            work_dir_snapshot = snapshot_existing_paths(work_dir)
            cleanup_targets.append(work_dir)
            preexisting_cleanup_targets.update(path.resolve() for path in cleanup_targets if path.exists())
            work_dir.mkdir(parents=True, exist_ok=True)
            mode_flag = work_dir / "full_context.enabled"
            if self.job_config.mode == "full_context":
                mode_flag.write_text("1", encoding="utf-8")
            elif mode_flag.exists():
                mode_flag.unlink()

            reference_context = ""
            if is_free_machine_translation_config(config):
                if self.job_config.reference_paths:
                    self.log("免费机翻不使用参考文件上下文，已忽略参考文件。")
                translated = machine_translate.translate_markdown_document(
                    markdown,
                    self.job_config.target_language,
                    provider_id=config.provider_id,
                    source_language=self.job_config.source_language,
                    base_url=config.base_url,
                    api_key=config.api_key,
                    parallelism=self.job_config.local_machine_parallelism,
                    log=self.log,
                    live_update=self.preview,
                    should_stop=self.isInterruptionRequested,
                    edge_download_consent=self.request_edge_model_download,
                )
            elif self.job_config.reference_paths:
                self.log("正在解析参考文件并构建专业语料库…")
                # 参考语料缓存使用工作目录下的共享缓存，而不是每篇文档自己的 translation_chunks。
                # 这样批量翻译、多次翻译、分块翻译都能复用同一批参考文献解析结果和语料包。
                reference_cache_root = current_work_dir() / "_reference_corpus_cache"
                reference_context = build_reference_translation_context(
                    self.job_config.reference_paths,
                    config,
                    reference_cache_root,
                    self.log,
                    self.job_config.target_language,
                    self.isInterruptionRequested,
                )
                if reference_context.strip():
                    self.log("参考语料已加入翻译上下文（原文内容保持最高优先级）。")

            if not is_free_machine_translation_config(config):
                translated = translate_markdown_text(
                    markdown,
                    config,
                    self.log,
                    self.job_config.target_language,
                    work_dir,
                    self.preview,
                    self.reasoning,
                    reference_context,
                    self.job_config.request_concurrency,
                )
            # EPUB chapter comments are protected structural boundaries. Some
            # providers omit HTML comments even when preserving all visible text.
            translated = restore_epub_chapter_markers(markdown, translated)
            if is_epub_markdown(markdown):
                translated = repair_epub_markdown_attributes(translated)
            translated = normalize_translated_inline_html(translated)
            if self.isInterruptionRequested():
                raise MinerUError("用户已停止翻译。")
            self.preview(translated, force=True)
            out_path = translation_output_path_for_job(self.markdown_path, self.job_config)
            cleanup_targets.append(out_path)
            metadata_path = translation_job_metadata_path(out_path)
            cleanup_targets.append(metadata_path)
            safe_stem = re.sub(r"[^A-Za-z0-9_.-]+", "_", out_path.stem)
            cleanup_targets.append(out_path.with_name(f"preview.{safe_stem}.html"))
            preexisting_cleanup_targets.update(path.resolve() for path in cleanup_targets if path.exists())
            tmp_out_path = out_path.with_name(f".{out_path.name}.tmp")
            cleanup_targets.append(tmp_out_path)
            tmp_out_path.write_text(translated, encoding="utf-8")
            tmp_out_path.replace(out_path)
            write_translation_job_metadata(out_path, self.job_config)
            render_preview_html(out_path, self.log)
            self.finished_signal.emit(True, f"全文翻译已完成：{out_path}", str(out_path))
        except Exception as exc:
            remove_paths_quietly(cleanup_targets, self.log, preexisting_cleanup_targets)
            if work_dir is not None:
                remove_new_paths_under(work_dir, work_dir_snapshot, self.log)
            self.finished_signal.emit(False, str(exc), "")


__all__ = [name for name in globals() if not name.startswith("__")]
